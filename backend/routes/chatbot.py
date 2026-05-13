"""
Chatbot Routes - AI Assistant for the platform
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from typing import Optional, List
from datetime import datetime, timedelta, timezone
import os
import io
import base64
import logging
from pathlib import Path

from motor.motor_asyncio import AsyncIOMotorClient
from dotenv import load_dotenv
from pydantic import BaseModel

# Load environment
ROOT_DIR = Path(__file__).parent.parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Security
import jwt
JWT_SECRET = os.environ.get('JWT_SECRET', 'fleet-maintenance-secret-key-2024')

security = HTTPBearer()

def decode_token(token: str) -> dict:
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        return payload
    except:
        raise HTTPException(status_code=401, detail="Token inválido")

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    payload = decode_token(credentials.credentials)
    # Normaliza: token usa "user_id"; alguns trechos esperam "id"
    if "id" not in payload and "user_id" in payload:
        payload["id"] = payload["user_id"]
    return payload

# Create router
chatbot_router = APIRouter(prefix="/chatbot", tags=["Chatbot"])


# Models
class ChatMessage(BaseModel):
    message: str
    module: str = "gerenciamento"


class ChatResponse(BaseModel):
    response: str
    context_used: List[str] = []


# ============ BASE DE CONHECIMENTO RH ============
# Documentos normativos (PCMSO, PGR, LTCAT, CCT) ficam permanentemente acessíveis
# ao Chat IA do RH para responder sobre exames, EPIs, riscos, pisos, jornadas, etc.

_KB_CACHE = {"loaded_at": None, "context_text": ""}
_KB_TTL = timedelta(minutes=10)


async def _build_knowledge_base_context() -> str:
    """Concatena os textos extraídos dos documentos normativos para o system_prompt.
    Cacheia em memória por 10 minutos."""
    now = datetime.now(timezone.utc)
    if (
        _KB_CACHE["loaded_at"]
        and (now - _KB_CACHE["loaded_at"]) < _KB_TTL
        and _KB_CACHE["context_text"]
    ):
        return _KB_CACHE["context_text"]

    docs = await db.chat_knowledge_base.find(
        {"category": "rh_normativos"}, {"_id": 0}
    ).sort("name", 1).to_list(20)

    if not docs:
        _KB_CACHE.update({"loaded_at": now, "context_text": ""})
        return ""

    parts = ["=" * 60, "DOCUMENTOS NORMATIVOS DE RH (consulte SEMPRE para perguntas sobre",
             "exames ocupacionais, EPIs, riscos, pisos salariais, jornadas, benefícios CCT)",
             "=" * 60]
    for d in docs:
        title = d.get("title") or d.get("name", "")
        text = d.get("extracted_text") or ""
        parts.append(f"\n\n>>> DOCUMENTO: {d.get('name','?')} — {title}")
        parts.append(f">>> ({d.get('pages', '?')} páginas, atualizado em "
                     f"{d.get('created_at','')[:10]})")
        parts.append("-" * 60)
        parts.append(text)
        parts.append("-" * 60)

    final = "\n".join(parts)
    _KB_CACHE.update({"loaded_at": now, "context_text": final})
    return final


def _invalidate_kb_cache():
    _KB_CACHE.update({"loaded_at": None, "context_text": ""})


# ========================================================================
# UPLOAD DE ANEXOS NO CHAT — Suporta PDF, imagem, Excel/CSV, Word, texto
# ========================================================================

CHAT_UPLOADS_DIR = Path("/app/backend/uploads/chat_attachments")
CHAT_UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

# MIMEs que o Gemini lê nativamente (passar direto via FileContentWithMimeType)
GEMINI_NATIVE_MIMES = {
    "application/pdf",
    "image/jpeg", "image/png", "image/webp", "image/heic", "image/heif",
    "text/plain", "text/csv", "text/html", "text/markdown",
    "audio/mpeg", "audio/wav", "audio/aac", "audio/ogg", "audio/flac",
    "video/mp4", "video/mpeg", "video/mov", "video/avi", "video/wmv",
}


def _extrair_texto_excel(path: str, max_rows: int = 2000) -> str:
    """Converte .xlsx para texto tabular legível pela IA."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        partes = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            partes.append(f"\n═══ Planilha: {sheet_name} ═══")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                row_count += 1
                if row_count > max_rows:
                    partes.append(f"... (planilha truncada em {max_rows} linhas)")
                    break
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    partes.append(" | ".join(cells))
        wb.close()
        return "\n".join(partes)
    except Exception as e:
        return f"[Erro ao ler Excel: {e}]"


def _extrair_texto_docx(path: str) -> str:
    """Converte .docx para texto."""
    try:
        from docx import Document
        doc = Document(path)
        partes = []
        for p in doc.paragraphs:
            if p.text.strip():
                partes.append(p.text)
        for table in doc.tables:
            partes.append("\n--- Tabela ---")
            for row in table.rows:
                partes.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(partes)
    except Exception as e:
        return f"[Erro ao ler Word: {e}]"


def _extrair_texto_csv(path: str, max_rows: int = 5000) -> str:
    """Lê CSV e retorna como texto tabular."""
    import csv as _csv
    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            reader = _csv.reader(f)
            rows = []
            for i, r in enumerate(reader):
                if i >= max_rows:
                    rows.append(f"... (truncado em {max_rows} linhas)")
                    break
                rows.append(" | ".join(r))
        return "\n".join(rows)
    except Exception as e:
        return f"[Erro ao ler CSV: {e}]"


async def _processar_anexo_chat(upload_file, conv_id: str) -> dict:
    """Salva o arquivo no disco e retorna metadados + estratégia de inclusão na IA.

    Retorna dict com:
      - filename: nome original
      - mime: detectado
      - path: caminho salvo
      - inline_text: se != None, deve ser concatenado ao prompt
      - gemini_file: se != None, deve ser passado em file_contents (FileContentWithMimeType)
    """
    import mimetypes
    raw = await upload_file.read()
    safe_name = upload_file.filename.replace("/", "_").replace("\\", "_")[:120]
    file_id = uuid.uuid4().hex[:12]
    target = CHAT_UPLOADS_DIR / f"{conv_id}_{file_id}_{safe_name}"
    with open(target, "wb") as f:
        f.write(raw)

    mime = upload_file.content_type or mimetypes.guess_type(safe_name)[0] or "application/octet-stream"
    ext = (safe_name.rsplit(".", 1)[-1] or "").lower()

    info = {
        "id": file_id,
        "filename": safe_name,
        "mime": mime,
        "path": str(target),
        "size": len(raw),
        "inline_text": None,
        "gemini_file": None,
    }

    # Estratégia por tipo
    if mime in GEMINI_NATIVE_MIMES or ext in ("pdf", "jpg", "jpeg", "png", "webp", "txt", "md"):
        info["gemini_file"] = {"path": str(target), "mime": mime if mime in GEMINI_NATIVE_MIMES else "application/pdf"}
        if ext == "txt" or ext == "md":
            try:
                with open(target, "r", encoding="utf-8", errors="replace") as f:
                    info["inline_text"] = f.read()[:50000]
                info["gemini_file"] = None  # texto inline é mais barato
            except Exception:
                pass
    elif ext in ("xlsx", "xlsm", "xls"):
        info["inline_text"] = _extrair_texto_excel(str(target))
    elif ext == "csv":
        info["inline_text"] = _extrair_texto_csv(str(target))
    elif ext in ("docx", "doc"):
        info["inline_text"] = _extrair_texto_docx(str(target))
    else:
        # Fallback: tenta ler como texto puro
        try:
            with open(target, "r", encoding="utf-8", errors="replace") as f:
                info["inline_text"] = f.read()[:30000]
        except Exception:
            info["inline_text"] = f"[Arquivo não suportado para leitura: {ext}]"

    return info


def _detectar_pedido_alteracao(texto: str) -> bool:
    """Heurística simples: detecta se o usuário pediu alteração do arquivo."""
    if not texto:
        return False
    t = texto.lower()
    palavras = [
        "altere", "alterar", "modifique", "modificar", "atualize", "atualizar",
        "edite", "editar", "corrija", "corrigir", "ajuste", "ajustar",
        "remova", "remover", "exclua", "excluir", "delete",
        "adicione", "adicionar", "insira", "inserir",
        "substitua", "substituir", "troque", "trocar",
        "gere uma versão", "gere o arquivo", "gerar novo arquivo",
        "salve as alterações", "salve em", "gere um novo", "exporte",
    ]
    return any(p in t for p in palavras)


# ============ AUTO-BOOTSTRAP DA BASE DE CONHECIMENTO ============
# URLs fixas dos 4 documentos normativos da CRA Construtora.
# Na primeira inicialização de qualquer ambiente (preview, produção, staging),
# os PDFs são baixados, o texto é extraído (com OCR via Gemini para escaneados)
# e gravados no Mongo. Idempotente: se um documento já existe na coleção, pula.
_KB_BOOTSTRAP_URLS = [
    {
        "name": "PCMSO",
        "title": "PCMSO - Programa de Controle Médico de Saúde Ocupacional (CRA Apoio Administrativo)",
        "url": "https://customer-assets.emergentagent.com/job_21e279ba-21c4-411a-94e8-db609ecbdb3a/artifacts/5u1qpo4w_PCMSO-%20CRA%20APOIO%20ADMINISTRATIVO.pdf",
    },
    {
        "name": "PGR",
        "title": "PGR - Programa de Gerenciamento de Riscos (CRA Apoio Administrativo)",
        "url": "https://customer-assets.emergentagent.com/job_21e279ba-21c4-411a-94e8-db609ecbdb3a/artifacts/5jitxhuv_PGR-CRA%20-APOIO%20ADMINISTRATIVO.pdf",
    },
    {
        "name": "LTCAT",
        "title": "LTCAT - Laudo Técnico de Condições Ambientais (CRA Apoio Administrativo)",
        "url": "https://customer-assets.emergentagent.com/job_21e279ba-21c4-411a-94e8-db609ecbdb3a/artifacts/qhcqhqew_LTCAT-APOIO%20ADMINISTRATIVO%20%281%29.pdf",
    },
    {
        "name": "CCT",
        "title": "Convenção Coletiva de Trabalho 2025/2026 - Construção Pesada TO",
        "url": "https://customer-assets.emergentagent.com/job_21e279ba-21c4-411a-94e8-db609ecbdb3a/artifacts/6xw3fvgt_CONVEN%C3%87%C3%83O%20COLETIVA%20%281%29.pdf",
    },
]


async def bootstrap_knowledge_base() -> dict:
    """Garante que os 4 documentos normativos padrão estão carregados na coleção.
    Idempotente: cada documento só é baixado/processado se ainda não existir."""
    import httpx
    storage_dir = "/app/backend/storage/rh_normativos"
    os.makedirs(storage_dir, exist_ok=True)

    summary = {"already_present": [], "added": [], "errors": []}

    for spec in _KB_BOOTSTRAP_URLS:
        name = spec["name"]
        try:
            existing = await db.chat_knowledge_base.find_one(
                {"category": "rh_normativos", "name": name}, {"_id": 0, "extracted_text": 0}
            )
            if existing:
                summary["already_present"].append(name)
                continue

            logging.info(f"[KB Bootstrap] Baixando {name} ...")
            async with httpx.AsyncClient(timeout=60.0) as cli:
                resp = await cli.get(spec["url"])
                resp.raise_for_status()
                pdf_bytes = resp.content

            dst = os.path.join(storage_dir, f"{name}.pdf")
            with open(dst, "wb") as f:
                f.write(pdf_bytes)

            # Extração de texto
            pages = 0
            text = ""
            try:
                from pypdf import PdfReader
                r = PdfReader(dst)
                pages = len(r.pages)
                text = "\n".join((p.extract_text() or "") for p in r.pages)
            except Exception as e:
                logging.warning(f"[KB Bootstrap] PyPDF falhou para {name}: {e}")

            # Fallback OCR via Gemini se texto for pobre (ex: PDF escaneado como CCT)
            if len(text.strip()) < 100:
                try:
                    from emergentintegrations.llm.chat import (
                        LlmChat, UserMessage, FileContentWithMimeType,
                    )
                    chat = LlmChat(
                        api_key=os.environ["EMERGENT_LLM_KEY"],
                        session_id=f"kb-bootstrap-{name}-{uuid.uuid4().hex[:6]}",
                        system_message="Extrator de texto de PDFs.",
                    ).with_model("gemini", "gemini-2.5-flash")
                    fc = FileContentWithMimeType(file_path=dst, mime_type="application/pdf")
                    msg = UserMessage(
                        text=f"Extraia TODO o texto deste documento ({spec['title']}). "
                             "Mantenha cláusulas, títulos, tabelas, valores e datas.",
                        file_contents=[fc],
                    )
                    text = await chat.send_message(msg)
                    logging.info(f"[KB Bootstrap] OCR Gemini extraiu {len(text)} chars para {name}")
                except Exception as ocr_err:
                    logging.error(f"[KB Bootstrap] OCR falhou para {name}: {ocr_err}")

            await db.chat_knowledge_base.insert_one({
                "id": str(uuid.uuid4()),
                "category": "rh_normativos",
                "name": name,
                "title": spec["title"],
                "extracted_text": text,
                "pdf_path": dst,
                "pdf_size": os.path.getsize(dst),
                "pages": pages,
                "created_at": datetime.now(timezone.utc).isoformat(),
                "uploaded_by": "system_bootstrap",
            })
            summary["added"].append(f"{name} ({len(text)} chars, {pages}p)")
            logging.info(f"[KB Bootstrap] {name} inserido com sucesso.")
        except Exception as e:
            logging.error(f"[KB Bootstrap] Falha em {name}: {e}")
            summary["errors"].append(f"{name}: {e}")

    if summary["added"]:
        _invalidate_kb_cache()
    return summary


@chatbot_router.post("/knowledge-base/bootstrap")
async def trigger_bootstrap_knowledge_base(current_user: dict = Depends(get_current_user)):
    """Endpoint manual para forçar o carregamento dos 4 documentos padrão.
    Útil se o startup falhou ou os documentos foram removidos."""
    summary = await bootstrap_knowledge_base()
    return summary




async def get_full_platform_context() -> str:
    """Coleta TODAS as informações de TODAS as coleções do banco de dados.
    Inclui agregações úteis para Ponto (faltas/atrasos), Folha, NF-e, OS, etc."""
    from collections import defaultdict
    context_parts = []
    hoje = datetime.now()
    inicio_30 = (hoje - timedelta(days=30)).strftime("%Y-%m-%d")
    inicio_60 = (hoje - timedelta(days=60)).strftime("%Y-%m-%d")
    mes_passado_dt = hoje.replace(day=1) - timedelta(days=1)
    mes_passado_inicio = mes_passado_dt.replace(day=1).strftime("%Y-%m-%d")
    mes_passado_fim = mes_passado_dt.strftime("%Y-%m-%d")

    context_parts.append("=" * 60)
    context_parts.append(f"BANCO DE DADOS COMPLETO - CRA CONSTRUTORA  |  Hoje: {hoje.strftime('%d/%m/%Y')}")
    context_parts.append("=" * 60)

    # ============ USUÁRIOS ============
    users = await db.users.find({}, {"_id": 0, "password": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nUSUÁRIOS ({len(users)})\n{'='*40}")
    for u in users:
        context_parts.append(f"- {u.get('name')} | {u.get('email')} | role={u.get('role','gerenciamento')}")

    # ============ MÁQUINAS / FROTA ============
    categories = await db.categories.find({}, {"_id": 0}).to_list(200)
    machines = await db.machines.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nFROTA ({len(machines)} máquinas, {len(categories)} categorias)\n{'='*40}")
    for m in machines:
        context_parts.append(
            f"- {m.get('name')} | placa={m.get('plate','-')} | marca={m.get('brand','-')} | status={m.get('status','-')}"
        )

    # ============ MANUTENÇÕES ============
    maintenances = await db.maintenances.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    total_valor_manut = sum(m.get('part_value', 0) for m in maintenances)
    context_parts.append(f"\n\n{'='*40}\nMANUTENÇÕES ({len(maintenances)})\n{'='*40}")
    context_parts.append(f"TOTAL GASTO: R$ {total_valor_manut:,.2f}")
    for m in maintenances[:30]:
        context_parts.append(
            f"- {m.get('part_name','-')} | R$ {m.get('part_value',0):,.2f} | data={m.get('replacement_date','-')} | máquina={m.get('machine_id','-')}"
        )

    # ============ ESTOQUE ============
    stock_items = await db.stock_items.find({}, {"_id": 0}).to_list(500)
    low_stock = [i for i in stock_items if i.get("quantity", 0) <= i.get("min_quantity", 0)]
    context_parts.append(f"\n\n{'='*40}\nESTOQUE ({len(stock_items)} itens, {len(low_stock)} em alerta)\n{'='*40}")
    for i in stock_items:
        context_parts.append(f"- {i.get('name')} | qtd={i.get('quantity',0)} | min={i.get('min_quantity',0)}")

    # ============ OBRAS ============
    obras = await db.obras.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nOBRAS ({len(obras)})\n{'='*40}")
    for o in obras:
        context_parts.append(f"- {o.get('name','-')} | {o.get('location','-')} | status={o.get('status','-')}")

    # ============ CADASTROS (clientes/fornecedores) ============
    cadastros = await db.cadastros.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nCADASTROS ({len(cadastros)})\n{'='*40}")
    for c in cadastros[:80]:
        context_parts.append(f"- {c.get('tipo_cadastro','-')}: {c.get('nome_razao','-')} | CNPJ/CPF={c.get('cnpj_cpf','-')}")

    # ============ FINANCEIRO ============
    contas_pagar = await db.contas_pagar.find({}, {"_id": 0}).to_list(2000)
    contas_receber = await db.contas_receber.find({}, {"_id": 0}).to_list(2000)
    pagar_aberto = [c for c in contas_pagar if c.get("status") in ("em_aberto", "pendente", "parcial")]
    receber_aberto = [c for c in contas_receber if c.get("status") in ("em_aberto", "pendente", "parcial")]
    pagar_vencidas = [c for c in pagar_aberto if (c.get("data_vencimento") or "9999") < hoje.strftime("%Y-%m-%d")]
    receber_vencidas = [c for c in receber_aberto if (c.get("data_vencimento") or "9999") < hoje.strftime("%Y-%m-%d")]
    context_parts.append(f"\n\n{'='*40}\nCONTAS A PAGAR ({len(contas_pagar)})\n{'='*40}")
    context_parts.append(f"EM ABERTO: {len(pagar_aberto)} | VENCIDAS: {len(pagar_vencidas)}")
    context_parts.append(f"TOTAL EM ABERTO: R$ {sum(c.get('valor',0) for c in pagar_aberto):,.2f}")
    for c in pagar_aberto[:40]:
        context_parts.append(
            f"- {c.get('descricao','-')[:50]} | R$ {c.get('valor',0):,.2f} | venc={c.get('data_vencimento','-')} | forma={c.get('forma_pagamento','-')} | fornecedor={c.get('cadastro_nome','-')}"
        )
    context_parts.append(f"\n\n{'='*40}\nCONTAS A RECEBER ({len(contas_receber)})\n{'='*40}")
    context_parts.append(f"EM ABERTO: {len(receber_aberto)} | VENCIDAS: {len(receber_vencidas)}")
    context_parts.append(f"TOTAL EM ABERTO: R$ {sum(c.get('valor',0) for c in receber_aberto):,.2f}")
    for c in receber_aberto[:40]:
        context_parts.append(
            f"- {c.get('descricao','-')[:50]} | R$ {c.get('valor',0):,.2f} | venc={c.get('data_vencimento','-')} | cliente={c.get('cadastro_nome','-')}"
        )

    # ============ ORDENS DE SERVIÇO ============
    ordens_servico = await db.ordens_servico.find({}, {"_id": 0}).sort("data_emissao", -1).to_list(200)
    context_parts.append(f"\n\n{'='*40}\nORDENS DE SERVIÇO ({len(ordens_servico)})\n{'='*40}")
    for os_ in ordens_servico[:30]:
        context_parts.append(
            f"- OS #{os_.get('numero','?')} | {os_.get('cliente_nome','-')} | R$ {os_.get('valor_total',0):,.2f} | status={os_.get('status','-')} | emissão={os_.get('data_emissao','-')}"
        )

    # ============ ALUGUEIS ============
    alugueis = await db.alugueis.find({}, {"_id": 0}).to_list(200)
    context_parts.append(f"\n\n{'='*40}\nALUGUEIS ({len(alugueis)})\n{'='*40}")
    for a in alugueis[:30]:
        context_parts.append(
            f"- {a.get('descricao','-')[:50]} | R$ {a.get('valor_mensal',0):,.2f}/mês | início={a.get('data_inicio','-')} | status={a.get('status','-')}"
        )

    # ============ NF-e e NFS-e IMPORTADAS ============
    nfes = await db.nfes_importadas.find({}, {"_id": 0}).sort("data_emissao", -1).to_list(50)
    nfses = await db.nfse_importadas.find({}, {"_id": 0}).sort("data_emissao", -1).to_list(50)
    context_parts.append(f"\n\n{'='*40}\nNF-e ({len(nfes)} amostra) / NFS-e ({len(nfses)} amostra)\n{'='*40}")
    for nf in nfes[:15]:
        context_parts.append(
            f"- NF-e #{nf.get('numero_nota','?')} | {nf.get('razao_social_emitente','-')} | R$ {nf.get('valor_total',0):,.2f} | {nf.get('data_emissao','-')}"
        )
    for nf in nfses[:15]:
        context_parts.append(
            f"- NFS-e #{nf.get('numero_nota','?')} | {nf.get('razao_social_emitente','-')} | R$ {nf.get('valor_total',0):,.2f} | {nf.get('data_emissao','-')}"
        )

    # ============ FUNCIONÁRIOS RH ============
    funcionarios = await db.funcionarios.find({}, {"_id": 0}).to_list(500)
    func_id_to_nome = {f.get("id"): f.get("nome") for f in funcionarios}
    ativos = [f for f in funcionarios if (f.get("status") or "").lower() == "ativo"]
    context_parts.append(f"\n\n{'='*40}\nFUNCIONÁRIOS RH ({len(funcionarios)} total, {len(ativos)} ativos)\n{'='*40}")
    for f in funcionarios:
        context_parts.append(
            f"- ID:{f.get('id','-')} | {f.get('nome','-')} | {f.get('cargo','-')} | R$ {f.get('salario',0):,.2f} | status={f.get('status','-')}"
        )

    # ============ PONTO ELETRÔNICO — agregações por funcionário ============
    ponto_30 = await db.ponto_registros.find(
        {"data": {"$gte": inicio_30, "$lte": hoje.strftime("%Y-%m-%d")}}, {"_id": 0}
    ).to_list(5000)
    ponto_mes_passado = await db.ponto_registros.find(
        {"data": {"$gte": mes_passado_inicio, "$lte": mes_passado_fim}}, {"_id": 0}
    ).to_list(5000)
    ponto_60 = await db.ponto_registros.find(
        {"data": {"$gte": inicio_60, "$lte": hoje.strftime("%Y-%m-%d")}}, {"_id": 0}
    ).to_list(8000)
    context_parts.append(
        f"\n\n{'='*40}\nPONTO ELETRÔNICO (últ. 30 dias: {len(ponto_30)} regs / 60 dias: {len(ponto_60)})\n{'='*40}"
    )
    context_parts.append(
        f"Período mês passado considerado: {mes_passado_inicio} a {mes_passado_fim}"
    )

    def _agregar(registros, label):
        faltas = defaultdict(int)
        atrasos = defaultdict(int)
        abonados = defaultdict(int)
        trabalhados = defaultdict(int)
        for r in registros:
            fid = r.get("funcionario_id")
            if not fid:
                continue
            status_dia = (r.get("status_dia") or "").lower()
            if r.get("abono") or status_dia == "abonado":
                abonados[fid] += 1
            elif status_dia in ("sem_registro", "faltou", "ausente"):
                faltas[fid] += 1
            elif status_dia in ("atrasado", "atraso"):
                atrasos[fid] += 1
                trabalhados[fid] += 1
            elif r.get("batidas"):
                trabalhados[fid] += 1
        if faltas or atrasos or abonados:
            context_parts.append(f"\n>>> RESUMO {label}:")
            top_faltas = sorted(faltas.items(), key=lambda x: -x[1])[:10]
            for fid, qtd in top_faltas:
                context_parts.append(
                    f"- {func_id_to_nome.get(fid,'?')} → {qtd} falta(s) | {atrasos.get(fid,0)} atraso(s) | {abonados.get(fid,0)} abono(s) | {trabalhados.get(fid,0)} dia(s) trabalhado(s)"
                )
            if not top_faltas:
                # Se ninguém tem faltas, ainda mostre quem mais trabalhou
                top_t = sorted(trabalhados.items(), key=lambda x: -x[1])[:5]
                for fid, qtd in top_t:
                    context_parts.append(f"- {func_id_to_nome.get(fid,'?')} → {qtd} dia(s) trabalhado(s) | {atrasos.get(fid,0)} atraso(s)")
        else:
            context_parts.append(f"(nenhuma falta/atraso/abono registrado em {label})")

    _agregar(ponto_mes_passado, "MÊS PASSADO")
    _agregar(ponto_30, "ÚLTIMOS 30 DIAS")

    # ============ FOLHA DE PAGAMENTO ============
    folhas = await db.folha_pagamento.find({}, {"_id": 0}).sort("competencia", -1).to_list(200)
    context_parts.append(f"\n\n{'='*40}\nFOLHA DE PAGAMENTO ({len(folhas)})\n{'='*40}")
    by_comp = defaultdict(list)
    for fp in folhas:
        by_comp[fp.get("competencia", "?")].append(fp)
    for comp in list(sorted(by_comp.keys(), reverse=True))[:6]:
        items = by_comp[comp]
        total_liq = sum(i.get("liquido", 0) for i in items)
        context_parts.append(
            f"- Competência {comp}: {len(items)} funcionário(s) | total líquido R$ {total_liq:,.2f}"
        )

    # ============ FÉRIAS ============
    ferias = await db.ferias.find({}, {"_id": 0}).to_list(200)
    context_parts.append(f"\n\n{'='*40}\nFÉRIAS ({len(ferias)})\n{'='*40}")
    for fe in ferias[:20]:
        context_parts.append(
            f"- {func_id_to_nome.get(fe.get('funcionario_id'), '?')} | {fe.get('data_inicio','-')} → {fe.get('data_fim','-')} | status={fe.get('status','-')}"
        )

    # ============ EPI ============
    epi = await db.epi_fichas.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nEPI ({len(epi)} fichas)\n{'='*40}")
    for e in epi[:20]:
        context_parts.append(
            f"- {func_id_to_nome.get(e.get('funcionario_id'), '?')} | EPI={e.get('item','-')} | entrega={e.get('data_entrega','-')}"
        )

    # ============ JORNADAS ============
    jornadas = await db.jornadas.find({}, {"_id": 0}).to_list(50)
    context_parts.append(f"\n\n{'='*40}\nJORNADAS DE TRABALHO ({len(jornadas)})\n{'='*40}")
    for j in jornadas:
        context_parts.append(f"- {j.get('nome','-')} | padrão={j.get('is_padrao',False)}")

    # ============ NOTIFICAÇÕES RH ============
    notifs = await db.rh_notificacoes.find({}, {"_id": 0}).sort("created_at", -1).to_list(50)
    context_parts.append(f"\n\n{'='*40}\nNOTIFICAÇÕES RH ({len(notifs)})\n{'='*40}")
    for n in notifs[:20]:
        context_parts.append(
            f"- [{n.get('tipo','info')}] {n.get('titulo','-')} → {n.get('funcionario_nome','-')} | lida={n.get('lida',False)}"
        )

    # ============ CONTAS BANCÁRIAS / FORMAS / PLANO ============
    contas_bancarias = await db.contas_bancarias.find({}, {"_id": 0}).to_list(50)
    formas_pag = await db.formas_pagamento.find({}, {"_id": 0}).to_list(50)
    plano_contas = await db.plano_contas.find({}, {"_id": 0}).to_list(200)
    centros = await db.centros_custo.find({}, {"_id": 0}).to_list(50)
    context_parts.append(f"\n\n{'='*40}\nCONTAS BANCÁRIAS ({len(contas_bancarias)})\n{'='*40}")
    for cb in contas_bancarias:
        context_parts.append(f"- {cb.get('nome','-')} | banco={cb.get('banco','-')} | saldo R$ {cb.get('saldo_atual',0):,.2f}")
    context_parts.append(f"\nFORMAS DE PAGAMENTO: {', '.join(f.get('nome','-') for f in formas_pag) or '(nenhuma)'}")
    context_parts.append(f"\nCENTROS DE CUSTO ({len(centros)}): {', '.join(c.get('nome','-') for c in centros) or '(nenhum)'}")
    context_parts.append(f"\nPLANO DE CONTAS: {len(plano_contas)} contas cadastradas")

    return "\n".join(context_parts)


@chatbot_router.post("/chat", response_model=ChatResponse)
async def chat_with_assistant(chat_message: ChatMessage, current_user: dict = Depends(get_current_user)):
    """Chat com o assistente de IA"""
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        platform_context = await get_full_platform_context()
        
        system_message = f"""Você é o assistente virtual inteligente da CRA Construtora.
Você tem ACESSO COMPLETO a TODAS as informações do banco de dados.

DADOS DO SISTEMA:
{platform_context}

INSTRUÇÕES:
1. SEMPRE responda em português brasileiro
2. Use quebras de linha para separar parágrafos
3. Use listas com "•" para enumerar itens
4. Formate valores monetários como R$ 1.234,56
5. NÃO use markdown com asteriscos
6. Seja útil e forneça dados específicos quando perguntado
"""
        
        llm_key = os.environ.get("EMERGENT_LLM_KEY")
        
        llm_chat = LlmChat(
            api_key=llm_key,
            session_id=f"chatbot-{current_user['id']}-{datetime.now().strftime('%Y%m%d%H%M%S')}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.5-flash")
        
        user_message = UserMessage(text=chat_message.message)
        response = await llm_chat.send_message(user_message)
        
        return ChatResponse(
            response=response,
            context_used=[chat_message.module, "database"]
        )
        
    except Exception as e:
        logging.error(f"Erro no chatbot: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erro no assistente: {str(e)}")


@chatbot_router.post("/chat-with-files", response_model=ChatResponse)
async def chat_with_files(
    message: str = Form(default=""),
    files: List[UploadFile] = File(default=[]),
    current_user: dict = Depends(get_current_user)
):
    """Chat com arquivos anexados"""
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage
        
        files_info = []
        file_contents = []
        
        for file in files:
            content = await file.read()
            file_size = len(content)
            filename = file.filename or "arquivo"
            content_type = file.content_type
            
            files_info.append({
                "nome": filename,
                "tipo": content_type,
                "tamanho": f"{file_size / 1024:.1f} KB"
            })
            
            extracted_content = ""
            
            # Arquivos de texto
            if content_type and content_type.startswith('text/'):
                try:
                    text = content.decode('utf-8', errors='ignore')
                    extracted_content = f"📄 CONTEÚDO DE {filename}:\n{text[:10000]}"
                except:
                    extracted_content = f"📄 Arquivo {filename}: não foi possível extrair texto"
            
            # PDFs
            elif content_type == 'application/pdf' or filename.lower().endswith('.pdf'):
                try:
                    from PyPDF2 import PdfReader
                    pdf_reader = PdfReader(io.BytesIO(content))
                    pdf_text = ""
                    for page in pdf_reader.pages[:10]:
                        pdf_text += page.extract_text() + "\n"
                    if pdf_text.strip():
                        extracted_content = f"📑 CONTEÚDO DO PDF {filename}:{pdf_text[:8000]}"
                    else:
                        extracted_content = f"📑 PDF {filename}: {len(pdf_reader.pages)} páginas"
                except Exception as pdf_err:
                    extracted_content = f"⚠️ PDF {filename}: erro ao extrair ({str(pdf_err)[:100]})"
            
            # Excel
            elif filename.lower().endswith(('.xlsx', '.xls')):
                try:
                    import pandas as pd
                    df = pd.read_excel(io.BytesIO(content))
                    preview = df.head(20).to_string()
                    extracted_content = f"📊 PLANILHA {filename}:\nColunas: {list(df.columns)}\nLinhas: {len(df)}\n\nPrimeiras linhas:\n{preview}"
                except Exception as xl_err:
                    extracted_content = f"⚠️ Planilha {filename}: erro ({str(xl_err)[:100]})"
            
            # Imagens
            elif content_type and content_type.startswith('image/'):
                try:
                    from PIL import Image
                    img = Image.open(io.BytesIO(content))
                    extracted_content = f"🖼️ IMAGEM {filename}: {img.format}, {img.width}x{img.height} pixels"
                except:
                    extracted_content = f"🖼️ Imagem {filename}: {file_size / 1024:.1f} KB"
            
            else:
                extracted_content = f"📁 Arquivo {filename}: {content_type or 'desconhecido'} ({file_size / 1024:.1f} KB)"
            
            if extracted_content:
                file_contents.append(extracted_content)
            
            await file.seek(0)
        
        files_context = ""
        if files_info:
            files_context = "\n\nARQUIVOS ANEXADOS:\n"
            for info in files_info:
                files_context += f"• {info['nome']} ({info['tipo']}, {info['tamanho']})\n"
            if file_contents:
                files_context += "\n" + "\n\n".join(file_contents)
        
        platform_context = await get_full_platform_context()
        
        system_message = f"""Você é o assistente virtual da CRA Construtora.

DADOS DO SISTEMA:
{platform_context}

{files_context}

INSTRUÇÕES:
1. SEMPRE responda em português brasileiro
2. Analise os arquivos anexados e faça comentários úteis
3. Relacione os dados dos arquivos com os dados da plataforma
4. NÃO use markdown com asteriscos
"""
        
        llm_key = os.environ.get("EMERGENT_LLM_KEY")
        
        llm_chat = LlmChat(
            api_key=llm_key,
            session_id=f"chatbot-files-{current_user['id']}-{datetime.now().strftime('%Y%m%d%H%M%S')}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.5-flash")
        
        user_text = message if message else "Analise os arquivos que anexei"
        if files_info:
            user_text += f"\n\n[Arquivos: {', '.join([f['nome'] for f in files_info])}]"
        
        user_message = UserMessage(text=user_text)
        response = await llm_chat.send_message(user_message)
        
        return ChatResponse(
            response=response,
            context_used=["arquivos", "todos_os_modulos"]
        )
        
    except Exception as e:
        logging.error(f"Erro no chatbot com arquivos: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erro ao processar arquivos: {str(e)}")



# ============================================================
# Conversas persistentes (estilo ChatGPT) — usadas pela tela principal do RH
# ============================================================
import uuid


# ============ KNOWLEDGE BASE: ENDPOINTS ADMIN ============
from fastapi.responses import FileResponse


@chatbot_router.get("/knowledge-base")
async def list_knowledge_base(current_user: dict = Depends(get_current_user)):
    """Lista os documentos normativos disponíveis no Chat IA do RH."""
    docs = await db.chat_knowledge_base.find(
        {"category": "rh_normativos"}, {"_id": 0, "extracted_text": 0}
    ).sort("name", 1).to_list(50)
    return docs


@chatbot_router.post("/knowledge-base/upload")
async def upload_knowledge_base(
    name: str = Form(...),
    title: str = Form(...),
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
):
    """Faz upload (ou substitui) um documento normativo (PCMSO, PGR, LTCAT, CCT, etc).
    Aceita PDF. Extrai texto via PyPDF e, se vier vazio (PDF escaneado), faz OCR via Gemini."""
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Apenas PDFs são aceitos.")
    storage_dir = "/app/backend/storage/rh_normativos"
    os.makedirs(storage_dir, exist_ok=True)
    name_safe = "".join(c for c in name if c.isalnum() or c in "_-").upper()[:32]
    if not name_safe:
        raise HTTPException(status_code=400, detail="Nome inválido.")
    dst = os.path.join(storage_dir, f"{name_safe}.pdf")
    raw = await file.read()
    with open(dst, "wb") as f:
        f.write(raw)

    # Extrai texto
    pages = 0
    text = ""
    try:
        from pypdf import PdfReader
        r = PdfReader(dst)
        pages = len(r.pages)
        text = "\n".join((p.extract_text() or "") for p in r.pages)
    except Exception as e:
        logging.warning(f"PyPDF falhou: {e}")

    # Se PDF é escaneado (texto vazio), usa Gemini para OCR
    if len(text.strip()) < 100:
        try:
            from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
            chat = LlmChat(
                api_key=os.environ["EMERGENT_LLM_KEY"],
                session_id=f"ocr-{name_safe}-{uuid.uuid4().hex[:6]}",
                system_message="Extrator de texto de PDFs."
            ).with_model("gemini", "gemini-2.5-flash")
            fc = FileContentWithMimeType(file_path=dst, mime_type="application/pdf")
            msg = UserMessage(
                text=f"Extraia TODO o texto deste documento ({title}). Mantenha cláusulas, "
                     "títulos, tabelas, valores e datas. Não invente, apenas extraia.",
                file_contents=[fc],
            )
            text = await chat.send_message(msg)
        except Exception as e:
            logging.error(f"OCR Gemini falhou para {name_safe}: {e}")

    # Substitui ou insere no DB
    await db.chat_knowledge_base.delete_many(
        {"category": "rh_normativos", "name": name_safe}
    )
    doc = {
        "id": str(uuid.uuid4()),
        "category": "rh_normativos",
        "name": name_safe,
        "title": title,
        "extracted_text": text,
        "pdf_path": dst,
        "pdf_size": os.path.getsize(dst),
        "pages": pages,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "uploaded_by": current_user.get("id"),
    }
    await db.chat_knowledge_base.insert_one(dict(doc))
    _invalidate_kb_cache()
    return {
        "id": doc["id"],
        "name": doc["name"],
        "title": doc["title"],
        "pages": doc["pages"],
        "extracted_chars": len(text),
        "pdf_size": doc["pdf_size"],
    }


@chatbot_router.get("/knowledge-base/{doc_id}/download")
async def download_knowledge_base(doc_id: str, current_user: dict = Depends(get_current_user)):
    """Baixa o PDF original do documento normativo."""
    d = await db.chat_knowledge_base.find_one({"id": doc_id}, {"_id": 0})
    if not d:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    pdf_path = d.get("pdf_path")
    if not pdf_path or not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail="Arquivo PDF original indisponível")
    return FileResponse(pdf_path, media_type="application/pdf",
                        filename=f"{d.get('name','documento')}.pdf")


@chatbot_router.delete("/knowledge-base/{doc_id}")
async def delete_knowledge_base(doc_id: str, current_user: dict = Depends(get_current_user)):
    """Remove um documento da base de conhecimento."""
    d = await db.chat_knowledge_base.find_one({"id": doc_id})
    if not d:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    if d.get("pdf_path") and os.path.exists(d["pdf_path"]):
        try:
            os.remove(d["pdf_path"])
        except OSError:
            pass
    await db.chat_knowledge_base.delete_one({"id": doc_id})
    _invalidate_kb_cache()
    return {"deleted": True}


class ConversationCreate(BaseModel):
    title: Optional[str] = "Nova conversa"
    module: str = "rh"


class ConversationOut(BaseModel):
    id: str
    title: str
    module: str
    created_at: str
    updated_at: str
    last_message_preview: Optional[str] = None


class MessageIn(BaseModel):
    content: str


class MessageOut(BaseModel):
    id: str
    role: str  # "user" | "assistant"
    content: str
    created_at: str
    artifact: Optional[dict] = None  # {download_url, label, type} quando IA gera arquivo
    attachments: Optional[List[dict]] = None  # anexos enviados pelo usuário


@chatbot_router.get("/conversations", response_model=List[ConversationOut])
async def list_conversations(
    module: Optional[str] = None,
    current_user: dict = Depends(get_current_user),
):
    """Lista conversas do usuário autenticado, ordenadas por última atualização (mais recentes primeiro)."""
    query = {"user_id": current_user["id"]}
    if module:
        query["module"] = module
    convs = await db.chat_conversations.find(query, {"_id": 0}).sort("updated_at", -1).to_list(200)
    return [ConversationOut(**c) for c in convs]


@chatbot_router.post("/conversations", response_model=ConversationOut)
async def create_conversation(
    payload: ConversationCreate,
    current_user: dict = Depends(get_current_user),
):
    now_iso = datetime.now(timezone.utc).isoformat()
    conv = {
        "id": str(uuid.uuid4()),
        "user_id": current_user["id"],
        "title": payload.title or "Nova conversa",
        "module": payload.module,
        "created_at": now_iso,
        "updated_at": now_iso,
        "last_message_preview": None,
    }
    await db.chat_conversations.insert_one(dict(conv))
    return ConversationOut(**{k: v for k, v in conv.items() if k != "user_id"})


@chatbot_router.get("/conversations/{conv_id}/messages", response_model=List[MessageOut])
async def list_messages(conv_id: str, current_user: dict = Depends(get_current_user)):
    conv = await db.chat_conversations.find_one(
        {"id": conv_id, "user_id": current_user["id"]}, {"_id": 0}
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Conversa não encontrada")
    msgs = await db.chat_messages.find(
        {"conversation_id": conv_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(2000)
    return [MessageOut(**m) for m in msgs]


@chatbot_router.delete("/conversations/{conv_id}")
async def delete_conversation(conv_id: str, current_user: dict = Depends(get_current_user)):
    conv = await db.chat_conversations.find_one({"id": conv_id, "user_id": current_user["id"]})
    if not conv:
        raise HTTPException(status_code=404, detail="Conversa não encontrada")
    await db.chat_messages.delete_many({"conversation_id": conv_id})
    await db.chat_conversations.delete_one({"id": conv_id})
    return {"deleted": True}


# ============ Execução de ferramentas (tool calling) ============

async def _execute_chat_tool(action: str, params: dict, current_user: dict):
    """Executa a ação solicitada pela IA. Retorna (artifact_dict|None, result_text)."""

    if action == "criar_notificacao":
        funcionario_id = params.get("funcionario_id")
        titulo = (params.get("titulo") or "").strip()
        mensagem = (params.get("mensagem") or "").strip()
        tipo = (params.get("tipo") or "info").strip()
        if not funcionario_id or not titulo or not mensagem:
            raise ValueError("Parâmetros obrigatórios: funcionario_id, titulo, mensagem")
        func = await db.funcionarios.find_one({"id": funcionario_id}, {"_id": 0})
        if not func:
            raise ValueError(f"Funcionário {funcionario_id} não encontrado")
        notif_id = str(uuid.uuid4())
        notif_doc = {
            "id": notif_id,
            "tipo": tipo if tipo in ("info", "alerta", "urgente") else "info",
            "titulo": titulo,
            "mensagem": mensagem,
            "categoria": "rh",
            "funcionario_id": funcionario_id,
            "funcionario_nome": func.get("nome"),
            "lida": False,
            "criada_por_ia": True,
            "criada_por_user_id": current_user.get("id"),
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        await db.rh_notificacoes.insert_one(dict(notif_doc))
        result_text = (
            f"✅ Notificação criada com sucesso para {func.get('nome')}.\n"
            f"• Título: {titulo}\n"
            f"• Tipo: {tipo}\n"
            f"• ID: {notif_id}"
        )
        return None, result_text

    if action == "gerar_pdf_notificacao":
        funcionario_id = params.get("funcionario_id")
        tipo_notif = (params.get("tipo_notificacao") or "comunicado").strip().lower()
        motivo = (params.get("motivo") or "Não informado").strip()
        data_ocorrencia = params.get("data_ocorrencia") or datetime.now().strftime("%Y-%m-%d")
        texto_complementar = (params.get("texto_complementar") or "").strip()
        if not funcionario_id:
            raise ValueError("funcionario_id é obrigatório")
        func = await db.funcionarios.find_one({"id": funcionario_id}, {"_id": 0})
        if not func:
            raise ValueError(f"Funcionário {funcionario_id} não encontrado")
        pdf_bytes = _gerar_pdf_notificacao_formal(func, tipo_notif, motivo, data_ocorrencia, texto_complementar)
        artifact_id = str(uuid.uuid4())
        nome_safe = (func.get("nome") or "func").replace(" ", "_")[:30]
        await db.chat_artifacts.insert_one({
            "id": artifact_id,
            "user_id": current_user.get("id"),
            "filename": f"notificacao_{tipo_notif}_{nome_safe}.pdf",
            "content_type": "application/pdf",
            "content_b64": base64.b64encode(pdf_bytes).decode("ascii"),
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        artifact = {
            "type": "pdf",
            "label": f"Baixar Notificação ({tipo_notif.title()}) - {func.get('nome')}",
            "download_url": f"/api/chatbot/artifacts/{artifact_id}",
        }
        # Também grava registro interno
        await db.rh_notificacoes.insert_one({
            "id": str(uuid.uuid4()),
            "tipo": "alerta" if tipo_notif in ("falta", "advertencia") else "info",
            "titulo": f"{tipo_notif.title()}: {motivo[:80]}",
            "mensagem": f"Data: {data_ocorrencia}. {texto_complementar}".strip(),
            "categoria": "rh",
            "funcionario_id": funcionario_id,
            "funcionario_nome": func.get("nome"),
            "lida": False,
            "criada_por_ia": True,
            "criada_por_user_id": current_user.get("id"),
            "tem_pdf": True,
            "pdf_artifact_id": artifact_id,
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        return artifact, (
            f"📄 Notificação de {tipo_notif} gerada para {func.get('nome')}.\n"
            f"• Data da ocorrência: {data_ocorrencia}\n"
            f"• Motivo: {motivo}"
        )

    if action == "gerar_pdf_funcionario":
        funcionario_id = params.get("funcionario_id")
        if not funcionario_id:
            raise ValueError("funcionario_id é obrigatório")
        func = await db.funcionarios.find_one({"id": funcionario_id}, {"_id": 0})
        if not func:
            raise ValueError(f"Funcionário {funcionario_id} não encontrado")
        pdf_bytes = _gerar_pdf_funcionario(func)
        # Persistir o PDF como artefato baixável (rota dedicada abaixo)
        artifact_id = str(uuid.uuid4())
        await db.chat_artifacts.insert_one({
            "id": artifact_id,
            "user_id": current_user.get("id"),
            "filename": f"funcionario_{(func.get('nome') or 'desconhecido').replace(' ', '_')}.pdf",
            "content_type": "application/pdf",
            "content_b64": base64.b64encode(pdf_bytes).decode("ascii"),
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        artifact = {
            "type": "pdf",
            "label": f"Baixar PDF de {func.get('nome')}",
            "download_url": f"/api/chatbot/artifacts/{artifact_id}",
        }
        return artifact, f"📄 PDF gerado para {func.get('nome')}. Clique em 'Baixar' abaixo."

    if action == "gerar_pdf_lista_funcionarios":
        funcionarios = await db.funcionarios.find({}, {"_id": 0}).to_list(500)
        pdf_bytes = _gerar_pdf_lista_funcionarios(funcionarios)
        artifact_id = str(uuid.uuid4())
        await db.chat_artifacts.insert_one({
            "id": artifact_id,
            "user_id": current_user.get("id"),
            "filename": "lista_funcionarios.pdf",
            "content_type": "application/pdf",
            "content_b64": base64.b64encode(pdf_bytes).decode("ascii"),
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        artifact = {
            "type": "pdf",
            "label": f"Baixar lista ({len(funcionarios)} funcionários)",
            "download_url": f"/api/chatbot/artifacts/{artifact_id}",
        }
        return artifact, f"📄 Lista de funcionários gerada ({len(funcionarios)} registros)."

    if action == "gerar_holerite":
        funcionario_id = params.get("funcionario_id")
        mes = params.get("mes")
        ano = params.get("ano")
        if not funcionario_id or not mes or not ano:
            raise ValueError("Parâmetros obrigatórios: funcionario_id, mes, ano")
        try:
            mes_i, ano_i = int(mes), int(ano)
        except (TypeError, ValueError):
            raise ValueError("mes e ano devem ser números inteiros")
        func = await db.funcionarios.find_one({"id": funcionario_id}, {"_id": 0})
        if not func:
            raise ValueError(f"Funcionário {funcionario_id} não encontrado")
        folha = await db.folha_pagamento.find_one(
            {"funcionario_id": funcionario_id, "mes": mes_i, "ano": ano_i}, {"_id": 0}
        )
        if not folha:
            meses_pt = ["", "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
                        "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
            raise ValueError(
                f"Não há folha de pagamento de {meses_pt[mes_i]}/{ano_i} para {func.get('nome')}. "
                f"Gere a folha primeiro em RH ▸ Folha de Pagamento."
            )
        from routes.rh import _build_holerite_pdf
        pdf_bytes = _build_holerite_pdf(folha, func)
        artifact_id = str(uuid.uuid4())
        nome_safe = (func.get("nome") or "func").replace(" ", "_")[:30]
        await db.chat_artifacts.insert_one({
            "id": artifact_id,
            "user_id": current_user.get("id"),
            "filename": f"Holerite_{nome_safe}_{mes_i:02d}_{ano_i}.pdf",
            "content_type": "application/pdf",
            "content_b64": base64.b64encode(pdf_bytes).decode("ascii"),
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        artifact = {
            "type": "pdf",
            "label": f"Baixar Holerite {mes_i:02d}/{ano_i} - {func.get('nome')}",
            "download_url": f"/api/chatbot/artifacts/{artifact_id}",
        }
        liq_brl = f"R$ {float(folha.get('salario_liquido', 0)):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        return artifact, (
            f"💰 Holerite de {func.get('nome')} gerado.\n"
            f"• Competência: {mes_i:02d}/{ano_i}\n"
            f"• Salário Líquido: {liq_brl}"
        )

    if action == "gerar_espelho_ponto":
        mes = params.get("mes")
        ano = params.get("ano")
        funcionario_id = params.get("funcionario_id")  # opcional
        if not mes or not ano:
            raise ValueError("Parâmetros obrigatórios: mes, ano")
        try:
            mes_i, ano_i = int(mes), int(ano)
        except (TypeError, ValueError):
            raise ValueError("mes e ano devem ser números inteiros")
        # Reusa a função que já calcula os dados consolidados de ponto
        from routes.rh import get_ponto_dashboard_mensal, _build_espelho_ponto_pdf
        dashboard = await get_ponto_dashboard_mensal(mes=mes_i, ano=ano_i)
        funcionarios = dashboard.get("funcionarios", [])
        nome_arquivo = "TODOS"
        nome_label = f"todos os funcionários de {mes_i:02d}/{ano_i}"
        if funcionario_id:
            funcionarios = [f for f in funcionarios if f.get("funcionario_id") == funcionario_id]
            if not funcionarios:
                func = await db.funcionarios.find_one({"id": funcionario_id}, {"_id": 0})
                nome = (func or {}).get("nome", funcionario_id)
                raise ValueError(
                    f"Não há registros de ponto de {nome} em {mes_i:02d}/{ano_i}."
                )
            nome_arquivo = (funcionarios[0].get("nome") or "func").replace(" ", "_")[:30]
            nome_label = f"{funcionarios[0].get('nome')} ({mes_i:02d}/{ano_i})"
        if not funcionarios:
            raise ValueError(f"Nenhum registro de ponto encontrado para {mes_i:02d}/{ano_i}.")
        pdf_bytes = _build_espelho_ponto_pdf(funcionarios, mes_i, ano_i)
        artifact_id = str(uuid.uuid4())
        await db.chat_artifacts.insert_one({
            "id": artifact_id,
            "user_id": current_user.get("id"),
            "filename": f"EspelhoPonto_{nome_arquivo}_{mes_i:02d}_{ano_i}.pdf",
            "content_type": "application/pdf",
            "content_b64": base64.b64encode(pdf_bytes).decode("ascii"),
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        artifact = {
            "type": "pdf",
            "label": f"Baixar Espelho de Ponto - {nome_label}",
            "download_url": f"/api/chatbot/artifacts/{artifact_id}",
        }
        return artifact, (
            f"⏰ Espelho de Ponto gerado.\n"
            f"• Período: {mes_i:02d}/{ano_i}\n"
            f"• Abrange: {len(funcionarios)} funcionário(s)"
        )

    raise ValueError(f"Ação desconhecida: {action}")


def _gerar_pdf_notificacao_formal(func: dict, tipo: str, motivo: str, data_ocorrencia: str, texto_complementar: str) -> bytes:
    """Gera PDF formal de notificação usando o template corporativo padrão da plataforma."""
    from utils.pdf_template import (
        create_corporate_doc, add_corporate_header, add_footer,
        get_corporate_styles, build_data_table, build_signatures_table,
    )
    from reportlab.platypus import Paragraph, Spacer

    buf = io.BytesIO()
    doc = create_corporate_doc(buf, title=f"Notificação - {func.get('nome','-')}")
    styles = get_corporate_styles()

    titulos_pt = {
        "falta": "NOTIFICAÇÃO DE FALTA",
        "atraso": "NOTIFICAÇÃO DE ATRASO",
        "advertencia": "ADVERTÊNCIA FORMAL",
        "comunicado": "COMUNICADO INTERNO",
        "outro": "NOTIFICAÇÃO INTERNA",
    }
    titulo_doc = titulos_pt.get(tipo, "NOTIFICAÇÃO INTERNA")

    elements = []
    add_corporate_header(elements, doc_title=titulo_doc)

    # Bloco de dados do colaborador
    def _f(v):
        return str(v) if v not in (None, "") else "-"
    data_oc_fmt = data_ocorrencia
    try:
        if len(data_ocorrencia) >= 10 and data_ocorrencia[4] == "-":
            data_oc_fmt = f"{data_ocorrencia[8:10]}/{data_ocorrencia[5:7]}/{data_ocorrencia[0:4]}"
    except Exception:
        pass

    elements.append(build_data_table([
        ("Colaborador:", _f(func.get("nome"))),
        ("CPF:", _f(func.get("cpf"))),
        ("Cargo:", _f(func.get("cargo"))),
        ("Departamento:", _f(func.get("departamento"))),
        ("Data de admissão:", _f(func.get("data_admissao"))),
        ("Data da ocorrência:", data_oc_fmt),
    ]))
    elements.append(Spacer(1, 16))

    elements.append(Paragraph("Motivo / Descrição", styles["section"]))
    elements.append(Paragraph(motivo, styles["body"]))

    if texto_complementar:
        elements.append(Paragraph("Observações complementares", styles["section"]))
        elements.append(Paragraph(texto_complementar.replace("\n", "<br/>"), styles["body"]))

    rodape_textos = {
        "falta": (
            "Pela presente, comunicamos formalmente o registro da falta acima descrita. "
            "Solicitamos a apresentação de justificativa por escrito (atestado médico ou outro documento) "
            "no prazo de 48 horas a partir do recebimento, sob pena de configuração de falta injustificada "
            "com os respectivos descontos previstos em lei."
        ),
        "atraso": (
            "Comunicamos o registro do atraso conforme descrito acima. "
            "Reforçamos a importância do cumprimento da jornada de trabalho conforme acordado. "
            "Atrasos recorrentes poderão configurar advertência formal."
        ),
        "advertencia": (
            "Pela presente, fica o(a) colaborador(a) ADVERTIDO(A) formalmente em razão dos fatos acima descritos. "
            "Esta advertência ficará arquivada em sua ficha funcional. "
            "A reincidência poderá ensejar a aplicação de penalidades mais severas, conforme a legislação vigente."
        ),
        "comunicado": (
            "Este comunicado tem caráter informativo. Solicitamos a leitura atenta e a observância das informações descritas."
        ),
    }
    elements.append(Spacer(1, 8))
    elements.append(Paragraph(rodape_textos.get(tipo, rodape_textos["comunicado"]), styles["body"]))
    elements.append(Spacer(1, 36))

    elements.append(build_signatures_table())
    add_footer(elements, "Documento gerado pelo Assistente IA do RH · CRA Construtora")

    doc.build(elements)
    return buf.getvalue()


def _gerar_pdf_funcionario(func: dict) -> bytes:
    """Gera PDF da ficha do funcionário usando template corporativo."""
    from utils.pdf_template import (
        create_corporate_doc, add_corporate_header, add_footer,
        get_corporate_styles, build_data_table,
    )
    from reportlab.platypus import Paragraph, Spacer

    buf = io.BytesIO()
    doc = create_corporate_doc(buf, title=f"Ficha - {func.get('nome','-')}")
    styles = get_corporate_styles()
    elements = []
    add_corporate_header(elements, doc_title="FICHA DO FUNCIONÁRIO")

    def _f(v):
        return str(v) if v not in (None, "") else "-"

    elements.append(Paragraph("Dados Pessoais", styles["section"]))
    elements.append(build_data_table([
        ("Nome:", _f(func.get("nome"))),
        ("CPF:", _f(func.get("cpf"))),
        ("RG:", _f(func.get("rg"))),
        ("Data de nascimento:", _f(func.get("data_nascimento"))),
        ("Telefone:", _f(func.get("telefone"))),
        ("E-mail:", _f(func.get("email"))),
        ("Endereço:", _f(func.get("endereco"))),
    ]))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Dados Funcionais", styles["section"]))
    elements.append(build_data_table([
        ("Cargo:", _f(func.get("cargo"))),
        ("Departamento:", _f(func.get("departamento"))),
        ("Data de admissão:", _f(func.get("data_admissao"))),
        ("Tipo de contrato:", _f(func.get("tipo_contrato"))),
        ("Salário:", f"R$ {(func.get('salario') or 0):,.2f}"),
        ("Status:", _f(func.get("status"))),
    ]))

    add_footer(elements)
    doc.build(elements)
    return buf.getvalue()


def _gerar_pdf_lista_funcionarios(funcionarios: list) -> bytes:
    """Gera PDF com lista geral de funcionários usando template corporativo."""
    from utils.pdf_template import (
        create_corporate_doc, add_corporate_header, add_footer,
        get_corporate_styles, header_table_style,
    )
    from reportlab.platypus import Paragraph, Spacer, Table
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = create_corporate_doc(buf, title="Lista de Funcionários")
    styles = get_corporate_styles()
    elements = []
    add_corporate_header(
        elements,
        doc_title="LISTA DE FUNCIONÁRIOS",
        subtitle=f"{len(funcionarios)} colaboradores · {datetime.now().strftime('%d/%m/%Y às %H:%M')}",
    )

    if not funcionarios:
        elements.append(Paragraph("Nenhum funcionário cadastrado.", styles["body"]))
    else:
        rows = [["Nome", "Cargo", "Departamento", "Salário"]]
        for f in funcionarios:
            rows.append([
                (f.get("nome") or "-")[:36],
                (f.get("cargo") or "-")[:24],
                (f.get("departamento") or "-")[:18],
                f"R$ {(f.get('salario') or 0):,.2f}",
            ])
        t = Table(rows, colWidths=[6 * cm, 4.5 * cm, 3.5 * cm, 3 * cm], repeatRows=1)
        t.setStyle(header_table_style())
        elements.append(t)

    elements.append(Spacer(1, 12))
    add_footer(elements, "Documento gerado pelo Assistente IA do RH · CRA Construtora")
    doc.build(elements)
    return buf.getvalue()


@chatbot_router.get("/artifacts/{artifact_id}")
async def download_chat_artifact(artifact_id: str, current_user: dict = Depends(get_current_user)):
    """Download de um artefato gerado pela IA (PDF, etc.)."""
    from fastapi.responses import StreamingResponse as _SR
    art = await db.chat_artifacts.find_one(
        {"id": artifact_id, "user_id": current_user.get("id")}, {"_id": 0}
    )
    if not art:
        raise HTTPException(status_code=404, detail="Artefato não encontrado")
    content = base64.b64decode(art["content_b64"])
    return _SR(
        io.BytesIO(content),
        media_type=art.get("content_type", "application/octet-stream"),
        headers={
            "Content-Disposition": f'attachment; filename="{art.get("filename", "arquivo.pdf")}"',
        },
    )


@chatbot_router.post("/conversations/{conv_id}/messages", response_model=MessageOut)
async def send_message_in_conversation(
    conv_id: str,
    payload: MessageIn,
    current_user: dict = Depends(get_current_user),
):
    """Envia mensagem do usuário, gera resposta com Gemini 2.5 Flash + contexto da plataforma e persiste tudo."""
    from emergentintegrations.llm.chat import LlmChat, UserMessage

    conv = await db.chat_conversations.find_one(
        {"id": conv_id, "user_id": current_user["id"]}, {"_id": 0}
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Conversa não encontrada")

    now_iso = datetime.now(timezone.utc).isoformat()

    # Persistir mensagem do usuário
    user_msg = {
        "id": str(uuid.uuid4()),
        "conversation_id": conv_id,
        "role": "user",
        "content": payload.content,
        "created_at": now_iso,
    }
    await db.chat_messages.insert_one(dict(user_msg))

    # Construir histórico recente para contexto (últimas 30 mensagens) — emergentintegrations
    # mantém histórico por session_id; aqui re-injetamos o histórico textual no system prompt.
    msgs_anteriores = await db.chat_messages.find(
        {"conversation_id": conv_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(60)
    historico_txt = "\n".join(
        f"[{m['role'].upper()}]: {m['content']}" for m in msgs_anteriores[:-1]  # exclui a recém-inserida (vai no UserMessage)
    )

    # Contexto completo da plataforma
    platform_context = await get_full_platform_context()

    # Base de conhecimento (PCMSO, PGR, LTCAT, CCT) — sempre disponível para o RH
    kb_context = await _build_knowledge_base_context()

    # Module-aware system prompt (RH em destaque)
    foco_modulo = ""
    if conv.get("module") == "rh":
        foco_modulo = (
            "Você é o assistente principal do MÓDULO DE RECURSOS HUMANOS desta plataforma. "
            "Tem acesso a TODOS os dados da plataforma (financeiro, frota, obras etc.) mas seu "
            "foco principal é responder com profundidade sobre Funcionários, Ponto Eletrônico, "
            "Folha de Pagamento, Férias, EPI, Custos de RH e Jornadas. Use os dados reais do "
            "banco para todas as respostas — nunca invente. Quando perguntarem sobre dados de "
            "outros módulos, responda com igual precisão.\n\n"
        )

    system_message = f"""{foco_modulo}Você é o assistente virtual inteligente da CRA Construtora.
Você tem ACESSO COMPLETO a TODAS as informações do banco de dados.

DADOS DO SISTEMA:
{platform_context}

{kb_context}

HISTÓRICO RECENTE DESTA CONVERSA:
{historico_txt or "(início da conversa)"}

INSTRUÇÕES:
1. SEMPRE responda em português brasileiro.
2. Use quebras de linha para separar parágrafos e listas com "•".
3. Formate valores monetários como R$ 1.234,56.
4. NÃO use markdown com asteriscos.
5. Seja útil, direto e específico — cite números e nomes reais quando possível.
6. Se a pergunta for ambígua, pergunte qual recorte o usuário quer.
7. Para perguntas sobre EXAMES (admissional, periódico, mudança de função, demissional),
   EPIs por função, RISCOS ocupacionais, CARGOS / pisos salariais, ADICIONAIS, BENEFÍCIOS
   da convenção coletiva, JORNADAS regulamentares, FÉRIAS por CCT — SEMPRE consulte os
   documentos normativos da seção "DOCUMENTOS NORMATIVOS DE RH" acima e cite o documento
   de origem (PCMSO, PGR, LTCAT ou CCT) na resposta.

FERRAMENTAS DISPONÍVEIS (você PODE executar ações reais na plataforma):

Quando o usuário pedir para criar uma notificação, gerar um PDF, ou executar uma ação,
emita uma chamada de ferramenta ANTES da sua mensagem em texto, no formato:

<<TOOL>>{{"action":"NOME_DA_ACAO","params":{{...}}}}<<END>>

Ações suportadas:
• action="criar_notificacao" — params: {{"funcionario_id":"<id>","titulo":"...","mensagem":"...","tipo":"info|alerta|urgente"}}
  Use apenas para registrar uma notificação INTERNA na plataforma (NÃO gera PDF).
• action="gerar_pdf_notificacao" — params: {{"funcionario_id":"<id>","tipo_notificacao":"falta|atraso|advertencia|comunicado|outro","motivo":"...","data_ocorrencia":"YYYY-MM-DD","texto_complementar":"..."}}
  GERA UM PDF FORMAL TIMBRADO com cabeçalho da CRA, dados completos do funcionário,
  motivo, data da ocorrência, texto explicativo e linhas para assinatura. USE ESTA AÇÃO
  quando o usuário pedir "notificação de falta", "advertência", "comunicado" em PDF.
• action="gerar_pdf_funcionario" — params: {{"funcionario_id":"<id>"}}
  Gera PDF com a FICHA COMPLETA do funcionário (cargo, salário, admissão).
• action="gerar_pdf_lista_funcionarios" — params: {{}}
  Gera PDF com a lista geral de funcionários ativos.
• action="gerar_holerite" — params: {{"funcionario_id":"<id>","mes":<1-12>,"ano":<YYYY>}}
  Gera o HOLERITE em PDF (proventos × descontos, salário líquido em destaque, FGTS, INSS).
  Use quando o usuário pedir "holerite", "contracheque" ou "demonstrativo de pagamento".
  Exige folha já gerada para o período. Se o usuário não informar mês/ano, pergunte.
• action="gerar_espelho_ponto" — params: {{"mes":<1-12>,"ano":<YYYY>,"funcionario_id":"<id>"}}
  Gera o ESPELHO DE PONTO em PDF (batidas, faltas, banco de horas, abonos).
  funcionario_id é OPCIONAL: se omitido, gera o espelho consolidado de TODOS os funcionários
  do período. Use quando o usuário pedir "espelho de ponto", "relatório de ponto",
  "frequência" ou "registro de ponto" em PDF.

Após o bloco <<TOOL>>...<<END>>, escreva uma mensagem natural em português confirmando
o que está sendo feito. NÃO emita uma ação se não tiver os parâmetros necessários —
nesse caso, pergunte ao usuário primeiro (ex.: "Para qual funcionário e qual a data da falta?").

Os IDs de funcionários estão na seção FUNCIONÁRIOS RH do contexto. Use o ID correto.
Quando o usuário pedir "gerar uma notificação de falta para o João", use a ação
`gerar_pdf_notificacao` com tipo_notificacao="falta", incluindo motivo e data.
Se o usuário não informou a data, use a data de hoje. Se não informou o motivo,
use "Falta sem justificativa apresentada".

Exemplos para HOLERITE/ESPELHO:
- "Gere o holerite do João de fevereiro" → use `gerar_holerite` com mes=2 e ano do contexto.
- "Quero o espelho de ponto de abril" → use `gerar_espelho_ponto` com mes=4 (sem funcionario_id = consolidado).
- "Espelho de ponto da Maria em março/2026" → use `gerar_espelho_ponto` com funcionario_id, mes=3, ano=2026.
Se o usuário não informou o ano, assuma o ano atual.
Se a folha de pagamento ainda não existir para o período, INFORME ao usuário sem chamar a tool.
"""

    llm_key = os.environ.get("EMERGENT_LLM_KEY")
    llm_chat = LlmChat(
        api_key=llm_key,
        session_id=f"conv-{conv_id}",
        system_message=system_message,
    ).with_model("gemini", "gemini-2.5-flash")

    try:
        ai_response = await llm_chat.send_message(UserMessage(text=payload.content))
    except Exception as e:
        logging.error(f"Erro Gemini: {e}")
        raise HTTPException(status_code=500, detail=f"Erro na IA: {str(e)[:200]}")

    # ============ Detecção e execução de ferramentas ============
    artifact: Optional[dict] = None
    import re as _re_tools
    import json as _json_tools

    tool_match = _re_tools.search(
        r"<<TOOL>>\s*(\{.*?\})\s*<<END>>", ai_response, _re_tools.DOTALL
    )
    if tool_match:
        try:
            tool_data = _json_tools.loads(tool_match.group(1))
            tool_action = (tool_data.get("action") or "").strip()
            tool_params = tool_data.get("params") or {}
            artifact, tool_result_text = await _execute_chat_tool(
                tool_action, tool_params, current_user
            )
            # Remove o bloco <<TOOL>>...<<END>> da resposta final e prefixa com resultado
            ai_response = _re_tools.sub(
                r"<<TOOL>>.*?<<END>>", "", ai_response, flags=_re_tools.DOTALL
            ).strip()
            if tool_result_text:
                ai_response = f"{tool_result_text}\n\n{ai_response}".strip()
        except Exception as et:
            logging.warning(f"Falha ao executar tool do chatbot: {et}")
            ai_response = _re_tools.sub(
                r"<<TOOL>>.*?<<END>>", "", ai_response, flags=_re_tools.DOTALL
            ).strip()
            ai_response = (
                f"⚠️ Falha ao executar a ação solicitada: {str(et)[:200]}\n\n{ai_response}"
            ).strip()

    # Persistir resposta do assistente
    assistant_msg = {
        "id": str(uuid.uuid4()),
        "conversation_id": conv_id,
        "role": "assistant",
        "content": ai_response,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "artifact": artifact,
    }
    await db.chat_messages.insert_one(dict(assistant_msg))

    # Atualizar conversa: timestamp + preview + título auto se ainda for "Nova conversa"
    update_fields = {
        "updated_at": assistant_msg["created_at"],
        "last_message_preview": (ai_response or "")[:120],
    }
    if conv.get("title") in (None, "", "Nova conversa"):
        # Pega as primeiras palavras da pergunta como título
        title_auto = (payload.content or "").strip().split("\n")[0][:60]
        if title_auto:
            update_fields["title"] = title_auto
    await db.chat_conversations.update_one({"id": conv_id}, {"$set": update_fields})

    return MessageOut(**{k: v for k, v in assistant_msg.items() if k != "conversation_id"})


# ========================================================================
# CHAT COM ANEXOS (multipart) — leitura e alteração de arquivos
# ========================================================================

@chatbot_router.post("/conversations/{conv_id}/messages-with-files", response_model=MessageOut)
async def send_message_with_files(
    conv_id: str,
    content: str = Form(""),
    files: List[UploadFile] = File(default=[]),
    current_user: dict = Depends(get_current_user),
):
    """Envia mensagem com anexos. Suporta PDF, imagem, Excel/CSV, Word, texto.

    A IA lê o conteúdo dos arquivos junto com o prompt e responde. Se o usuário
    pedir alteração no arquivo, a IA gera uma nova versão e a disponibiliza
    para download (artifact com URL).
    """
    from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType

    conv = await db.chat_conversations.find_one(
        {"id": conv_id, "user_id": current_user["id"]}, {"_id": 0}
    )
    if not conv:
        raise HTTPException(status_code=404, detail="Conversa não encontrada")

    if not content and not files:
        raise HTTPException(status_code=400, detail="Envie uma mensagem ou anexe um arquivo")

    now_iso = datetime.now(timezone.utc).isoformat()

    # Processa cada anexo
    anexos_info: List[dict] = []
    for f in files:
        if not f.filename:
            continue
        info = await _processar_anexo_chat(f, conv_id)
        anexos_info.append(info)

    # Lista de file_contents para Gemini (PDFs, imagens nativas)
    gemini_files = [
        FileContentWithMimeType(file_path=a["gemini_file"]["path"], mime_type=a["gemini_file"]["mime"])
        for a in anexos_info if a.get("gemini_file")
    ]

    # Texto inline (Excel, CSV, Word, txt)
    blocos_inline = []
    for a in anexos_info:
        if a.get("inline_text"):
            blocos_inline.append(
                f"\n\n════ ARQUIVO ANEXADO: {a['filename']} ({a['mime']}) ════\n{a['inline_text']}"
            )
    texto_inline = "".join(blocos_inline)

    # Persistir mensagem do usuário (com anexos referenciados)
    user_msg = {
        "id": str(uuid.uuid4()),
        "conversation_id": conv_id,
        "role": "user",
        "content": content or "(Sem texto — apenas arquivos anexados)",
        "attachments": [
            {"id": a["id"], "filename": a["filename"], "mime": a["mime"], "size": a["size"]}
            for a in anexos_info
        ],
        "created_at": now_iso,
    }
    await db.chat_messages.insert_one(dict(user_msg))

    # Histórico recente (textual)
    msgs_anteriores = await db.chat_messages.find(
        {"conversation_id": conv_id}, {"_id": 0}
    ).sort("created_at", 1).to_list(60)
    historico_txt = "\n".join(
        f"[{m['role'].upper()}]: {m.get('content','')[:1000]}" for m in msgs_anteriores[:-1]
    )

    # Detecta pedido de alteração
    pediu_alteracao = _detectar_pedido_alteracao(content)
    arquivo_alvo = next(
        (a for a in anexos_info if a["filename"].rsplit(".", 1)[-1].lower() in ("csv", "txt", "md", "xlsx", "xls", "docx")),
        None,
    )

    # System prompt
    instrucoes_alteracao = ""
    if pediu_alteracao and arquivo_alvo:
        ext = arquivo_alvo["filename"].rsplit(".", 1)[-1].lower()
        formato_saida = {
            "csv": "CSV (separador vírgula)",
            "txt": "texto puro",
            "md": "Markdown",
            "xlsx": "CSV (separador vírgula) que será convertido para Excel",
            "xls": "CSV (separador vírgula) que será convertido para Excel",
            "docx": "texto puro (será convertido para Word)",
        }.get(ext, "texto puro")
        instrucoes_alteracao = f"""
INSTRUÇÃO ESPECIAL — O USUÁRIO PEDIU ALTERAÇÃO NO ARQUIVO `{arquivo_alvo['filename']}`:
1. Explique brevemente em português o que você alterou.
2. Em seguida, forneça o ARQUIVO COMPLETO ALTERADO entre os marcadores:
   <<<ARTEFATO filename="{arquivo_alvo['filename']}" format="{ext}">>>
   ...conteúdo completo em {formato_saida}, sem nenhum comentário extra...
   <<<FIM_ARTEFATO>>>
3. NÃO use blocos de código markdown dentro do artefato. Apenas o conteúdo cru.
4. Mantenha cabeçalhos/estrutura originais quando possível.
"""

    system_message = f"""Você é o assistente virtual inteligente da CRA Construtora.
Você lê arquivos anexados pelo usuário (PDF, imagem, Excel, CSV, Word, texto) e responde
com base no conteúdo deles JUNTO com o prompt textual. NÃO faça alterações no sistema —
você apenas lê e responde. SE o usuário pedir explicitamente para alterar o arquivo,
gere a NOVA VERSÃO dentro dos marcadores especificados abaixo.

HISTÓRICO RECENTE DA CONVERSA:
{historico_txt or "(início)"}

{instrucoes_alteracao}

INSTRUÇÕES:
- Responda em português brasileiro.
- Cite trechos/valores do arquivo quando relevante.
- Se vários arquivos foram anexados, mencione cada um.
- Se o usuário não pediu alteração, apenas responda às perguntas sobre o conteúdo.
"""

    llm_key = os.environ.get("EMERGENT_LLM_KEY")
    llm_chat = LlmChat(
        api_key=llm_key,
        session_id=f"conv-files-{conv_id}",
        system_message=system_message,
    ).with_model("gemini", "gemini-2.5-flash")

    prompt_completo = (content or "(Analise os arquivos anexados)") + texto_inline

    try:
        if gemini_files:
            user_message = UserMessage(text=prompt_completo, file_contents=gemini_files)
        else:
            user_message = UserMessage(text=prompt_completo)
        ai_response = await llm_chat.send_message(user_message)
    except Exception as e:
        logging.error(f"Erro Gemini (multipart): {e}")
        raise HTTPException(status_code=500, detail=f"Erro na IA: {str(e)[:200]}")

    # Detecta artefato (arquivo alterado) na resposta
    artifact = None
    import re as _re
    art_match = _re.search(
        r'<<<ARTEFATO\s+filename="([^"]+)"\s+format="([^"]+)">>>(.*?)<<<FIM_ARTEFATO>>>',
        ai_response,
        _re.DOTALL,
    )
    if art_match:
        art_filename = art_match.group(1)
        art_format = art_match.group(2).lower()
        art_content = art_match.group(3).strip()

        # Salva o artefato no disco e cria URL pública
        art_id = uuid.uuid4().hex[:12]
        new_filename = f"alterado_{art_filename}"

        try:
            if art_format in ("xlsx", "xls"):
                # Recebe CSV → converte para XLSX
                import openpyxl as _xl, csv as _csv, io as _io
                wb = _xl.Workbook()
                ws = wb.active
                ws.title = "Alterado"
                reader = _csv.reader(_io.StringIO(art_content))
                for row in reader:
                    ws.append(row)
                target = CHAT_UPLOADS_DIR / f"art_{conv_id}_{art_id}_{new_filename}"
                wb.save(str(target))
            elif art_format == "docx":
                # Texto puro → Word
                from docx import Document
                doc = Document()
                for line in art_content.split("\n"):
                    doc.add_paragraph(line)
                target = CHAT_UPLOADS_DIR / f"art_{conv_id}_{art_id}_{new_filename}"
                doc.save(str(target))
            else:
                # csv / txt / md → grava direto
                target = CHAT_UPLOADS_DIR / f"art_{conv_id}_{art_id}_{new_filename}"
                with open(target, "w", encoding="utf-8") as f:
                    f.write(art_content)
            artifact = {
                "type": "file",
                "filename": new_filename,
                "format": art_format,
                "download_url": f"/api/chatbot/artifacts-file/{art_id}",
                "size": target.stat().st_size,
            }
            # Persistir caminho para download
            await db.chat_artifacts.insert_one({
                "id": art_id,
                "conversation_id": conv_id,
                "filename": new_filename,
                "path": str(target),
                "mime": {
                    "csv": "text/csv",
                    "txt": "text/plain",
                    "md": "text/markdown",
                    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    "xls": "application/vnd.ms-excel",
                    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                }.get(art_format, "application/octet-stream"),
                "created_at": now_iso,
            })
            # Remove o bloco do texto exibido ao usuário
            ai_response = _re.sub(
                r'<<<ARTEFATO[^>]+>>>.*?<<<FIM_ARTEFATO>>>', "", ai_response, flags=_re.DOTALL
            ).strip()
            ai_response += f"\n\n📎 Arquivo alterado pronto para download: **{new_filename}**"
        except Exception as e:
            logging.error(f"Erro ao gerar artefato: {e}")
            ai_response += f"\n\n⚠️ Falha ao gerar o arquivo alterado: {str(e)[:200]}"

    # Persistir resposta
    assistant_msg = {
        "id": str(uuid.uuid4()),
        "conversation_id": conv_id,
        "role": "assistant",
        "content": ai_response,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "artifact": artifact,
    }
    await db.chat_messages.insert_one(dict(assistant_msg))

    update_fields = {
        "updated_at": assistant_msg["created_at"],
        "last_message_preview": (ai_response or "")[:120],
    }
    if conv.get("title") in (None, "", "Nova conversa"):
        title_auto = (content or "Anexo enviado").strip().split("\n")[0][:60]
        update_fields["title"] = title_auto
    await db.chat_conversations.update_one({"id": conv_id}, {"$set": update_fields})

    return MessageOut(**{k: v for k, v in assistant_msg.items() if k != "conversation_id"})


@chatbot_router.get("/artifacts-file/{art_id}")
async def download_chat_artifact_file(art_id: str, current_user: dict = Depends(get_current_user)):
    """Faz download de um arquivo gerado pelo chat (alteração de arquivo do usuário)."""
    from fastapi.responses import FileResponse
    art = await db.chat_artifacts.find_one({"id": art_id}, {"_id": 0})
    if not art:
        raise HTTPException(status_code=404, detail="Artefato não encontrado")
    path = art.get("path")
    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Arquivo expirado/removido")
    return FileResponse(path, filename=art["filename"], media_type=art.get("mime"))
