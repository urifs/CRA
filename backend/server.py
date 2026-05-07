from fastapi import FastAPI, APIRouter, HTTPException, Depends, UploadFile, File, Form, Response, Body, Query
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
"""
================================================================================
CRA CONSTRUTORA - SISTEMA DE GESTÃO EMPRESARIAL
================================================================================
Backend API - FastAPI + MongoDB

ÍNDICE DO ARQUIVO:
------------------
Linha ~40   : MODELS - Modelos Pydantic
Linha ~350  : AUTH HELPERS - Funções de autenticação
Linha ~450  : AUTH ROUTES - Rotas de autenticação
Linha ~520  : CATEGORY ROUTES - Rotas de categorias de máquinas
Linha ~600  : MACHINE ROUTES - Rotas de máquinas
Linha ~800  : MAINTENANCE ROUTES - Rotas de manutenções
Linha ~1000 : STOCK ROUTES - Rotas de estoque
Linha ~1300 : OBRA ROUTES - Rotas de obras/projetos
Linha ~1500 : DASHBOARD ROUTES - Rotas do dashboard
Linha ~1800 : ADMIN FINANCIAL ROUTES - Rotas administrativas financeiras
Linha ~3200 : NOTIFICATION ROUTES - Rotas de notificações
Linha ~3400 : ADMIN PANEL ROUTES - Rotas do painel admin (usuários, auditoria, DB)
Linha ~4200 : EXPORT ROUTES - Rotas de exportação PDF
Linha ~4600 : CHATBOT ROUTES - Rotas do chatbot com IA
Linha ~4900 : FILE ROUTES - Rotas de upload de arquivos
Linha ~5100 : TASK ROUTES - Rotas do sistema de tarefas
Linha ~5350 : APP STARTUP/SHUTDOWN - Eventos da aplicação

ARQUITETURA:
------------
- Frontend: React + TailwindCSS + Shadcn/UI
- Backend: FastAPI + MongoDB (Motor)
- Auth: JWT com bcrypt para senhas
- Integração: Gemini AI para chatbot

DEPENDÊNCIAS PRINCIPAIS:
------------------------
- fastapi, motor (async MongoDB)
- pydantic, python-jose (JWT)
- bcrypt, reportlab (PDF)
- emergentintegrations (Gemini AI)
================================================================================
"""

from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict, EmailStr
from typing import List, Optional
import uuid
from datetime import datetime, timezone, timedelta
import jwt
import bcrypt
import base64

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# Import modular routes
from routes.rh import rh_router
from routes.admin import admin_router
from routes.chatbot import chatbot_router
from routes.folha_importacao import folha_router as folha_importacao_router, fin_folha_router as fin_folha_solicitacoes_router
from routes.storage import storage_router
from routes.exports import export_router
from routes.stock import stock_router
from routes.obras import obras_router
from routes.conciliacao import conciliacao_router
from routes.nfse import nfse_router
from routes.nfe import nfe_router
from routes.financeiro import financeiro_router
from routes.emissao_nf import emissao_router
from routes.importacao_nf import importacao_router, importar_nfe, importar_nfse
from routes.exports_all import exports_all_router
from routes.dashboard import dashboard_router
from routes.medicoes import medicoes_router

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'fleet-maintenance-secret-key-2024')
JWT_ALGORITHM = "HS256"

# Create the main app
app = FastAPI(title="CRA Construtora")

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer()
optional_security = HTTPBearer(auto_error=False)

# ============ MODELS ============

class UserCreate(BaseModel):
    name: str
    email: EmailStr
    password: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    email: str
    role: str = "gerenciamento"
    created_at: str

class TokenResponse(BaseModel):
    token: str
    user: UserResponse

class CategoryCreate(BaseModel):
    name: str
    description: Optional[str] = ""
    color: Optional[str] = "#E31A1A"  # Cor padrão vermelho

class CategoryResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    description: str
    color: str = "#E31A1A"
    subcategories_count: int = 0
    created_at: str

# Subcategorias de máquinas
class SubcategoryCreate(BaseModel):
    name: str
    category_id: str
    description: Optional[str] = ""

class SubcategoryResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    category_id: str
    category_name: str = ""
    description: str
    created_at: str

class MachineCreate(BaseModel):
    name: str
    plate: Optional[str] = ""
    category_id: str
    subcategory_id: Optional[str] = None
    brand: Optional[str] = ""
    model: Optional[str] = ""
    year: Optional[int] = None
    notes: Optional[str] = ""
    obra_id: Optional[str] = None
    fleet_id: Optional[str] = None
    subfleet_id: Optional[str] = None
    operator_id: Optional[str] = None  # ID do funcionário/operador
    identificador_tipo: Optional[str] = None  # 'chassi' ou 'serie'
    identificador_numero: Optional[str] = None
    status: Optional[str] = "patio"  # patio, operacional, manutencao

class MachineResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    plate: Optional[str] = ""
    category_id: str
    category_name: Optional[str] = ""
    subcategory_id: Optional[str] = None
    subcategory_name: Optional[str] = ""
    brand: str
    model: str
    year: Optional[int] = None
    notes: str
    status: str
    obra_id: Optional[str] = None
    obra_name: Optional[str] = ""
    fleet_id: Optional[str] = None
    fleet_name: Optional[str] = ""
    subfleet_id: Optional[str] = None
    subfleet_name: Optional[str] = ""
    operator_id: Optional[str] = None
    operator_name: Optional[str] = ""
    identificador_tipo: Optional[str] = None
    identificador_numero: Optional[str] = None
    horimetro_atual: Optional[float] = None
    created_at: str

class MaintenanceCreate(BaseModel):
    machine_id: str
    part_name: str
    replacement_date: str
    part_value: float
    maintenance_type: str  # preventiva ou corretiva
    description: Optional[str] = ""
    is_oil_change: bool = False  # marca se é troca de óleo

class MaintenanceResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    machine_id: str
    machine_name: Optional[str] = ""
    machine_plate: Optional[str] = ""
    part_name: str
    replacement_date: str
    part_value: float
    maintenance_type: str
    description: str
    photos: List[str]
    is_oil_change: bool = False
    created_at: str

class CategoryMachineCount(BaseModel):
    category_id: str
    category_name: str
    category_color: str
    count: int

class DashboardStats(BaseModel):
    total_machines: int
    total_maintenances: int
    preventive_count: int
    corrective_count: int
    total_spent: float
    recent_maintenances: List[MaintenanceResponse]
    low_stock_count: int = 0
    oil_change_alerts: int = 0
    machines_by_category: List[CategoryMachineCount] = []

# ============ NFE IMPORT MODELS ============

class NFeCertificadoCreate(BaseModel):
    cnpj: str
    razao_social: str
    uf: str = "SP"
    ambiente: str = "producao"  # producao ou homologacao
    certificado_base64: str  # Certificado .pfx em base64
    senha_certificado: str
    ativo: bool = True
    inscricao_municipal: Optional[str] = None  # Inscrição municipal para NFS-e
    url_nfse: Optional[str] = None             # URL do webservice NFS-e da prefeitura

class NFeCertificadoResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    cnpj: str
    razao_social: str
    uf: str
    ambiente: str
    ativo: bool
    ultimo_nsu: str = "000000000000000"
    created_at: str
    updated_at: Optional[str] = None
    inscricao_municipal: Optional[str] = None
    url_nfse: Optional[str] = None

class NFeItemResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    codigo: str
    descricao: str
    ncm: str
    cfop: str
    unidade: str
    quantidade: float
    valor_unitario: float
    valor_total: float

class NFeImportadaResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    certificado_id: str
    cnpj_destinatario: str
    chave_acesso: str
    numero_nf: str
    serie: str
    data_emissao: str
    cnpj_emitente: str
    razao_social_emitente: str
    valor_total: float
    itens: List[NFeItemResponse] = []
    xml_base64: Optional[str] = None
    nsu: str
    conta_pagar_id: Optional[str] = None  # Vinculada a conta a pagar
    status: str = "nova"  # nova, processada, ignorada
    created_at: str

# ============ EMISSÃO DE NF-e / NFS-e MODELS ============

class NFeItemEmissao(BaseModel):
    """Item para emissão de NF-e"""
    produto_id: Optional[str] = None  # ID do produto cadastrado
    codigo: str
    descricao: str
    ncm: str = "00000000"
    cfop: str = "5102"  # Default: Venda de mercadoria
    unidade: str = "UN"
    quantidade: float
    valor_unitario: float
    valor_total: float
    # Tributos
    origem: str = "0"  # 0 = Nacional
    cst_icms: str = "00"
    aliquota_icms: float = 0
    valor_icms: float = 0
    cst_pis: str = "01"
    aliquota_pis: float = 0
    valor_pis: float = 0
    cst_cofins: str = "01"
    aliquota_cofins: float = 0
    valor_cofins: float = 0
    cst_ipi: str = "50"
    aliquota_ipi: float = 0
    valor_ipi: float = 0

class NFeEmissaoCreate(BaseModel):
    """Modelo para emissão de NF-e (Nota Fiscal de Produtos)"""
    certificado_id: str
    # Destinatário
    dest_cpf_cnpj: str
    dest_razao_social: str
    dest_ie: Optional[str] = None  # Inscrição Estadual
    dest_email: Optional[str] = None
    dest_telefone: Optional[str] = None
    # Endereço do destinatário
    dest_cep: str
    dest_logradouro: str
    dest_numero: str
    dest_complemento: Optional[str] = None
    dest_bairro: str
    dest_cidade: str
    dest_uf: str
    dest_codigo_municipio: Optional[str] = None
    # Dados da nota
    natureza_operacao: str = "Venda de Mercadoria"
    tipo_operacao: str = "1"  # 0=Entrada, 1=Saída
    finalidade: str = "1"  # 1=Normal, 2=Complementar, 3=Ajuste, 4=Devolução
    consumidor_final: str = "1"  # 0=Normal, 1=Consumidor Final
    presenca_comprador: str = "1"  # 1=Presencial, 9=Não presencial
    # Forma de pagamento
    forma_pagamento: str = "01"  # 01=Dinheiro, 02=Cheque, 03=Cartão Crédito, etc.
    valor_pagamento: Optional[float] = None
    # Transporte
    modalidade_frete: str = "9"  # 9=Sem frete
    transportador_cnpj: Optional[str] = None
    transportador_razao: Optional[str] = None
    # Itens
    itens: List[NFeItemEmissao]
    # Totais
    valor_produtos: float
    valor_frete: float = 0
    valor_seguro: float = 0
    valor_desconto: float = 0
    valor_outros: float = 0
    valor_total: float
    # Informações adicionais
    info_complementar: Optional[str] = None

class NFSeItemEmissao(BaseModel):
    """Item/Serviço para emissão de NFS-e"""
    produto_id: Optional[str] = None  # ID do produto/serviço cadastrado
    codigo_servico: str  # Código do serviço LC 116/2003
    descricao: str
    quantidade: float = 1
    valor_unitario: float
    valor_total: float
    aliquota_iss: float = 0
    valor_iss: float = 0

class NFSeEmissaoCreate(BaseModel):
    """Modelo para emissão de NFS-e (Nota Fiscal de Serviços) - Palmas/TO"""
    certificado_id: str
    # Tomador do serviço
    tomador_cpf_cnpj: str
    tomador_razao_social: str
    tomador_ie: Optional[str] = None
    tomador_im: Optional[str] = None  # Inscrição Municipal
    tomador_email: Optional[str] = None
    tomador_telefone: Optional[str] = None
    # Endereço do tomador
    tomador_cep: str
    tomador_logradouro: str
    tomador_numero: str
    tomador_complemento: Optional[str] = None
    tomador_bairro: str
    tomador_cidade: str
    tomador_uf: str
    tomador_codigo_municipio: Optional[str] = None
    # Dados do serviço
    codigo_cnae: Optional[str] = None
    codigo_tributario_municipio: str
    item_lista_servico: str  # Código do serviço LC 116/2003
    discriminacao: str  # Descrição detalhada do serviço
    # Valores
    valor_servicos: float
    valor_deducoes: float = 0
    valor_pis: float = 0
    valor_cofins: float = 0
    valor_inss: float = 0
    valor_ir: float = 0
    valor_csll: float = 0
    outras_retencoes: float = 0
    valor_iss: float = 0
    aliquota_iss: float = 0
    valor_liquido: float
    # ISS
    iss_retido: bool = False
    # Informações adicionais
    info_complementar: Optional[str] = None
    # Itens (opcional, para detalhamento)
    itens: List[NFSeItemEmissao] = []

class NotaFiscalEmitidaResponse(BaseModel):
    """Resposta após emissão de nota fiscal"""
    model_config = ConfigDict(extra="ignore")
    id: str
    tipo: str  # "nfe" ou "nfse"
    numero: str
    serie: str
    chave_acesso: Optional[str] = None
    protocolo: Optional[str] = None
    status: str  # "autorizada", "rejeitada", "pendente", "rascunho"
    mensagem: Optional[str] = None
    xml_base64: Optional[str] = None
    pdf_base64: Optional[str] = None
    created_at: str

# ============ OIL CHANGE / USAGE MODELS ============

class UsageLogCreate(BaseModel):
    machine_id: str
    hours: float
    notes: Optional[str] = ""

class UsageLogResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    machine_id: str
    machine_name: Optional[str] = ""
    machine_plate: Optional[str] = ""
    hours: float
    notes: str
    created_at: str

class OilChangeStatusResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    machine_id: str
    machine_name: str
    machine_plate: str
    last_oil_change_date: Optional[str] = None
    hours_since_change: float
    hours_remaining: float
    days_since_change: int
    days_remaining: int
    needs_alert: bool
    alert_reason: Optional[str] = None

# ============ FLEET MODELS ============

class FleetCreate(BaseModel):
    name: str
    description: Optional[str] = ""

class FleetUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None

class FleetResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    description: str
    machines_count: int = 0
    subfleets_count: int = 0
    created_at: str

class SubfleetCreate(BaseModel):
    name: str
    fleet_id: str
    description: Optional[str] = ""

class SubfleetUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None

class SubfleetResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    fleet_id: str
    fleet_name: str = ""
    description: str
    machines_count: int = 0
    created_at: str

class NotificationResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    machine_id: str
    machine_name: str
    machine_plate: str
    notification_type: str
    message: str
    hours_remaining: Optional[float] = None
    days_remaining: Optional[int] = None
    created_at: str

# ============ STOCK MODELS ============

class StockCategoryCreate(BaseModel):
    name: str

class StockCategoryResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    subcategories_count: int = 0
    created_at: str

# Subcategorias de estoque
class StockSubcategoryCreate(BaseModel):
    name: str
    category_id: str

class StockSubcategoryResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    category_id: str
    category_name: str = ""
    created_at: str

class StockItemCreate(BaseModel):
    name: str
    code: Optional[str] = ""
    category: Optional[str] = ""  # filtro, óleo, correia, etc.
    subcategory_id: Optional[str] = None
    unit: str = "un"  # un, L, kg, etc.
    quantity: float = 0
    min_quantity: float = 0
    unit_price: Optional[float] = 0
    location: Optional[str] = ""
    notes: Optional[str] = ""
    machine_ids: Optional[List[str]] = []  # IDs das máquinas vinculadas

class StockItemResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    code: str
    category: str
    subcategory_id: Optional[str] = None
    subcategory_name: Optional[str] = ""
    unit: str
    quantity: float
    min_quantity: float
    unit_price: float
    location: str
    notes: str
    is_low_stock: bool
    created_at: str
    machine_ids: Optional[List[str]] = []
    machine_names: Optional[List[str]] = []

class StockMovementCreate(BaseModel):
    item_id: str
    movement_type: str  # entrada ou saida
    quantity: float
    reason: Optional[str] = ""
    notes: Optional[str] = ""

class StockMovementResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    item_id: str
    item_name: Optional[str] = ""
    movement_type: str
    quantity: float
    previous_quantity: float
    new_quantity: float
    reason: str
    notes: str
    created_at: str

# ============ OBRA (PROJECT) MODELS ============

class ObraCreate(BaseModel):
    name: str
    description: Optional[str] = ""
    location: Optional[str] = ""
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    status: str = "em_andamento"  # em_andamento, concluida, pausada

class ObraResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    description: str
    location: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    status: str
    machine_count: int = 0
    total_maintenance_cost: float = 0
    created_at: str

class ObraDetailResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    name: str
    description: str
    location: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    status: str
    machines: List[MachineResponse] = []
    maintenances: List[MaintenanceResponse] = []
    total_maintenance_cost: float = 0
    preventive_cost: float = 0
    corrective_cost: float = 0
    created_at: str

class MachineObraUpdate(BaseModel):
    obra_id: Optional[str] = None  # None to remove from obra

# ============ HORIMETRO MODELS ============

class HorimetroCreate(BaseModel):
    machine_id: str
    data: str
    hora_inicial: float
    hora_final: float
    horas_trabalhadas: Optional[float] = None
    operador: Optional[str] = None
    observacoes: Optional[str] = None
    tipo_medicao: Optional[str] = "hora"  # "hora" ou "km"

class HorimetroResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    machine_id: str
    machine_name: Optional[str] = None
    data: str
    hora_inicial: float
    hora_final: float
    horas_trabalhadas: float
    operador: Optional[str] = None
    observacoes: Optional[str] = None
    tipo_medicao: str = "hora"
    created_by: str
    created_at: str

# ============ COMBUSTIVEL MODELS ============

# Modelo para veículo abastecedor (tanque)
class CompartimentoOleo(BaseModel):
    """Modelo para compartimento de óleo dinâmico"""
    id: Optional[str] = None
    item_estoque_id: str  # ID do item do estoque (óleo)
    item_nome: Optional[str] = None
    unidade_medida: str = "L"  # L, KG, ML, etc.
    capacidade: float = 0
    quantidade_atual: float = 0

class VeiculoAbastecedorCreate(BaseModel):
    machine_id: str  # ID da máquina que é abastecedora
    capacidade_diesel: float = 0  # Capacidade total de diesel em litros
    capacidade_oleo: float = 0  # Capacidade total de óleo em litros (legado)
    capacidade_graxa: float = 0  # Capacidade total de graxa em litros
    litros_diesel: float = 0  # Litros de diesel atual
    litros_oleo: float = 0  # Litros de óleo atual (legado)
    litros_graxa: float = 0  # Litros de graxa atual
    operador_id: Optional[str] = None
    # Novos campos para compartimentos dinâmicos de óleo
    compartimentos_oleo: Optional[List[dict]] = None  # Lista de compartimentos de óleo

class VeiculoAbastecedorResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    machine_id: str
    machine_name: Optional[str] = None
    capacidade_diesel: float = 0
    capacidade_oleo: float = 0
    capacidade_graxa: float = 0
    litros_diesel: float = 0
    litros_oleo: float = 0
    litros_graxa: float = 0
    operador_id: Optional[str] = None
    operador_nome: Optional[str] = None
    compartimentos_oleo: Optional[List[dict]] = None  # Lista de compartimentos de óleo
    created_by: str
    created_at: str
    updated_at: Optional[str] = None

# Modelo para registro de combustível (abastecimento)
class CombustivelCreate(BaseModel):
    machine_id: str
    data: str
    tipo_registro: str = "abastecido"  # "abastecedor" ou "abastecido"
    tipo_medicao: str = "litros"  # 'litros_hora' ou 'litros_km' ou 'litros'
    hora_km_inicial: Optional[float] = None
    litros_diesel: float = 0
    litros_oleo: float = 0
    litros_graxa: float = 0
    # Para abastecido
    fonte_abastecimento: Optional[str] = None  # "interno", "externo" ou "posto"
    veiculo_abastecedor_id: Optional[str] = None  # ID do veículo abastecedor (se interno)
    posto_id: Optional[str] = None  # ID do posto/fornecedor parceiro (cadastro)
    operador_id: Optional[str] = None
    observacoes: Optional[str] = None
    # Para óleos específicos
    oleos_utilizados: Optional[List[dict]] = None  # Lista de óleos com item_id, quantidade, unidade

class CombustivelResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    machine_id: str
    machine_name: Optional[str] = None
    data: str
    tipo_registro: str = "abastecido"
    tipo_medicao: str = "litros"
    hora_km_inicial: Optional[float] = None
    litros_diesel: float = 0
    litros_oleo: float = 0
    litros_graxa: float = 0
    fonte_abastecimento: Optional[str] = None
    veiculo_abastecedor_id: Optional[str] = None
    veiculo_abastecedor_nome: Optional[str] = None
    operador_id: Optional[str] = None
    operador_nome: Optional[str] = None
    observacoes: Optional[str] = None
    created_by: str
    created_at: str

# ============ AUDIT LOG MODELS ============

class AuditLogResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    user_id: str
    user_name: str
    user_email: str
    action: str  # criar, editar, excluir
    entity_type: str  # maquina, manutencao, obra, estoque, categoria, etc
    entity_id: str
    entity_name: str
    details: Optional[str] = ""
    created_at: str

# ============ AUTH HELPERS ============

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode('utf-8'), hashed.encode('utf-8'))

def create_token(user_id: str, email: str) -> str:
    payload = {
        "user_id": user_id,
        "email": email,
        "exp": datetime.now(timezone.utc).timestamp() + 86400 * 7  # 7 days
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        user = await db.users.find_one({"id": payload["user_id"]}, {"_id": 0})
        if not user:
            raise HTTPException(status_code=401, detail="Usuário não encontrado")
        return user
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expirado")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Token inválido")

# ============ AUDIT LOG HELPER ============

async def create_audit_log(
    user: dict,
    action: str,
    entity_type: str,
    entity_id: str,
    entity_name: str,
    details: str = "",
    module: str = ""
):
    """Create an audit log entry for tracking user actions."""
    # Mapear ações para descrições em português
    action_labels = {
        "criar": "Criou",
        "create": "Criou",
        "editar": "Editou",
        "update": "Editou",
        "excluir": "Excluiu",
        "delete": "Excluiu",
        "login": "Fez login",
        "logout": "Fez logout"
    }
    
    # Mapear tipos de entidade para descrições
    entity_labels = {
        "categoria": "Categoria de Máquina",
        "máquina": "Máquina",
        "maquina": "Máquina",
        "manutenção": "Manutenção",
        "manutencao": "Manutenção",
        "registro de uso": "Registro de Uso (Horímetro)",
        "categoria de estoque": "Categoria de Estoque",
        "item de estoque": "Item de Estoque",
        "movimentação de estoque (entrada)": "Entrada de Estoque",
        "movimentação de estoque (saída)": "Saída de Estoque",
        "obra": "Obra/Projeto",
        "cadastro": "Cadastro (Cliente/Fornecedor)",
        "conta_pagar": "Conta a Pagar",
        "conta_receber": "Conta a Receber",
        "produto": "Produto",
        "ordem_servico": "Ordem de Serviço",
        "plano_contas": "Plano de Contas",
        "centro_custo": "Centro de Custo",
        "forma_pagamento": "Forma de Pagamento",
        "aluguel": "Aluguel de Máquina",
        "documento (users)": "Documento (Usuários)",
        "documento (machines)": "Documento (Máquinas)",
        "documento (maintenances)": "Documento (Manutenções)",
        "usuario": "Usuário",
        "user": "Usuário"
    }
    
    # Determinar o módulo automaticamente se não especificado
    if not module:
        admin_entities = ["cadastro", "conta_pagar", "conta_receber", "produto", 
                         "ordem_servico", "plano_contas", "centro_custo", 
                         "forma_pagamento", "aluguel"]
        gerenciamento_entities = ["categoria", "máquina", "maquina", "manutenção", 
                                  "manutencao", "registro de uso", "categoria de estoque",
                                  "item de estoque", "obra"]
        admin_panel_entities = ["usuario", "user"]
        
        entity_lower = entity_type.lower()
        if any(e in entity_lower for e in admin_entities):
            module = "Administrativo"
        elif any(e in entity_lower for e in gerenciamento_entities) or "estoque" in entity_lower:
            module = "Gerenciamento"
        elif any(e in entity_lower for e in admin_panel_entities) or "documento" in entity_lower:
            module = "Painel Admin"
        else:
            module = "Sistema"
    
    # Criar descrição detalhada da ação
    action_label = action_labels.get(action.lower(), action.capitalize())
    entity_label = entity_labels.get(entity_type.lower(), entity_type)
    
    # Construir descrição completa
    full_action = f"{action_label} {entity_label}"
    
    # Adicionar detalhes formatados
    detailed_info = f"Item: {entity_name}"
    if details:
        detailed_info += f" | {details}"
    
    audit_id = str(uuid.uuid4())
    audit_doc = {
        "id": audit_id,
        "user_id": user["id"],
        "user_name": user["name"],
        "user_email": user["email"],
        "action": full_action,
        "entity_type": entity_type,
        "entity_id": entity_id,
        "entity_name": entity_name,
        "module": module,
        "details": detailed_info,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.audit_logs.insert_one(audit_doc)

# ============ AUTH ROUTES ============

@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user: UserCreate):
    existing = await db.users.find_one({"email": user.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email já cadastrado")
    
    user_id = str(uuid.uuid4())
    user_doc = {
        "id": user_id,
        "name": user.name,
        "email": user.email,
        "password": hash_password(user.password),
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.users.insert_one(user_doc)
    
    token = create_token(user_id, user.email)
    user_response = UserResponse(
        id=user_id,
        name=user.name,
        email=user.email,
        created_at=user_doc["created_at"]
    )
    return TokenResponse(token=token, user=user_response)

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user or not verify_password(credentials.password, user["password"]):
        raise HTTPException(status_code=401, detail="Credenciais inválidas")
    
    # Atualizar último login
    await db.users.update_one(
        {"id": user["id"]},
        {"$set": {"last_login": datetime.now(timezone.utc).isoformat()}}
    )
    
    token = create_token(user["id"], user["email"])
    user_response = UserResponse(
        id=user["id"],
        name=user["name"],
        email=user["email"],
        role=user.get("role", "gerenciamento"),
        created_at=user["created_at"]
    )
    return TokenResponse(token=token, user=user_response)

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: dict = Depends(get_current_user)):
    return UserResponse(
        id=current_user["id"],
        name=current_user["name"],
        email=current_user["email"],
        role=current_user.get("role", "gerenciamento"),
        created_at=current_user["created_at"]
    )

# ============ CREATE ADMIN ACCOUNT (PUBLIC) ============

class AdminCreate(BaseModel):
    name: str
    email: str
    password: str

@api_router.post("/auth/create-admin")
async def create_admin_account(data: AdminCreate):
    """
    Cria uma conta administrador com acesso total ao sistema.
    Disponível publicamente na tela de login.
    """
    # Verificar se email já existe
    existing = await db.users.find_one({"email": data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Este email já está cadastrado")
    
    if len(data.password) < 6:
        raise HTTPException(status_code=400, detail="A senha deve ter pelo menos 6 caracteres")
    
    user_id = str(uuid.uuid4())
    user_doc = {
        "id": user_id,
        "name": data.name,
        "email": data.email,
        "password": hash_password(data.password),
        "role": "admin",  # Acesso total
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    return {"message": "Conta administrador criada com sucesso!", "email": data.email}

# ============ SPECIAL ADMIN SETUP ENDPOINT ============
# Endpoint especial para promover usuário para admin (usar uma vez para configurar)

class AdminSetupRequest(BaseModel):
    email: str
    secret_key: str

@api_router.post("/auth/setup-admin")
async def setup_admin(request: AdminSetupRequest):
    """
    Endpoint especial para promover um usuário existente para admin.
    Usar a chave secreta: CRA-SETUP-2026-ADMIN
    """
    SETUP_SECRET = "CRA-SETUP-2026-ADMIN"
    
    if request.secret_key != SETUP_SECRET:
        raise HTTPException(status_code=403, detail="Chave secreta inválida")
    
    user = await db.users.find_one({"email": request.email})
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")
    
    # Promover para admin com acesso total
    await db.users.update_one(
        {"email": request.email},
        {"$set": {"role": "admin"}}
    )
    
    return {
        "message": f"Usuário {request.email} promovido para administrador com sucesso!",
        "email": request.email,
        "new_role": "admin"
    }

# ============ CATEGORY ROUTES ============

@api_router.post("/categories", response_model=CategoryResponse)
async def create_category(category: CategoryCreate, current_user: dict = Depends(get_current_user)):
    category_id = str(uuid.uuid4())
    category_doc = {
        "id": category_id,
        "name": category.name,
        "description": category.description or "",
        "color": category.color or "#E31A1A",
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.categories.insert_one(category_doc)
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="categoria",
        entity_id=category_id,
        entity_name=category.name
    )
    
    return CategoryResponse(
        id=category_id,
        name=category.name,
        description=category.description or "",
        color=category.color or "#E31A1A",
        created_at=category_doc["created_at"]
    )

@api_router.get("/categories", response_model=List[CategoryResponse])
async def get_categories(current_user: dict = Depends(get_current_user)):
    categories = await db.categories.find({}, {"_id": 0}).to_list(100)
    return [CategoryResponse(
        id=c["id"],
        name=c["name"],
        description=c.get("description", ""),
        color=c.get("color", "#E31A1A"),
        created_at=c["created_at"]
    ) for c in categories]

@api_router.delete("/categories/{category_id}")
async def delete_category(category_id: str, current_user: dict = Depends(get_current_user)):
    existing = await db.categories.find_one({"id": category_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Categoria não encontrada")
    
    await db.categories.delete_one({"id": category_id})
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="categoria",
        entity_id=category_id,
        entity_name=existing["name"]
    )
    
    return {"message": "Categoria removida com sucesso"}

@api_router.put("/categories/{category_id}", response_model=CategoryResponse)
async def update_category(category_id: str, category: CategoryCreate, current_user: dict = Depends(get_current_user)):
    existing = await db.categories.find_one({"id": category_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Categoria não encontrada")
    
    await db.categories.update_one(
        {"id": category_id},
        {"$set": {"name": category.name, "description": category.description or "", "color": category.color or "#E31A1A"}}
    )
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="editar",
        entity_type="categoria",
        entity_id=category_id,
        entity_name=category.name,
        details=f"Nome anterior: {existing['name']}"
    )
    
    return CategoryResponse(
        id=category_id,
        name=category.name,
        description=category.description or "",
        color=category.color or "#E31A1A",
        created_at=existing["created_at"]
    )

# ============ SUBCATEGORY ROUTES (Machine Subcategories) ============

@api_router.post("/subcategories", response_model=SubcategoryResponse)
async def create_subcategory(subcategory: SubcategoryCreate, current_user: dict = Depends(get_current_user)):
    """Create a new subcategory for machines"""
    # Verify category exists
    category = await db.categories.find_one({"id": subcategory.category_id}, {"_id": 0})
    if not category:
        raise HTTPException(status_code=404, detail="Categoria não encontrada")
    
    subcategory_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    doc = {
        "id": subcategory_id,
        "name": subcategory.name,
        "category_id": subcategory.category_id,
        "description": subcategory.description or "",
        "created_at": now
    }
    
    await db.subcategories.insert_one(doc)
    
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="subcategoria",
        entity_id=subcategory_id,
        entity_name=subcategory.name,
        module="Gerenciamento"
    )
    
    return SubcategoryResponse(**doc, category_name=category["name"])

@api_router.get("/subcategories", response_model=List[SubcategoryResponse])
async def list_subcategories(category_id: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    """List all subcategories, optionally filtered by category"""
    query = {"category_id": category_id} if category_id else {}
    subcategories = await db.subcategories.find(query, {"_id": 0}).to_list(500)
    
    # Get category names
    categories = await db.categories.find({}, {"_id": 0}).to_list(100)
    category_map = {c["id"]: c["name"] for c in categories}
    
    return [SubcategoryResponse(**s, category_name=category_map.get(s["category_id"], "")) for s in subcategories]

@api_router.put("/subcategories/{subcategory_id}", response_model=SubcategoryResponse)
async def update_subcategory(subcategory_id: str, data: SubcategoryCreate, current_user: dict = Depends(get_current_user)):
    """Update a subcategory"""
    existing = await db.subcategories.find_one({"id": subcategory_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Subcategoria não encontrada")
    
    await db.subcategories.update_one(
        {"id": subcategory_id},
        {"$set": {"name": data.name, "description": data.description or ""}}
    )
    
    updated = await db.subcategories.find_one({"id": subcategory_id}, {"_id": 0})
    category = await db.categories.find_one({"id": updated["category_id"]}, {"_id": 0})
    
    return SubcategoryResponse(**updated, category_name=category["name"] if category else "")

@api_router.delete("/subcategories/{subcategory_id}")
async def delete_subcategory(subcategory_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a subcategory"""
    subcategory = await db.subcategories.find_one({"id": subcategory_id}, {"_id": 0})
    if not subcategory:
        raise HTTPException(status_code=404, detail="Subcategoria não encontrada")
    
    # Remove subcategory from machines
    await db.machines.update_many({"subcategory_id": subcategory_id}, {"$set": {"subcategory_id": None}})
    
    await db.subcategories.delete_one({"id": subcategory_id})
    
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="subcategoria",
        entity_id=subcategory_id,
        entity_name=subcategory["name"],
        module="Gerenciamento"
    )
    
    return {"message": "Subcategoria excluída com sucesso"}

# ============ FLEET ROUTES ============

@api_router.post("/fleets", response_model=FleetResponse)
async def create_fleet(fleet: FleetCreate, current_user: dict = Depends(get_current_user)):
    """Create a new fleet"""
    fleet_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    fleet_doc = {
        "id": fleet_id,
        "name": fleet.name,
        "description": fleet.description or "",
        "created_at": now
    }
    
    await db.fleets.insert_one(fleet_doc)
    
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="frota",
        entity_id=fleet_id,
        entity_name=fleet.name,
        module="Gerenciamento"
    )
    
    return FleetResponse(**fleet_doc, machines_count=0, subfleets_count=0)

@api_router.get("/fleets", response_model=List[FleetResponse])
async def list_fleets(current_user: dict = Depends(get_current_user)):
    """List all fleets"""
    fleets = await db.fleets.find({}, {"_id": 0}).to_list(500)
    
    result = []
    for f in fleets:
        # Count machines in this fleet
        machines_count = await db.machines.count_documents({"fleet_id": f["id"]})
        # Count subfleets
        subfleets_count = await db.subfleets.count_documents({"fleet_id": f["id"]})
        result.append(FleetResponse(
            **f,
            machines_count=machines_count,
            subfleets_count=subfleets_count
        ))
    
    return result

@api_router.get("/fleets/{fleet_id}", response_model=FleetResponse)
async def get_fleet(fleet_id: str, current_user: dict = Depends(get_current_user)):
    """Get a specific fleet"""
    fleet = await db.fleets.find_one({"id": fleet_id}, {"_id": 0})
    if not fleet:
        raise HTTPException(status_code=404, detail="Frota não encontrada")
    
    machines_count = await db.machines.count_documents({"fleet_id": fleet_id})
    subfleets_count = await db.subfleets.count_documents({"fleet_id": fleet_id})
    
    return FleetResponse(**fleet, machines_count=machines_count, subfleets_count=subfleets_count)

@api_router.put("/fleets/{fleet_id}", response_model=FleetResponse)
async def update_fleet(fleet_id: str, fleet: FleetUpdate, current_user: dict = Depends(get_current_user)):
    """Update a fleet"""
    existing = await db.fleets.find_one({"id": fleet_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Frota não encontrada")
    
    update_data = {k: v for k, v in fleet.dict().items() if v is not None}
    if update_data:
        await db.fleets.update_one({"id": fleet_id}, {"$set": update_data})
    
    updated = await db.fleets.find_one({"id": fleet_id}, {"_id": 0})
    machines_count = await db.machines.count_documents({"fleet_id": fleet_id})
    subfleets_count = await db.subfleets.count_documents({"fleet_id": fleet_id})
    
    await create_audit_log(
        user=current_user,
        action="atualizar",
        entity_type="frota",
        entity_id=fleet_id,
        entity_name=updated["name"],
        module="Gerenciamento"
    )
    
    return FleetResponse(**updated, machines_count=machines_count, subfleets_count=subfleets_count)

@api_router.delete("/fleets/{fleet_id}")
async def delete_fleet(fleet_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a fleet"""
    fleet = await db.fleets.find_one({"id": fleet_id}, {"_id": 0})
    if not fleet:
        raise HTTPException(status_code=404, detail="Frota não encontrada")
    
    # Remove fleet reference from machines
    await db.machines.update_many({"fleet_id": fleet_id}, {"$set": {"fleet_id": None, "subfleet_id": None}})
    
    # Delete subfleets
    await db.subfleets.delete_many({"fleet_id": fleet_id})
    
    # Delete fleet
    await db.fleets.delete_one({"id": fleet_id})
    
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="frota",
        entity_id=fleet_id,
        entity_name=fleet["name"],
        module="Gerenciamento"
    )
    
    return {"message": "Frota excluída com sucesso"}

# ============ SUBFLEET ROUTES ============

@api_router.post("/subfleets", response_model=SubfleetResponse)
async def create_subfleet(subfleet: SubfleetCreate, current_user: dict = Depends(get_current_user)):
    """Create a new subfleet"""
    # Check if fleet exists
    fleet = await db.fleets.find_one({"id": subfleet.fleet_id}, {"_id": 0})
    if not fleet:
        raise HTTPException(status_code=404, detail="Frota não encontrada")
    
    subfleet_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    subfleet_doc = {
        "id": subfleet_id,
        "name": subfleet.name,
        "fleet_id": subfleet.fleet_id,
        "description": subfleet.description or "",
        "created_at": now
    }
    
    await db.subfleets.insert_one(subfleet_doc)
    
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="subfrota",
        entity_id=subfleet_id,
        entity_name=subfleet.name,
        module="Gerenciamento"
    )
    
    return SubfleetResponse(**subfleet_doc, fleet_name=fleet["name"], machines_count=0)

@api_router.get("/subfleets", response_model=List[SubfleetResponse])
async def list_subfleets(fleet_id: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    """List subfleets, optionally filtered by fleet"""
    query = {"fleet_id": fleet_id} if fleet_id else {}
    subfleets = await db.subfleets.find(query, {"_id": 0}).to_list(500)
    
    # Get fleet names
    fleet_ids = list(set(s["fleet_id"] for s in subfleets))
    fleets = await db.fleets.find({"id": {"$in": fleet_ids}}, {"_id": 0}).to_list(500)
    fleet_map = {f["id"]: f["name"] for f in fleets}
    
    result = []
    for s in subfleets:
        machines_count = await db.machines.count_documents({"subfleet_id": s["id"]})
        result.append(SubfleetResponse(
            **s,
            fleet_name=fleet_map.get(s["fleet_id"], ""),
            machines_count=machines_count
        ))
    
    return result

@api_router.put("/subfleets/{subfleet_id}", response_model=SubfleetResponse)
async def update_subfleet(subfleet_id: str, subfleet: SubfleetUpdate, current_user: dict = Depends(get_current_user)):
    """Update a subfleet"""
    existing = await db.subfleets.find_one({"id": subfleet_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Subfrota não encontrada")
    
    update_data = {k: v for k, v in subfleet.dict().items() if v is not None}
    if update_data:
        await db.subfleets.update_one({"id": subfleet_id}, {"$set": update_data})
    
    updated = await db.subfleets.find_one({"id": subfleet_id}, {"_id": 0})
    fleet = await db.fleets.find_one({"id": updated["fleet_id"]}, {"_id": 0})
    machines_count = await db.machines.count_documents({"subfleet_id": subfleet_id})
    
    await create_audit_log(
        user=current_user,
        action="atualizar",
        entity_type="subfrota",
        entity_id=subfleet_id,
        entity_name=updated["name"],
        module="Gerenciamento"
    )
    
    return SubfleetResponse(**updated, fleet_name=fleet["name"] if fleet else "", machines_count=machines_count)

@api_router.delete("/subfleets/{subfleet_id}")
async def delete_subfleet(subfleet_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a subfleet"""
    subfleet = await db.subfleets.find_one({"id": subfleet_id}, {"_id": 0})
    if not subfleet:
        raise HTTPException(status_code=404, detail="Subfrota não encontrada")
    
    # Remove subfleet reference from machines
    await db.machines.update_many({"subfleet_id": subfleet_id}, {"$set": {"subfleet_id": None}})
    
    # Delete subfleet
    await db.subfleets.delete_one({"id": subfleet_id})
    
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="subfrota",
        entity_id=subfleet_id,
        entity_name=subfleet["name"],
        module="Gerenciamento"
    )
    
    return {"message": "Subfrota excluída com sucesso"}

# ============ MACHINE ROUTES ============

@api_router.post("/machines", response_model=MachineResponse)
async def create_machine(machine: MachineCreate, current_user: dict = Depends(get_current_user)):
    # Check if plate already exists (only if plate is provided)
    if machine.plate:
        existing = await db.machines.find_one({"plate": machine.plate.upper()})
        if existing:
            raise HTTPException(status_code=400, detail="Placa já cadastrada")
    
    # Get category name
    category = await db.categories.find_one({"id": machine.category_id}, {"_id": 0})
    category_name = category["name"] if category else ""
    
    # Get obra name if specified
    obra_name = ""
    if machine.obra_id:
        obra = await db.obras.find_one({"id": machine.obra_id}, {"_id": 0})
        obra_name = obra["name"] if obra else ""
    
    # Get fleet and subfleet names
    fleet_name = ""
    subfleet_name = ""
    if machine.fleet_id:
        fleet = await db.fleets.find_one({"id": machine.fleet_id}, {"_id": 0})
        fleet_name = fleet["name"] if fleet else ""
    if machine.subfleet_id:
        subfleet = await db.subfleets.find_one({"id": machine.subfleet_id}, {"_id": 0})
        subfleet_name = subfleet["name"] if subfleet else ""
    
    # Get operator name
    operator_name = ""
    if machine.operator_id:
        operator = await db.cadastros.find_one({"id": machine.operator_id}, {"_id": 0})
        operator_name = operator.get("nome_razao", "") if operator else ""
    
    # Get subcategory name
    subcategory_name = ""
    if machine.subcategory_id:
        subcategory = await db.subcategories.find_one({"id": machine.subcategory_id}, {"_id": 0})
        subcategory_name = subcategory["name"] if subcategory else ""
    
    machine_id = str(uuid.uuid4())
    machine_doc = {
        "id": machine_id,
        "name": machine.name,
        "plate": (machine.plate or "").upper(),
        "category_id": machine.category_id,
        "subcategory_id": machine.subcategory_id,
        "brand": machine.brand or "",
        "model": machine.model or "",
        "year": machine.year,
        "notes": machine.notes or "",
        "obra_id": machine.obra_id,
        "fleet_id": machine.fleet_id,
        "subfleet_id": machine.subfleet_id,
        "operator_id": machine.operator_id,
        "status": machine.status or "patio",  # patio, operacional, manutencao
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.machines.insert_one(machine_doc)
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="máquina",
        entity_id=machine_id,
        entity_name=f"{machine.name} ({(machine.plate or '').upper()})"
    )
    
    return MachineResponse(
        id=machine_id,
        name=machine.name,
        plate=(machine.plate or "").upper(),
        category_id=machine.category_id,
        category_name=category_name,
        subcategory_id=machine.subcategory_id,
        subcategory_name=subcategory_name,
        brand=machine.brand or "",
        model=machine.model or "",
        year=machine.year,
        notes=machine.notes or "",
        obra_id=machine.obra_id,
        obra_name=obra_name,
        fleet_id=machine.fleet_id,
        fleet_name=fleet_name,
        subfleet_id=machine.subfleet_id,
        subfleet_name=subfleet_name,
        operator_id=machine.operator_id,
        operator_name=operator_name,
        status="operational",
        created_at=machine_doc["created_at"]
    )

@api_router.get("/machines", response_model=List[MachineResponse])
async def get_machines(obra_id: Optional[str] = None, fleet_id: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    query = {}
    if obra_id:
        query["obra_id"] = obra_id
    if fleet_id:
        query["fleet_id"] = fleet_id
    
    machines = await db.machines.find(query, {"_id": 0}).to_list(1000)
    
    # Get all categories
    categories = await db.categories.find({}, {"_id": 0}).to_list(100)
    category_map = {c["id"]: c["name"] for c in categories}
    
    # Get all obras
    obras = await db.obras.find({}, {"_id": 0}).to_list(100)
    obra_map = {o["id"]: o["name"] for o in obras}
    
    # Get all fleets and subfleets
    fleets = await db.fleets.find({}, {"_id": 0}).to_list(100)
    fleet_map = {f["id"]: f["name"] for f in fleets}
    
    subfleets = await db.subfleets.find({}, {"_id": 0}).to_list(500)
    subfleet_map = {s["id"]: s["name"] for s in subfleets}
    
    # Get all cadastros (operators)
    cadastros = await db.cadastros.find({}, {"_id": 0}).to_list(1000)
    operator_map = {c["id"]: c.get("nome_razao", "") for c in cadastros}
    
    # Get all subcategories
    subcategories = await db.subcategories.find({}, {"_id": 0}).to_list(500)
    subcategory_map = {s["id"]: s["name"] for s in subcategories}
    
    return [MachineResponse(
        id=m["id"],
        name=m["name"],
        plate=m.get("plate", ""),
        category_id=m["category_id"],
        category_name=category_map.get(m["category_id"], ""),
        subcategory_id=m.get("subcategory_id"),
        subcategory_name=subcategory_map.get(m.get("subcategory_id", ""), ""),
        brand=m.get("brand", ""),
        model=m.get("model", ""),
        year=m.get("year"),
        notes=m.get("notes", ""),
        status=m.get("status", "operational"),
        obra_id=m.get("obra_id"),
        obra_name=obra_map.get(m.get("obra_id", ""), ""),
        fleet_id=m.get("fleet_id"),
        fleet_name=fleet_map.get(m.get("fleet_id", ""), ""),
        subfleet_id=m.get("subfleet_id"),
        subfleet_name=subfleet_map.get(m.get("subfleet_id", ""), ""),
        operator_id=m.get("operator_id"),
        operator_name=operator_map.get(m.get("operator_id", ""), ""),
        created_at=m["created_at"]
    ) for m in machines]

@api_router.get("/machines/{machine_id}", response_model=MachineResponse)
async def get_machine(machine_id: str, current_user: dict = Depends(get_current_user)):
    machine = await db.machines.find_one({"id": machine_id}, {"_id": 0})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    category = await db.categories.find_one({"id": machine["category_id"]}, {"_id": 0})
    category_name = category["name"] if category else ""
    
    obra_name = ""
    if machine.get("obra_id"):
        obra = await db.obras.find_one({"id": machine["obra_id"]}, {"_id": 0})
        obra_name = obra["name"] if obra else ""
    
    fleet_name = ""
    if machine.get("fleet_id"):
        fleet = await db.fleets.find_one({"id": machine["fleet_id"]}, {"_id": 0})
        fleet_name = fleet["name"] if fleet else ""
    
    subfleet_name = ""
    if machine.get("subfleet_id"):
        subfleet = await db.subfleets.find_one({"id": machine["subfleet_id"]}, {"_id": 0})
        subfleet_name = subfleet["name"] if subfleet else ""
    
    operator_name = ""
    if machine.get("operator_id"):
        operator = await db.cadastros.find_one({"id": machine["operator_id"]}, {"_id": 0})
        operator_name = operator.get("nome_razao", "") if operator else ""
    
    return MachineResponse(
        id=machine["id"],
        name=machine["name"],
        plate=machine.get("plate", ""),
        category_id=machine["category_id"],
        category_name=category_name,
        brand=machine.get("brand", ""),
        model=machine.get("model", ""),
        year=machine.get("year"),
        notes=machine.get("notes", ""),
        status=machine.get("status", "patio"),
        obra_id=machine.get("obra_id"),
        obra_name=obra_name,
        fleet_id=machine.get("fleet_id"),
        fleet_name=fleet_name,
        subfleet_id=machine.get("subfleet_id"),
        subfleet_name=subfleet_name,
        operator_id=machine.get("operator_id"),
        operator_name=operator_name,
        created_at=machine["created_at"]
    )

@api_router.put("/machines/{machine_id}", response_model=MachineResponse)
async def update_machine(machine_id: str, machine: MachineCreate, current_user: dict = Depends(get_current_user)):
    existing = await db.machines.find_one({"id": machine_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    # Get category name
    category = await db.categories.find_one({"id": machine.category_id}, {"_id": 0})
    category_name = category["name"] if category else ""
    
    # Get obra name if specified
    obra_name = ""
    if machine.obra_id:
        obra = await db.obras.find_one({"id": machine.obra_id}, {"_id": 0})
        obra_name = obra["name"] if obra else ""
    
    update_doc = {
        "name": machine.name,
        "plate": (machine.plate or "").upper(),
        "category_id": machine.category_id,
        "subcategory_id": machine.subcategory_id,
        "brand": machine.brand or "",
        "model": machine.model or "",
        "year": machine.year,
        "notes": machine.notes or "",
        "obra_id": machine.obra_id,
        "fleet_id": machine.fleet_id,
        "subfleet_id": machine.subfleet_id,
        "operator_id": machine.operator_id,
        "identificador_tipo": machine.identificador_tipo,
        "identificador_numero": machine.identificador_numero,
    }
    await db.machines.update_one({"id": machine_id}, {"$set": update_doc})
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="editar",
        entity_type="máquina",
        entity_id=machine_id,
        entity_name=f"{machine.name} ({(machine.plate or '').upper()})",
        details=f"Dados anteriores: {existing['name']} ({existing.get('plate','')})"
    )
    
    # Buscar nomes para o response
    fleet_name = ""
    if machine.fleet_id:
        f = await db.fleets.find_one({"id": machine.fleet_id}, {"_id": 0, "name": 1})
        fleet_name = f["name"] if f else ""
    subfleet_name = ""
    if machine.subfleet_id:
        sf = await db.subfleets.find_one({"id": machine.subfleet_id}, {"_id": 0, "name": 1})
        subfleet_name = sf["name"] if sf else ""

    return MachineResponse(
        id=machine_id,
        name=machine.name,
        plate=(machine.plate or "").upper(),
        category_id=machine.category_id,
        category_name=category_name,
        subcategory_id=machine.subcategory_id,
        brand=machine.brand or "",
        model=machine.model or "",
        year=machine.year,
        notes=machine.notes or "",
        status=existing.get("status", "operational"),
        obra_id=machine.obra_id,
        obra_name=obra_name,
        fleet_id=machine.fleet_id,
        fleet_name=fleet_name,
        subfleet_id=machine.subfleet_id,
        subfleet_name=subfleet_name,
        operator_id=machine.operator_id,
        identificador_tipo=machine.identificador_tipo,
        identificador_numero=machine.identificador_numero,
        created_at=existing["created_at"]
    )

# ============ UPDATE MACHINE OBRA (TAG) ============

@api_router.patch("/machines/{machine_id}/obra", response_model=MachineResponse)
async def update_machine_obra(machine_id: str, update: MachineObraUpdate, current_user: dict = Depends(get_current_user)):
    existing = await db.machines.find_one({"id": machine_id})
    if not existing:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    # Validate obra if specified
    obra_name = ""
    if update.obra_id:
        obra = await db.obras.find_one({"id": update.obra_id}, {"_id": 0})
        if not obra:
            raise HTTPException(status_code=404, detail="Obra não encontrada")
        obra_name = obra["name"]
    
    await db.machines.update_one({"id": machine_id}, {"$set": {"obra_id": update.obra_id}})
    
    # Audit log
    old_obra = existing.get("obra_id")
    if update.obra_id:
        action_detail = f"Vinculada à obra: {obra_name}"
    else:
        action_detail = "Removida da obra"
    
    await create_audit_log(
        user=current_user,
        action="editar",
        entity_type="máquina",
        entity_id=machine_id,
        entity_name=f"{existing['name']} ({existing['plate']})",
        details=action_detail
    )
    
    category = await db.categories.find_one({"id": existing["category_id"]}, {"_id": 0})
    category_name = category["name"] if category else ""
    
    return MachineResponse(
        id=machine_id,
        name=existing["name"],
        plate=existing["plate"],
        category_id=existing["category_id"],
        category_name=category_name,
        brand=existing.get("brand", ""),
        model=existing.get("model", ""),
        year=existing.get("year"),
        notes=existing.get("notes", ""),
        status=existing.get("status", "operational"),
        obra_id=update.obra_id,
        obra_name=obra_name,
        created_at=existing["created_at"]
    )

@api_router.delete("/machines/{machine_id}")
async def delete_machine(machine_id: str, current_user: dict = Depends(get_current_user)):
    existing = await db.machines.find_one({"id": machine_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    await db.machines.delete_one({"id": machine_id})
    # Delete related maintenances
    await db.maintenances.delete_many({"machine_id": machine_id})
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="máquina",
        entity_id=machine_id,
        entity_name=f"{existing['name']} ({existing['plate']})"
    )
    
    return {"message": "Máquina removida com sucesso"}

# Endpoint para alterar status da máquina manualmente
class MachineStatusUpdate(BaseModel):
    status: str  # patio, operacional, manutencao

@api_router.patch("/machines/{machine_id}/status")
async def update_machine_status(machine_id: str, status_update: MachineStatusUpdate, current_user: dict = Depends(get_current_user)):
    existing = await db.machines.find_one({"id": machine_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    valid_statuses = ["patio", "operacional", "manutencao"]
    if status_update.status not in valid_statuses:
        raise HTTPException(status_code=400, detail=f"Status inválido. Use: {', '.join(valid_statuses)}")
    
    await db.machines.update_one(
        {"id": machine_id},
        {"$set": {"status": status_update.status}}
    )
    
    # Audit log
    status_labels = {"patio": "Pátio", "operacional": "Operacional", "manutencao": "Manutenção"}
    await create_audit_log(
        user=current_user,
        action="alterar_status",
        entity_type="máquina",
        entity_id=machine_id,
        entity_name=f"{existing['name']} ({existing.get('plate', '')}) - Status: {status_labels.get(status_update.status, status_update.status)}"
    )
    
    return {"message": f"Status alterado para {status_labels.get(status_update.status, status_update.status)}"}

# ============ HORIMETRO ROUTES ============

@api_router.get("/horimetro", response_model=List[HorimetroResponse])
async def list_horimetro(
    machine_id: Optional[str] = None,
    skip: int = 0,
    limit: int = 100,
    current_user: dict = Depends(get_current_user)
):
    """Lista todos os registros de horímetro"""
    query = {}
    if machine_id:
        query["machine_id"] = machine_id
    
    registros = await db.horimetro.find(query, {"_id": 0}).sort("data", -1).skip(skip).limit(limit).to_list(limit)
    
    # Adicionar nome da máquina
    for registro in registros:
        machine = await db.machines.find_one({"id": registro["machine_id"]}, {"_id": 0, "name": 1})
        registro["machine_name"] = machine["name"] if machine else "Máquina não encontrada"
    
    return registros

@api_router.get("/horimetro/{registro_id}", response_model=HorimetroResponse)
async def get_horimetro(registro_id: str, current_user: dict = Depends(get_current_user)):
    """Obtém um registro de horímetro específico"""
    registro = await db.horimetro.find_one({"id": registro_id}, {"_id": 0})
    if not registro:
        raise HTTPException(status_code=404, detail="Registro não encontrado")
    
    machine = await db.machines.find_one({"id": registro["machine_id"]}, {"_id": 0, "name": 1})
    registro["machine_name"] = machine["name"] if machine else "Máquina não encontrada"
    
    return registro

@api_router.post("/horimetro", response_model=HorimetroResponse)
async def create_horimetro(data: HorimetroCreate, current_user: dict = Depends(get_current_user)):
    """Cria um novo registro de horímetro"""
    # Verificar se a máquina existe
    machine = await db.machines.find_one({"id": data.machine_id}, {"_id": 0})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    # Calcular horas trabalhadas se não fornecido
    horas_trabalhadas = data.horas_trabalhadas
    if horas_trabalhadas is None:
        horas_trabalhadas = data.hora_final - data.hora_inicial
    
    registro_id = str(uuid.uuid4())
    registro_doc = {
        "id": registro_id,
        "machine_id": data.machine_id,
        "data": data.data,
        "hora_inicial": data.hora_inicial,
        "hora_final": data.hora_final,
        "horas_trabalhadas": horas_trabalhadas,
        "operador": data.operador or "",
        "observacoes": data.observacoes or "",
        "tipo_medicao": data.tipo_medicao or "hora",
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.horimetro.insert_one(registro_doc)
    
    # Atualizar horímetro atual da máquina
    await db.machines.update_one(
        {"id": data.machine_id},
        {"$set": {"horimetro_atual": data.hora_final}}
    )
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="horímetro",
        entity_id=registro_id,
        entity_name=f"{machine['name']} - {data.data}"
    )
    
    registro_doc["machine_name"] = machine["name"]
    return registro_doc

@api_router.put("/horimetro/{registro_id}", response_model=HorimetroResponse)
async def update_horimetro(registro_id: str, data: HorimetroCreate, current_user: dict = Depends(get_current_user)):
    """Atualiza um registro de horímetro"""
    existing = await db.horimetro.find_one({"id": registro_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Registro não encontrado")
    
    machine = await db.machines.find_one({"id": data.machine_id}, {"_id": 0})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    horas_trabalhadas = data.horas_trabalhadas
    if horas_trabalhadas is None:
        horas_trabalhadas = data.hora_final - data.hora_inicial
    
    update_doc = {
        "machine_id": data.machine_id,
        "data": data.data,
        "hora_inicial": data.hora_inicial,
        "hora_final": data.hora_final,
        "horas_trabalhadas": horas_trabalhadas,
        "operador": data.operador or "",
        "observacoes": data.observacoes or "",
        "tipo_medicao": data.tipo_medicao or "hora"
    }
    
    await db.horimetro.update_one({"id": registro_id}, {"$set": update_doc})
    
    # Atualizar horímetro atual da máquina se for o registro mais recente
    await db.machines.update_one(
        {"id": data.machine_id},
        {"$set": {"horimetro_atual": data.hora_final}}
    )
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="editar",
        entity_type="horímetro",
        entity_id=registro_id,
        entity_name=f"{machine['name']} - {data.data}"
    )
    
    updated = await db.horimetro.find_one({"id": registro_id}, {"_id": 0})
    updated["machine_name"] = machine["name"]
    return updated

@api_router.delete("/horimetro/{registro_id}")
async def delete_horimetro(registro_id: str, current_user: dict = Depends(get_current_user)):
    """Exclui um registro de horímetro"""
    existing = await db.horimetro.find_one({"id": registro_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Registro não encontrado")
    
    machine = await db.machines.find_one({"id": existing["machine_id"]}, {"_id": 0, "name": 1})
    
    await db.horimetro.delete_one({"id": registro_id})
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="horímetro",
        entity_id=registro_id,
        entity_name=f"{machine['name'] if machine else 'Máquina'} - {existing['data']}"
    )
    
    return {"message": "Registro excluído com sucesso"}

@api_router.get("/horimetro/machine/{machine_id}", response_model=List[HorimetroResponse])
async def get_horimetro_by_machine(machine_id: str, limit: int = 50, current_user: dict = Depends(get_current_user)):
    """Obtém registros de horímetro de uma máquina específica"""
    machine = await db.machines.find_one({"id": machine_id}, {"_id": 0, "name": 1})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    registros = await db.horimetro.find({"machine_id": machine_id}, {"_id": 0}).sort("data", -1).limit(limit).to_list(limit)
    
    for registro in registros:
        registro["machine_name"] = machine["name"]
    
    return registros

# ============ COMBUSTIVEL ROUTES ============

# --- Rotas para Veículos Abastecedores ---

@api_router.get("/combustivel/abastecedores")
async def list_veiculos_abastecedores(current_user: dict = Depends(get_current_user)):
    """Lista todos os veículos abastecedores cadastrados"""
    abastecedores = await db.veiculos_abastecedores.find({}, {"_id": 0}).to_list(100)
    
    result = []
    for abast in abastecedores:
        machine = await db.machines.find_one({"id": abast["machine_id"]}, {"_id": 0, "name": 1})
        abast["machine_name"] = machine["name"] if machine else "Máquina não encontrada"
        
        # Buscar nome do operador
        if abast.get("operador_id"):
            # Tentar funcionário do RH
            funcionario = await db.funcionarios.find_one({"id": abast["operador_id"]}, {"_id": 0, "nome": 1})
            if funcionario:
                abast["operador_nome"] = funcionario["nome"]
            else:
                # Tentar cadastro financeiro
                cadastro = await db.cadastros.find_one({"id": abast["operador_id"]}, {"_id": 0, "nome_razao": 1})
                abast["operador_nome"] = cadastro["nome_razao"] if cadastro else None
        
        result.append(abast)
    
    return result

@api_router.post("/combustivel/abastecedores")
async def create_veiculo_abastecedor(data: VeiculoAbastecedorCreate, current_user: dict = Depends(get_current_user)):
    """Cadastra uma máquina como veículo abastecedor"""
    machine = await db.machines.find_one({"id": data.machine_id}, {"_id": 0})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    # Verificar se já existe
    existing = await db.veiculos_abastecedores.find_one({"machine_id": data.machine_id}, {"_id": 0})
    if existing:
        raise HTTPException(status_code=400, detail="Esta máquina já está cadastrada como abastecedor")
    
    # Processar compartimentos de óleo
    compartimentos_processados = []
    if data.compartimentos_oleo:
        for comp in data.compartimentos_oleo:
            item = await db.stock_items.find_one({"id": comp.get("item_estoque_id")}, {"_id": 0, "name": 1})
            compartimentos_processados.append({
                "id": str(uuid.uuid4()),
                "item_estoque_id": comp.get("item_estoque_id"),
                "item_nome": item["name"] if item else "Item não encontrado",
                "unidade_medida": comp.get("unidade_medida", "L"),
                "capacidade": float(comp.get("capacidade", 0)),
                "quantidade_atual": float(comp.get("quantidade_atual", 0))
            })
    
    abastecedor_id = str(uuid.uuid4())
    abastecedor_doc = {
        "id": abastecedor_id,
        "machine_id": data.machine_id,
        "capacidade_diesel": data.capacidade_diesel,
        "capacidade_oleo": data.capacidade_oleo,
        "capacidade_graxa": data.capacidade_graxa,
        "litros_diesel": data.litros_diesel,
        "litros_oleo": data.litros_oleo,
        "litros_graxa": data.litros_graxa,
        "operador_id": data.operador_id,
        "compartimentos_oleo": compartimentos_processados,
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.veiculos_abastecedores.insert_one(abastecedor_doc)
    
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="veículo abastecedor",
        entity_id=abastecedor_id,
        entity_name=machine["name"]
    )
    
    # Buscar o documento criado sem o _id para evitar erro de serialização
    created = await db.veiculos_abastecedores.find_one({"id": abastecedor_id}, {"_id": 0})
    created["machine_name"] = machine["name"]
    return created

@api_router.put("/combustivel/abastecedores/{abastecedor_id}")
async def update_veiculo_abastecedor(abastecedor_id: str, data: VeiculoAbastecedorCreate, current_user: dict = Depends(get_current_user)):
    """Atualiza um veículo abastecedor"""
    existing = await db.veiculos_abastecedores.find_one({"id": abastecedor_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Veículo abastecedor não encontrado")
    
    machine = await db.machines.find_one({"id": data.machine_id}, {"_id": 0})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    # Processar compartimentos de óleo
    compartimentos_processados = []
    if data.compartimentos_oleo:
        for comp in data.compartimentos_oleo:
            item = await db.stock_items.find_one({"id": comp.get("item_estoque_id")}, {"_id": 0, "name": 1})
            compartimentos_processados.append({
                "id": comp.get("id") or str(uuid.uuid4()),
                "item_estoque_id": comp.get("item_estoque_id"),
                "item_nome": item["name"] if item else "Item não encontrado",
                "unidade_medida": comp.get("unidade_medida", "L"),
                "capacidade": float(comp.get("capacidade", 0)),
                "quantidade_atual": float(comp.get("quantidade_atual", 0))
            })
    
    update_doc = {
        "machine_id": data.machine_id,
        "capacidade_diesel": data.capacidade_diesel,
        "capacidade_oleo": data.capacidade_oleo,
        "capacidade_graxa": data.capacidade_graxa,
        "litros_diesel": data.litros_diesel,
        "litros_oleo": data.litros_oleo,
        "litros_graxa": data.litros_graxa,
        "operador_id": data.operador_id,
        "compartimentos_oleo": compartimentos_processados,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.veiculos_abastecedores.update_one({"id": abastecedor_id}, {"$set": update_doc})
    
    await create_audit_log(
        user=current_user,
        action="editar",
        entity_type="veículo abastecedor",
        entity_id=abastecedor_id,
        entity_name=machine["name"]
    )
    
    updated = await db.veiculos_abastecedores.find_one({"id": abastecedor_id}, {"_id": 0})
    updated["machine_name"] = machine["name"]
    return updated

@api_router.delete("/combustivel/abastecedores/{abastecedor_id}")
async def delete_veiculo_abastecedor(abastecedor_id: str, current_user: dict = Depends(get_current_user)):
    """Remove um veículo abastecedor"""
    existing = await db.veiculos_abastecedores.find_one({"id": abastecedor_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Veículo abastecedor não encontrado")
    
    machine = await db.machines.find_one({"id": existing["machine_id"]}, {"_id": 0, "name": 1})
    
    await db.veiculos_abastecedores.delete_one({"id": abastecedor_id})
    
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="veículo abastecedor",
        entity_id=abastecedor_id,
        entity_name=machine["name"] if machine else "Máquina"
    )
    
    return {"message": "Veículo abastecedor removido com sucesso"}

# --- Rotas para Registros de Combustível ---

@api_router.get("/combustivel", response_model=List[CombustivelResponse])
async def list_combustivel(
    machine_id: Optional[str] = None,
    tipo_registro: Optional[str] = None,
    skip: int = 0,
    limit: int = 100,
    current_user: dict = Depends(get_current_user)
):
    """Lista todos os registros de combustível"""
    query = {}
    if machine_id:
        query["machine_id"] = machine_id
    if tipo_registro:
        query["tipo_registro"] = tipo_registro
    
    registros = await db.combustivel.find(query, {"_id": 0}).sort("data", -1).skip(skip).limit(limit).to_list(limit)
    
    for registro in registros:
        machine = await db.machines.find_one({"id": registro["machine_id"]}, {"_id": 0, "name": 1})
        registro["machine_name"] = machine["name"] if machine else "Máquina não encontrada"
        
        # Buscar nome do veículo abastecedor se houver
        if registro.get("veiculo_abastecedor_id"):
            abast = await db.veiculos_abastecedores.find_one({"id": registro["veiculo_abastecedor_id"]}, {"_id": 0})
            if abast:
                abast_machine = await db.machines.find_one({"id": abast["machine_id"]}, {"_id": 0, "name": 1})
                registro["veiculo_abastecedor_nome"] = abast_machine["name"] if abast_machine else None
        
        # Buscar nome do operador
        if registro.get("operador_id"):
            funcionario = await db.funcionarios.find_one({"id": registro["operador_id"]}, {"_id": 0, "nome": 1})
            if funcionario:
                registro["operador_nome"] = funcionario["nome"]
            else:
                cadastro = await db.cadastros.find_one({"id": registro["operador_id"]}, {"_id": 0, "nome_razao": 1})
                registro["operador_nome"] = cadastro["nome_razao"] if cadastro else None
    
    return registros

@api_router.get("/combustivel/{registro_id}", response_model=CombustivelResponse)
async def get_combustivel(registro_id: str, current_user: dict = Depends(get_current_user)):
    """Obtém um registro de combustível específico"""
    registro = await db.combustivel.find_one({"id": registro_id}, {"_id": 0})
    if not registro:
        raise HTTPException(status_code=404, detail="Registro não encontrado")
    
    machine = await db.machines.find_one({"id": registro["machine_id"]}, {"_id": 0, "name": 1})
    registro["machine_name"] = machine["name"] if machine else "Máquina não encontrada"
    
    return registro

@api_router.post("/combustivel", response_model=CombustivelResponse)
async def create_combustivel(data: CombustivelCreate, current_user: dict = Depends(get_current_user)):
    """Cria um novo registro de combustível"""
    machine = await db.machines.find_one({"id": data.machine_id}, {"_id": 0})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    registro_id = str(uuid.uuid4())
    registro_doc = {
        "id": registro_id,
        "machine_id": data.machine_id,
        "data": data.data,
        "tipo_registro": data.tipo_registro,
        "tipo_medicao": data.tipo_medicao,
        "hora_km_inicial": data.hora_km_inicial or 0,
        "litros_diesel": data.litros_diesel,
        "litros_oleo": data.litros_oleo,
        "litros_graxa": data.litros_graxa,
        "fonte_abastecimento": data.fonte_abastecimento,
        "veiculo_abastecedor_id": data.veiculo_abastecedor_id,
        "operador_id": data.operador_id,
        "observacoes": data.observacoes or "",
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.combustivel.insert_one(registro_doc)
    
    # Se foi abastecido por um veículo abastecedor interno, descontar do tanque
    if data.tipo_registro == "abastecido" and data.fonte_abastecimento == "interno" and data.veiculo_abastecedor_id:
        abastecedor = await db.veiculos_abastecedores.find_one({"id": data.veiculo_abastecedor_id}, {"_id": 0})
        if abastecedor:
            new_diesel = max(0, abastecedor.get("litros_diesel", 0) - data.litros_diesel)
            new_oleo = max(0, abastecedor.get("litros_oleo", 0) - data.litros_oleo)
            new_graxa = max(0, abastecedor.get("litros_graxa", 0) - data.litros_graxa)
            
            await db.veiculos_abastecedores.update_one(
                {"id": data.veiculo_abastecedor_id},
                {"$set": {
                    "litros_diesel": new_diesel,
                    "litros_oleo": new_oleo,
                    "litros_graxa": new_graxa,
                    "updated_at": datetime.now(timezone.utc).isoformat()
                }}
            )
    
    # Se for registro de abastecedor (entrada de combustível no tanque), adicionar ao tanque
    if data.tipo_registro == "abastecedor":
        # Encontrar o abastecedor pela máquina
        abastecedor = await db.veiculos_abastecedores.find_one({"machine_id": data.machine_id}, {"_id": 0})
        if abastecedor:
            new_diesel = abastecedor.get("litros_diesel", 0) + data.litros_diesel
            new_oleo = abastecedor.get("litros_oleo", 0) + data.litros_oleo
            new_graxa = abastecedor.get("litros_graxa", 0) + data.litros_graxa
            
            await db.veiculos_abastecedores.update_one(
                {"id": abastecedor["id"]},
                {"$set": {
                    "litros_diesel": new_diesel,
                    "litros_oleo": new_oleo,
                    "litros_graxa": new_graxa,
                    "updated_at": datetime.now(timezone.utc).isoformat()
                }}
            )
    
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="combustível",
        entity_id=registro_id,
        entity_name=f"{machine['name']} - {data.data}"
    )
    
    # Buscar o documento criado sem o _id para evitar erro de serialização
    created = await db.combustivel.find_one({"id": registro_id}, {"_id": 0})
    created["machine_name"] = machine["name"]
    return created


@api_router.put("/combustivel/{registro_id}")
async def update_combustivel(registro_id: str, data: CombustivelCreate, current_user: dict = Depends(get_current_user)):
    """Atualiza um registro de combustível existente"""
    existing = await db.combustivel.find_one({"id": registro_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Registro não encontrado")
    
    machine = await db.machines.find_one({"id": data.machine_id}, {"_id": 0, "name": 1})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    # Buscar nomes relacionados
    operador_nome = None
    if data.operador_id:
        operador = await db.funcionarios.find_one({"id": data.operador_id}, {"_id": 0, "nome": 1})
        if not operador:
            operador = await db.cadastros.find_one({"id": data.operador_id}, {"_id": 0, "nome_razao": 1})
            if operador:
                operador_nome = operador.get("nome_razao")
        else:
            operador_nome = operador.get("nome")
    
    veiculo_abastecedor_nome = None
    if data.veiculo_abastecedor_id:
        abastecedor = await db.veiculos_abastecedores.find_one({"id": data.veiculo_abastecedor_id}, {"_id": 0, "machine_id": 1})
        if abastecedor:
            abast_machine = await db.machines.find_one({"id": abastecedor["machine_id"]}, {"_id": 0, "name": 1})
            if abast_machine:
                veiculo_abastecedor_nome = abast_machine["name"]
    
    posto_nome = None
    if data.posto_id:
        posto = await db.cadastros.find_one({"id": data.posto_id}, {"_id": 0, "nome_razao": 1})
        if posto:
            posto_nome = posto.get("nome_razao")
    
    update_doc = {
        "machine_id": data.machine_id,
        "data": data.data,
        "tipo_registro": data.tipo_registro,
        "tipo_medicao": data.tipo_medicao,
        "hora_km_inicial": data.hora_km_inicial,
        "litros_diesel": data.litros_diesel,
        "litros_oleo": data.litros_oleo,
        "litros_graxa": data.litros_graxa,
        "fonte_abastecimento": data.fonte_abastecimento,
        "veiculo_abastecedor_id": data.veiculo_abastecedor_id,
        "veiculo_abastecedor_nome": veiculo_abastecedor_nome,
        "posto_id": data.posto_id,
        "posto_nome": posto_nome,
        "operador_id": data.operador_id,
        "operador_nome": operador_nome,
        "observacoes": data.observacoes,
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.combustivel.update_one({"id": registro_id}, {"$set": update_doc})
    
    await create_audit_log(
        user=current_user,
        action="editar",
        entity_type="combustível",
        entity_id=registro_id,
        entity_name=f"{machine['name']} - {data.data}"
    )
    
    updated = await db.combustivel.find_one({"id": registro_id}, {"_id": 0})
    updated["machine_name"] = machine["name"]
    return updated


@api_router.delete("/combustivel/{registro_id}")
async def delete_combustivel(registro_id: str, current_user: dict = Depends(get_current_user)):
    """Exclui um registro de combustível"""
    existing = await db.combustivel.find_one({"id": registro_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Registro não encontrado")
    
    machine = await db.machines.find_one({"id": existing["machine_id"]}, {"_id": 0, "name": 1})
    
    await db.combustivel.delete_one({"id": registro_id})
    
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="combustível",
        entity_id=registro_id,
        entity_name=f"{machine['name'] if machine else 'Máquina'} - {existing['data']}"
    )
    
    return {"message": "Registro excluído com sucesso"}

@api_router.get("/combustivel/machine/{machine_id}")
async def get_combustivel_by_machine(machine_id: str, limit: int = 50, current_user: dict = Depends(get_current_user)):
    """Obtém registros de combustível de uma máquina específica"""
    machine = await db.machines.find_one({"id": machine_id}, {"_id": 0, "name": 1})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    registros = await db.combustivel.find({"machine_id": machine_id}, {"_id": 0}).sort("data", -1).limit(limit).to_list(limit)
    
    for registro in registros:
        registro["machine_name"] = machine["name"]
    
    return registros

# Rota para buscar operadores (RH + Cadastros Financeiro)
@api_router.get("/operadores")
async def list_operadores(current_user: dict = Depends(get_current_user)):
    """Lista todos os operadores disponíveis (funcionários RH + cadastros financeiro)"""
    operadores = []
    
    # Buscar funcionários do RH
    funcionarios = await db.funcionarios.find({}, {"_id": 0, "id": 1, "nome": 1, "cargo": 1}).to_list(500)
    for f in funcionarios:
        operadores.append({
            "id": f["id"],
            "nome": f["nome"],
            "tipo": "rh",
            "cargo": f.get("cargo", "")
        })
    
    # Buscar cadastros do financeiro (nome_razao é o campo correto)
    cadastros = await db.cadastros.find({}, {"_id": 0, "id": 1, "nome_razao": 1, "tipo_cadastro": 1}).to_list(500)
    for c in cadastros:
        operadores.append({
            "id": c["id"],
            "nome": c.get("nome_razao", ""),
            "tipo": "cadastro",
            "cargo": c.get("tipo_cadastro", "")
        })
    
    return operadores

# ============ MAINTENANCE ROUTES ============

@api_router.post("/maintenances", response_model=MaintenanceResponse)
async def create_maintenance(maintenance: MaintenanceCreate, current_user: dict = Depends(get_current_user)):
    # Check if machine exists
    machine = await db.machines.find_one({"id": maintenance.machine_id}, {"_id": 0})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    maintenance_id = str(uuid.uuid4())
    maintenance_doc = {
        "id": maintenance_id,
        "machine_id": maintenance.machine_id,
        "part_name": maintenance.part_name,
        "replacement_date": maintenance.replacement_date,
        "part_value": maintenance.part_value,
        "maintenance_type": maintenance.maintenance_type,
        "description": maintenance.description or "",
        "is_oil_change": maintenance.is_oil_change,
        "photos": [],
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.maintenances.insert_one(maintenance_doc)
    
    # Update machine status to maintenance if corrective
    if maintenance.maintenance_type == "corretiva":
        await db.machines.update_one({"id": maintenance.machine_id}, {"$set": {"status": "maintenance"}})
    
    # If oil change, reset hours counter for this machine
    if maintenance.is_oil_change:
        await db.machines.update_one(
            {"id": maintenance.machine_id}, 
            {"$set": {"last_oil_change_date": maintenance.replacement_date, "hours_since_oil_change": 0}}
        )
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="manutenção",
        entity_id=maintenance_id,
        entity_name=f"{maintenance.part_name} - {machine['name']}",
        details=f"Valor: R$ {maintenance.part_value:.2f}, Tipo: {maintenance.maintenance_type}"
    )
    
    return MaintenanceResponse(
        id=maintenance_id,
        machine_id=maintenance.machine_id,
        machine_name=machine["name"],
        machine_plate=machine["plate"],
        part_name=maintenance.part_name,
        replacement_date=maintenance.replacement_date,
        part_value=maintenance.part_value,
        maintenance_type=maintenance.maintenance_type,
        description=maintenance.description or "",
        is_oil_change=maintenance.is_oil_change,
        photos=[],
        created_at=maintenance_doc["created_at"]
    )

@api_router.get("/maintenances", response_model=List[MaintenanceResponse])
async def get_maintenances(machine_id: Optional[str] = None, oil_changes_only: bool = False, current_user: dict = Depends(get_current_user)):
    query = {}
    if machine_id:
        query["machine_id"] = machine_id
    if oil_changes_only:
        query["is_oil_change"] = True
    
    maintenances = await db.maintenances.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    # Get all machines
    machines = await db.machines.find({}, {"_id": 0}).to_list(1000)
    machine_map = {m["id"]: {"name": m["name"], "plate": m["plate"]} for m in machines}
    
    return [MaintenanceResponse(
        id=m["id"],
        machine_id=m["machine_id"],
        machine_name=machine_map.get(m["machine_id"], {}).get("name", ""),
        machine_plate=machine_map.get(m["machine_id"], {}).get("plate", ""),
        part_name=m["part_name"],
        replacement_date=m["replacement_date"],
        part_value=m["part_value"],
        maintenance_type=m["maintenance_type"],
        description=m.get("description", ""),
        is_oil_change=m.get("is_oil_change", False),
        photos=m.get("photos", []),
        created_at=m["created_at"]
    ) for m in maintenances]

@api_router.get("/maintenances/{maintenance_id}", response_model=MaintenanceResponse)
async def get_maintenance(maintenance_id: str, current_user: dict = Depends(get_current_user)):
    maintenance = await db.maintenances.find_one({"id": maintenance_id}, {"_id": 0})
    if not maintenance:
        raise HTTPException(status_code=404, detail="Manutenção não encontrada")
    
    machine = await db.machines.find_one({"id": maintenance["machine_id"]}, {"_id": 0})
    
    return MaintenanceResponse(
        id=maintenance["id"],
        machine_id=maintenance["machine_id"],
        machine_name=machine["name"] if machine else "",
        machine_plate=machine["plate"] if machine else "",
        part_name=maintenance["part_name"],
        replacement_date=maintenance["replacement_date"],
        part_value=maintenance["part_value"],
        maintenance_type=maintenance["maintenance_type"],
        description=maintenance.get("description", ""),
        is_oil_change=maintenance.get("is_oil_change", False),
        photos=maintenance.get("photos", []),
        created_at=maintenance["created_at"]
    )

@api_router.delete("/maintenances/{maintenance_id}")
async def delete_maintenance(maintenance_id: str, current_user: dict = Depends(get_current_user)):
    existing = await db.maintenances.find_one({"id": maintenance_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Manutenção não encontrada")
    
    await db.maintenances.delete_one({"id": maintenance_id})
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="manutenção",
        entity_id=maintenance_id,
        entity_name=existing["part_name"]
    )
    
    return {"message": "Manutenção removida com sucesso"}

# ============ PHOTO UPLOAD ============

@api_router.post("/maintenances/{maintenance_id}/photos")
async def upload_photo(
    maintenance_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    maintenance = await db.maintenances.find_one({"id": maintenance_id})
    if not maintenance:
        raise HTTPException(status_code=404, detail="Manutenção não encontrada")
    
    # Read file and convert to base64
    contents = await file.read()
    base64_image = base64.b64encode(contents).decode('utf-8')
    photo_data = f"data:{file.content_type};base64,{base64_image}"
    
    # Add photo to maintenance
    photos = maintenance.get("photos", [])
    photos.append(photo_data)
    await db.maintenances.update_one({"id": maintenance_id}, {"$set": {"photos": photos}})
    
    return {"message": "Foto adicionada com sucesso", "photo": photo_data}

@api_router.delete("/maintenances/{maintenance_id}/photos/{photo_index}")
async def delete_photo(
    maintenance_id: str,
    photo_index: int,
    current_user: dict = Depends(get_current_user)
):
    maintenance = await db.maintenances.find_one({"id": maintenance_id})
    if not maintenance:
        raise HTTPException(status_code=404, detail="Manutenção não encontrada")
    
    photos = maintenance.get("photos", [])
    if photo_index < 0 or photo_index >= len(photos):
        raise HTTPException(status_code=400, detail="Índice de foto inválido")
    
    photos.pop(photo_index)
    await db.maintenances.update_one({"id": maintenance_id}, {"$set": {"photos": photos}})
    
    return {"message": "Foto removida com sucesso"}

# ============ USAGE LOG / OIL CHANGE ROUTES ============

@api_router.post("/usage-logs", response_model=UsageLogResponse)
async def create_usage_log(log: UsageLogCreate, current_user: dict = Depends(get_current_user)):
    # Check if machine exists
    machine = await db.machines.find_one({"id": log.machine_id}, {"_id": 0})
    if not machine:
        raise HTTPException(status_code=404, detail="Máquina não encontrada")
    
    log_id = str(uuid.uuid4())
    log_doc = {
        "id": log_id,
        "machine_id": log.machine_id,
        "hours": log.hours,
        "notes": log.notes or "",
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.usage_logs.insert_one(log_doc)
    
    # Update machine's hours since oil change
    current_hours = machine.get("hours_since_oil_change", 0)
    await db.machines.update_one(
        {"id": log.machine_id},
        {"$set": {"hours_since_oil_change": current_hours + log.hours}}
    )
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type="registro de uso",
        entity_id=log_id,
        entity_name=f"{machine['name']} - {log.hours}h"
    )
    
    return UsageLogResponse(
        id=log_id,
        machine_id=log.machine_id,
        machine_name=machine["name"],
        machine_plate=machine["plate"],
        hours=log.hours,
        notes=log.notes or "",
        created_at=log_doc["created_at"]
    )

@api_router.get("/usage-logs", response_model=List[UsageLogResponse])
async def get_usage_logs(machine_id: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    query = {}
    if machine_id:
        query["machine_id"] = machine_id
    
    logs = await db.usage_logs.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    # Get machines
    machines = await db.machines.find({}, {"_id": 0}).to_list(1000)
    machine_map = {m["id"]: {"name": m["name"], "plate": m["plate"]} for m in machines}
    
    return [UsageLogResponse(
        id=l["id"],
        machine_id=l["machine_id"],
        machine_name=machine_map.get(l["machine_id"], {}).get("name", ""),
        machine_plate=machine_map.get(l["machine_id"], {}).get("plate", ""),
        hours=l["hours"],
        notes=l.get("notes", ""),
        created_at=l["created_at"]
    ) for l in logs]

@api_router.delete("/usage-logs/{log_id}")
async def delete_usage_log(log_id: str, current_user: dict = Depends(get_current_user)):
    """Exclui um registro de uso e atualiza as horas da máquina"""
    # Find the usage log
    log = await db.usage_logs.find_one({"id": log_id}, {"_id": 0})
    if not log:
        raise HTTPException(status_code=404, detail="Registro de uso não encontrado")
    
    # Get machine info
    machine = await db.machines.find_one({"id": log["machine_id"]}, {"_id": 0})
    
    # Subtract hours from machine's total
    if machine:
        current_hours = machine.get("hours_since_oil_change", 0)
        new_hours = max(0, current_hours - log["hours"])  # Prevent negative
        await db.machines.update_one(
            {"id": log["machine_id"]},
            {"$set": {"hours_since_oil_change": new_hours}}
        )
    
    # Delete the usage log
    await db.usage_logs.delete_one({"id": log_id})
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type="registro de uso",
        entity_id=log_id,
        entity_name=f"{machine.get('name', 'N/A') if machine else 'N/A'} - {log['hours']}h"
    )
    
    return {"message": "Registro de uso excluído com sucesso"}

@api_router.get("/oil-change-status", response_model=List[OilChangeStatusResponse])
async def get_oil_change_status(current_user: dict = Depends(get_current_user)):
    machines = await db.machines.find({}, {"_id": 0}).to_list(1000)
    
    result = []
    today = datetime.now(timezone.utc)
    
    for machine in machines:
        hours_since_change = machine.get("hours_since_oil_change", 0)
        hours_remaining = 500 - hours_since_change
        
        last_oil_change_date = machine.get("last_oil_change_date")
        if last_oil_change_date:
            try:
                last_change = datetime.fromisoformat(last_oil_change_date.replace('Z', '+00:00'))
                if last_change.tzinfo is None:
                    last_change = last_change.replace(tzinfo=timezone.utc)
            except:
                last_change = datetime.strptime(last_oil_change_date, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            days_since_change = (today - last_change).days
            days_remaining = 365 - days_since_change
        else:
            days_since_change = 0
            days_remaining = 365
        
        # Check if needs alert
        needs_alert = False
        alert_reason = None
        
        if hours_remaining <= 50:
            needs_alert = True
            alert_reason = f"Faltam apenas {hours_remaining:.0f} horas para atingir 500h de uso"
        elif days_remaining <= 60 and hours_remaining > 50:
            needs_alert = True
            alert_reason = f"Faltam {days_remaining} dias para completar 1 ano desde a última troca"
        
        result.append(OilChangeStatusResponse(
            machine_id=machine["id"],
            machine_name=machine["name"],
            machine_plate=machine["plate"],
            last_oil_change_date=last_oil_change_date,
            hours_since_change=hours_since_change,
            hours_remaining=max(0, hours_remaining),
            days_since_change=days_since_change,
            days_remaining=max(0, days_remaining),
            needs_alert=needs_alert,
            alert_reason=alert_reason
        ))
    
    return result

@api_router.get("/notifications", response_model=List[NotificationResponse])
async def get_notifications(current_user: dict = Depends(get_current_user)):
    machines = await db.machines.find({}, {"_id": 0}).to_list(1000)
    
    notifications = []
    today = datetime.now(timezone.utc)
    
    for machine in machines:
        hours_since_change = machine.get("hours_since_oil_change", 0)
        hours_remaining = 500 - hours_since_change
        
        last_oil_change_date = machine.get("last_oil_change_date")
        if last_oil_change_date:
            try:
                last_change = datetime.fromisoformat(last_oil_change_date.replace('Z', '+00:00'))
                if last_change.tzinfo is None:
                    last_change = last_change.replace(tzinfo=timezone.utc)
            except:
                last_change = datetime.strptime(last_oil_change_date, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            days_since_change = (today - last_change).days
            days_remaining = 365 - days_since_change
        else:
            days_since_change = 0
            days_remaining = 365
        
        # Alert for hours
        if hours_remaining <= 50 and hours_remaining > 0:
            notifications.append(NotificationResponse(
                id=f"oil-hours-{machine['id']}",
                machine_id=machine["id"],
                machine_name=machine["name"],
                machine_plate=machine["plate"],
                notification_type="oil_change_hours",
                message=f"⚠️ Troca de óleo necessária em breve! Restam apenas {hours_remaining:.0f} horas de uso.",
                hours_remaining=hours_remaining,
                days_remaining=None,
                created_at=today.isoformat()
            ))
        elif hours_remaining <= 0:
            notifications.append(NotificationResponse(
                id=f"oil-hours-urgent-{machine['id']}",
                machine_id=machine["id"],
                machine_name=machine["name"],
                machine_plate=machine["plate"],
                notification_type="oil_change_urgent",
                message=f"🚨 URGENTE: Limite de 500 horas atingido! Troque o óleo imediatamente.",
                hours_remaining=0,
                days_remaining=None,
                created_at=today.isoformat()
            ))
        
        # Alert for time (only if hours haven't triggered yet)
        if hours_remaining > 50:
            if days_remaining <= 0:
                notifications.append(NotificationResponse(
                    id=f"oil-time-urgent-{machine['id']}",
                    machine_id=machine["id"],
                    machine_name=machine["name"],
                    machine_plate=machine["plate"],
                    notification_type="oil_change_time_urgent",
                    message=f"🚨 URGENTE: Completou 1 ano desde a última troca de óleo! Troque imediatamente.",
                    hours_remaining=hours_remaining,
                    days_remaining=0,
                    created_at=today.isoformat()
                ))
            elif days_remaining <= 60:
                notifications.append(NotificationResponse(
                    id=f"oil-time-{machine['id']}",
                    machine_id=machine["id"],
                    machine_name=machine["name"],
                    machine_plate=machine["plate"],
                    notification_type="oil_change_time",
                    message=f"⚠️ Faltam {days_remaining} dias para completar 1 ano desde a última troca de óleo.",
                    hours_remaining=hours_remaining,
                    days_remaining=days_remaining,
                    created_at=today.isoformat()
                ))
    
    # Add low stock alerts (quantity < 5)
    stock_items = await db.stock_items.find({}, {"_id": 0}).to_list(1000)
    for item in stock_items:
        if item["quantity"] < 5:
            if item["quantity"] <= 0:
                notifications.append(NotificationResponse(
                    id=f"stock-empty-{item['id']}",
                    machine_id="",
                    machine_name=item["name"],
                    machine_plate=item.get("code", ""),
                    notification_type="stock_empty",
                    message=f"🚨 ESTOQUE ZERADO: {item['name']} está sem estoque! Reponha imediatamente.",
                    hours_remaining=None,
                    days_remaining=None,
                    created_at=today.isoformat()
                ))
            else:
                notifications.append(NotificationResponse(
                    id=f"stock-low-{item['id']}",
                    machine_id="",
                    machine_name=item["name"],
                    machine_plate=item.get("code", ""),
                    notification_type="stock_low",
                    message=f"⚠️ Estoque baixo: {item['name']} tem apenas {item['quantity']:.0f} {item.get('unit', 'un')} restantes.",
                    hours_remaining=None,
                    days_remaining=None,
                    created_at=today.isoformat()
                ))
    
    # Add fuel low alerts for veiculos abastecedores (tankers)
    veiculos_abastecedores = await db.veiculos_abastecedores.find({}, {"_id": 0}).to_list(100)
    for veiculo in veiculos_abastecedores:
        capacidade = veiculo.get("capacidade_diesel", 0)
        nivel_atual = veiculo.get("litros_diesel", 0)  # Campo correto é litros_diesel
        
        # Buscar nome da máquina
        machine = await db.machines.find_one({"id": veiculo.get("machine_id")}, {"_id": 0})
        machine_name = machine.get("name", "Veículo Tanque") if machine else "Veículo Tanque"
        machine_plate = machine.get("plate", "") if machine else ""
        
        if capacidade > 0:
            porcentagem = (nivel_atual / capacidade) * 100
            
            # Critical: less than 10% fuel
            if porcentagem < 10:
                notifications.append(NotificationResponse(
                    id=f"fuel-critical-{veiculo['id']}",
                    machine_id=veiculo.get("machine_id", ""),
                    machine_name=machine_name,
                    machine_plate=machine_plate,
                    notification_type="fuel_critical",
                    message=f"🚨 COMBUSTÍVEL CRÍTICO: {machine_name} está com apenas {porcentagem:.0f}% ({nivel_atual:.0f}L de {capacidade:.0f}L). Reabasteça urgentemente!",
                    hours_remaining=None,
                    days_remaining=None,
                    created_at=today.isoformat()
                ))
            # Warning: less than 25% fuel
            elif porcentagem < 25:
                notifications.append(NotificationResponse(
                    id=f"fuel-low-{veiculo['id']}",
                    machine_id=veiculo.get("machine_id", ""),
                    machine_name=machine_name,
                    machine_plate=machine_plate,
                    notification_type="fuel_low",
                    message=f"⚠️ Combustível baixo: {machine_name} está com {porcentagem:.0f}% ({nivel_atual:.0f}L de {capacidade:.0f}L). Considere reabastecer.",
                    hours_remaining=None,
                    days_remaining=None,
                    created_at=today.isoformat()
                ))
    
    # Sort by urgency
    def sort_key(n):
        if "urgent" in n.notification_type or "empty" in n.notification_type or "critical" in n.notification_type:
            return 0
        return 1
    
    notifications.sort(key=sort_key)
    
    return notifications

# ============ BALANCE / FINANCIAL REPORTS ============

class MachineExpenseResponse(BaseModel):
    machine_id: str
    machine_name: str
    machine_plate: str
    category_name: str
    total_maintenances: int
    total_spent: float
    preventive_spent: float
    corrective_spent: float
    last_maintenance_date: Optional[str] = None

class MonthlyExpenseResponse(BaseModel):
    month: str
    year: int
    total_spent: float
    maintenance_count: int

class BalanceResponse(BaseModel):
    total_spent: float
    total_maintenances: int
    preventive_total: float
    corrective_total: float
    preventive_count: int
    corrective_count: int
    average_per_maintenance: float
    expenses_by_machine: List[MachineExpenseResponse]
    expenses_by_month: List[MonthlyExpenseResponse]

@api_router.get("/balance", response_model=BalanceResponse)
async def get_balance(current_user: dict = Depends(get_current_user)):
    # Get all maintenances
    maintenances = await db.maintenances.find({}, {"_id": 0}).to_list(10000)
    
    # Get all machines with categories
    machines = await db.machines.find({}, {"_id": 0}).to_list(1000)
    categories = await db.categories.find({}, {"_id": 0}).to_list(100)
    category_map = {c["id"]: c["name"] for c in categories}
    machine_map = {m["id"]: {
        "name": m["name"], 
        "plate": m["plate"],
        "category_name": category_map.get(m.get("category_id", ""), "")
    } for m in machines}
    
    # Calculate totals
    total_spent = sum(m["part_value"] for m in maintenances)
    total_maintenances = len(maintenances)
    preventive = [m for m in maintenances if m["maintenance_type"] == "preventiva"]
    corrective = [m for m in maintenances if m["maintenance_type"] == "corretiva"]
    preventive_total = sum(m["part_value"] for m in preventive)
    corrective_total = sum(m["part_value"] for m in corrective)
    average_per_maintenance = total_spent / total_maintenances if total_maintenances > 0 else 0
    
    # Expenses by machine
    machine_expenses = {}
    for m in maintenances:
        mid = m["machine_id"]
        if mid not in machine_expenses:
            machine_expenses[mid] = {
                "total_maintenances": 0,
                "total_spent": 0,
                "preventive_spent": 0,
                "corrective_spent": 0,
                "last_maintenance_date": None
            }
        machine_expenses[mid]["total_maintenances"] += 1
        machine_expenses[mid]["total_spent"] += m["part_value"]
        if m["maintenance_type"] == "preventiva":
            machine_expenses[mid]["preventive_spent"] += m["part_value"]
        else:
            machine_expenses[mid]["corrective_spent"] += m["part_value"]
        
        # Track last maintenance
        if machine_expenses[mid]["last_maintenance_date"] is None or m["replacement_date"] > machine_expenses[mid]["last_maintenance_date"]:
            machine_expenses[mid]["last_maintenance_date"] = m["replacement_date"]
    
    expenses_by_machine = [
        MachineExpenseResponse(
            machine_id=mid,
            machine_name=machine_map.get(mid, {}).get("name", "Máquina Removida"),
            machine_plate=machine_map.get(mid, {}).get("plate", "-"),
            category_name=machine_map.get(mid, {}).get("category_name", ""),
            total_maintenances=data["total_maintenances"],
            total_spent=data["total_spent"],
            preventive_spent=data["preventive_spent"],
            corrective_spent=data["corrective_spent"],
            last_maintenance_date=data["last_maintenance_date"]
        )
        for mid, data in machine_expenses.items()
    ]
    expenses_by_machine.sort(key=lambda x: x.total_spent, reverse=True)
    
    # Expenses by month
    month_expenses = {}
    for m in maintenances:
        try:
            date = datetime.fromisoformat(m["replacement_date"].replace('Z', '+00:00'))
        except:
            date = datetime.strptime(m["replacement_date"], "%Y-%m-%d")
        key = f"{date.year}-{date.month:02d}"
        if key not in month_expenses:
            month_expenses[key] = {"total_spent": 0, "count": 0, "year": date.year, "month": date.month}
        month_expenses[key]["total_spent"] += m["part_value"]
        month_expenses[key]["count"] += 1
    
    month_names = ["", "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", 
                   "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    
    expenses_by_month = [
        MonthlyExpenseResponse(
            month=month_names[data["month"]],
            year=data["year"],
            total_spent=data["total_spent"],
            maintenance_count=data["count"]
        )
        for key, data in sorted(month_expenses.items(), reverse=True)
    ]
    
    return BalanceResponse(
        total_spent=total_spent,
        total_maintenances=total_maintenances,
        preventive_total=preventive_total,
        corrective_total=corrective_total,
        preventive_count=len(preventive),
        corrective_count=len(corrective),
        average_per_maintenance=average_per_maintenance,
        expenses_by_machine=expenses_by_machine,
        expenses_by_month=expenses_by_month[:12]  # Last 12 months
    )

# ============ DASHBOARD / STOCK / OBRAS / MEDIÇÕES ============
# (Dashboard -> /app/backend/routes/dashboard.py)
# (Stock    -> /app/backend/routes/stock.py    - já existia, duplicados removidos)
# (Obras    -> /app/backend/routes/obras.py    - já existia, duplicados removidos)
# (Medições -> /app/backend/routes/medicoes.py)
# Refactor Sessão 32 Fase 2 Parte 2

# ============ AUDIT LOG ROUTES ============

@api_router.get("/audit-logs", response_model=List[AuditLogResponse])
async def get_audit_logs(
    entity_type: Optional[str] = None,
    user_id: Optional[str] = None,
    limit: int = 100,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if entity_type:
        query["entity_type"] = entity_type
    if user_id:
        query["user_id"] = user_id
    
    logs = await db.audit_logs.find(query, {"_id": 0}).sort("created_at", -1).limit(limit).to_list(limit)
    
    return [AuditLogResponse(
        id=l.get("id", ""),
        user_id=l.get("user_id", ""),
        user_name=l.get("user_name", "Usuário"),
        user_email=l.get("user_email", ""),
        action=l.get("action", ""),
        entity_type=l.get("entity_type", ""),
        entity_id=l.get("entity_id", ""),
        entity_name=l.get("entity_name", ""),
        details=l.get("details", ""),
        created_at=l.get("created_at", "")
    ) for l in logs]

# ============ ROOT ============

@api_router.get("/")
async def root():
    return {"message": "CRA Construtora API"}

# ============ ADMIN MODELS - COMPLETOS ============

# --- Cadastro Unificado de Clientes/Fornecedores ---
class CadastroCreate(BaseModel):
    # Tipo de cadastro
    tipo_cadastro: str = "cliente"  # cliente, fornecedor, cli_forn, transportador, funcionario
    tipo_pessoa: str = "PF"  # PF ou PJ
    status: str = "ativo"  # ativo, inativo
    
    # Dados principais
    nome_razao: str
    apelido_fantasia: Optional[str] = None
    cpf_cnpj: Optional[str] = None
    rg_ie: Optional[str] = None  # RG ou Inscrição Estadual
    
    # Contato
    telefone: Optional[str] = None
    celular: Optional[str] = None
    email: Optional[str] = None
    
    # Endereço
    cep: Optional[str] = None
    endereco: Optional[str] = None
    numero: Optional[str] = None
    complemento: Optional[str] = None
    bairro: Optional[str] = None
    cidade: Optional[str] = None
    uf: Optional[str] = None
    
    # Dados adicionais
    grupo: Optional[str] = None
    rota: Optional[str] = None
    vendedor: Optional[str] = None
    limite_credito: Optional[float] = None
    observacoes: Optional[str] = None

class CadastroResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    codigo: int
    tipo_cadastro: str
    tipo_pessoa: str
    status: str
    nome_razao: str
    apelido_fantasia: Optional[str] = None
    cpf_cnpj: Optional[str] = None
    rg_ie: Optional[str] = None
    telefone: Optional[str] = None
    celular: Optional[str] = None
    email: Optional[str] = None
    cep: Optional[str] = None
    endereco: Optional[str] = None
    numero: Optional[str] = None
    complemento: Optional[str] = None
    bairro: Optional[str] = None
    cidade: Optional[str] = None
    uf: Optional[str] = None
    grupo: Optional[str] = None
    rota: Optional[str] = None
    vendedor: Optional[str] = None
    limite_credito: Optional[float] = None
    observacoes: Optional[str] = None
    created_at: str

# --- Contas a Pagar (Completo) ---
class ContaPagarCreate(BaseModel):
    # Dados principais
    fornecedor_id: Optional[str] = None
    fornecedor_nome: Optional[str] = None
    documento: Optional[str] = None  # NF/NFS
    numero_doc: Optional[str] = None
    descricao: str
    valor: float
    valor_desconto: Optional[float] = 0
    valor_juros: Optional[float] = 0
    valor_multa: Optional[float] = 0
    valor_retencao: Optional[float] = 0  # Retenção (IRRF/INSS/ISS) — descontada do valor total
    
    # Parcelamento
    total_parcelas: Optional[int] = 1
    numero_parcela: Optional[int] = 1
    parcela_origem_id: Optional[str] = None  # ID da primeira parcela (para agrupar)
    
    # Datas
    data_emissao: Optional[str] = None
    data_vencimento: str
    data_pagamento: Optional[str] = None
    data_cancelamento: Optional[str] = None
    
    # Classificação
    plano_conta_id: Optional[str] = None
    plano_conta_nome: Optional[str] = None
    subconta_id: Optional[str] = None
    subconta_nome: Optional[str] = None
    centro_custo: Optional[str] = None
    frota_id: Optional[str] = None  # Frota associada
    frota_nome: Optional[str] = None
    maquina_id: Optional[str] = None  # Máquina específica da frota
    maquina_nome: Optional[str] = None
    
    # Pagamento
    forma_pagamento: str = "dinheiro"  # dinheiro, pix, cartao_debito, cartao_credito, boleto, cheque, transferencia
    conta_movimento: Optional[str] = None
    conta_bancaria_id: Optional[str] = None
    conta_bancaria_nome: Optional[str] = None
    
    # Status
    status: str = "em_aberto"  # em_aberto, quitada, cancelada, perdida, parcial
    
    observacoes: Optional[str] = None

class ContaPagarResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    numero: int
    fornecedor_id: Optional[str] = None
    fornecedor_nome: Optional[str] = None
    documento: Optional[str] = None
    numero_doc: Optional[str] = None
    descricao: str
    valor: float
    valor_desconto: Optional[float] = 0
    valor_juros: Optional[float] = 0
    valor_multa: Optional[float] = 0
    valor_retencao: Optional[float] = 0  # Retenção (IRRF/INSS/ISS) — descontada do valor total
    valor_final: Optional[float] = None
    # Parcelamento
    total_parcelas: Optional[int] = 1
    numero_parcela: Optional[int] = 1
    parcela_origem_id: Optional[str] = None
    # Datas
    data_emissao: Optional[str] = None
    data_vencimento: str
    data_pagamento: Optional[str] = None
    data_cancelamento: Optional[str] = None
    plano_conta_id: Optional[str] = None
    plano_conta_nome: Optional[str] = None
    subconta_id: Optional[str] = None
    subconta_nome: Optional[str] = None
    centro_custo: Optional[str] = None
    frota_id: Optional[str] = None
    frota_nome: Optional[str] = None
    maquina_id: Optional[str] = None
    maquina_nome: Optional[str] = None
    forma_pagamento: str
    conta_movimento: Optional[str] = None
    conta_bancaria_id: Optional[str] = None
    conta_bancaria_nome: Optional[str] = None
    status: str
    observacoes: Optional[str] = None
    created_at: str

# --- Contas a Receber (Completo) ---
class ContaReceberCreate(BaseModel):
    # Dados principais
    cliente_id: Optional[str] = None
    cliente_nome: Optional[str] = None
    documento: Optional[str] = None  # NF/NFS
    numero_doc: Optional[str] = None
    descricao: str
    valor: float
    valor_desconto: Optional[float] = 0
    valor_juros: Optional[float] = 0
    valor_multa: Optional[float] = 0
    valor_retencao: Optional[float] = 0  # Retenção (IRRF/INSS/ISS) — descontada do valor total
    
    # Parcelamento
    total_parcelas: Optional[int] = 1
    numero_parcela: Optional[int] = 1
    parcela_origem_id: Optional[str] = None  # ID da primeira parcela (para agrupar)
    
    # Datas
    data_emissao: Optional[str] = None
    data_vencimento: str
    data_recebimento: Optional[str] = None
    data_cancelamento: Optional[str] = None
    
    # Classificação
    plano_conta_id: Optional[str] = None
    plano_conta_nome: Optional[str] = None
    subconta_id: Optional[str] = None
    subconta_nome: Optional[str] = None
    centro_custo: Optional[str] = None
    frota_id: Optional[str] = None  # Frota associada
    frota_nome: Optional[str] = None
    maquina_id: Optional[str] = None  # Máquina específica da frota
    maquina_nome: Optional[str] = None
    
    # Pagamento
    forma_pagamento: str = "dinheiro"
    conta_movimento: Optional[str] = None
    conta_bancaria_id: Optional[str] = None
    conta_bancaria_nome: Optional[str] = None
    
    # Status
    status: str = "em_aberto"  # em_aberto, quitada, cancelada, perdida, parcial
    
    # Faturamento
    faturamento: Optional[str] = None
    
    observacoes: Optional[str] = None

class ContaReceberResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    numero: int
    cliente_id: Optional[str] = None
    cliente_nome: Optional[str] = None
    documento: Optional[str] = None
    numero_doc: Optional[str] = None
    descricao: str
    valor: float
    valor_desconto: Optional[float] = 0
    valor_juros: Optional[float] = 0
    valor_multa: Optional[float] = 0
    valor_retencao: Optional[float] = 0  # Retenção (IRRF/INSS/ISS) — descontada do valor total
    valor_final: Optional[float] = None
    # Parcelamento
    total_parcelas: Optional[int] = 1
    numero_parcela: Optional[int] = 1
    parcela_origem_id: Optional[str] = None
    # Datas
    data_emissao: Optional[str] = None
    data_vencimento: str
    data_recebimento: Optional[str] = None
    data_cancelamento: Optional[str] = None
    plano_conta_id: Optional[str] = None
    plano_conta_nome: Optional[str] = None
    subconta_id: Optional[str] = None
    subconta_nome: Optional[str] = None
    centro_custo: Optional[str] = None
    frota_id: Optional[str] = None
    frota_nome: Optional[str] = None
    maquina_id: Optional[str] = None
    maquina_nome: Optional[str] = None
    forma_pagamento: str
    conta_movimento: Optional[str] = None
    conta_bancaria_id: Optional[str] = None
    conta_bancaria_nome: Optional[str] = None
    status: str
    faturamento: Optional[str] = None
    observacoes: Optional[str] = None
    created_at: str

# --- Produtos (Completo com dados fiscais) ---
class ProdutoCreate(BaseModel):
    # Identificação
    codigo_interno: Optional[str] = None
    codigo_fabricante: Optional[str] = None
    codigo_barras: Optional[str] = None
    descricao: str
    
    # Classificação
    fabricante: Optional[str] = None
    aplicacao: Optional[str] = None
    grupo: Optional[str] = None
    subgrupo: Optional[str] = None
    
    # Unidades
    unidade_comercial: str = "UN"
    unidade_tributada: Optional[str] = None
    multiplo: Optional[float] = 1
    
    # Preços
    preco_custo: Optional[float] = 0
    preco_custo_final: Optional[float] = 0
    preco_venda: Optional[float] = 0
    margem_lucro: Optional[float] = 0
    
    # Estoque
    estoque_atual: Optional[float] = 0
    estoque_minimo: Optional[float] = 0
    estoque_maximo: Optional[float] = 0
    localizacao: Optional[str] = None  # Prateleira
    
    # Dados fiscais
    ncm: Optional[str] = None
    cst: Optional[str] = None
    cest: Optional[str] = None
    origem: Optional[str] = "0"  # 0 = Nacional
    ipi: Optional[float] = 0
    icms: Optional[float] = 0
    pis: Optional[float] = 0
    cofins: Optional[float] = 0
    
    # Tipo do item (NF-e)
    tipo_item: Optional[str] = "00"  # 00=Mercadoria, 01=Matéria-prima, etc
    
    # Status
    status: str = "ativo"
    em_promocao: bool = False
    preco_promocao: Optional[float] = None
    
    observacoes: Optional[str] = None

class ProdutoResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    codigo_interno: Optional[str] = None
    codigo_fabricante: Optional[str] = None
    codigo_barras: Optional[str] = None
    descricao: str
    fabricante: Optional[str] = None
    aplicacao: Optional[str] = None
    grupo: Optional[str] = None
    subgrupo: Optional[str] = None
    unidade_comercial: str
    unidade_tributada: Optional[str] = None
    multiplo: Optional[float] = 1
    preco_custo: Optional[float] = 0
    preco_custo_final: Optional[float] = 0
    preco_venda: Optional[float] = 0
    margem_lucro: Optional[float] = 0
    estoque_atual: Optional[float] = 0
    estoque_minimo: Optional[float] = 0
    estoque_maximo: Optional[float] = 0
    localizacao: Optional[str] = None
    ncm: Optional[str] = None
    cst: Optional[str] = None
    cest: Optional[str] = None
    origem: Optional[str] = "0"
    ipi: Optional[float] = 0
    icms: Optional[float] = 0
    pis: Optional[float] = 0
    cofins: Optional[float] = 0
    tipo_item: Optional[str] = "00"
    status: str
    em_promocao: bool = False
    preco_promocao: Optional[float] = None
    observacoes: Optional[str] = None
    created_at: str
    updated_at: Optional[str] = None

# --- Ordem de Serviço (Completo) ---
class OrdemServicoCreate(BaseModel):
    model_config = ConfigDict(extra="allow")  # aceita campos extras sem reclamar

    # Identificação
    numero_contrato: Optional[str] = None
    numero_documento_fiscal: Optional[str] = None

    # Cliente (todos opcionais)
    cliente_id: Optional[str] = None
    cliente_nome: Optional[str] = None
    cliente_fantasia: Optional[str] = None
    cliente_documento: Optional[str] = None
    cliente_email: Optional[str] = None
    cliente_telefone: Optional[str] = None
    cliente_celular: Optional[str] = None
    cliente_ie: Optional[str] = None
    cliente_endereco: Optional[str] = None
    cliente_bairro: Optional[str] = None
    cliente_cidade: Optional[str] = None
    cliente_uf: Optional[str] = None
    cliente_cep: Optional[str] = None

    # Obra / endereço de entrega
    obra: Optional[str] = None
    obra_id: Optional[str] = None
    endereco_entrega: Optional[str] = None
    prisma: Optional[str] = None
    periodo: Optional[str] = None
    periodicidade: Optional[str] = None  # diaria, semanal, quinzenal, mensal, semestral, anual
    km: Optional[str] = None  # KM rodado (opcional)

    # Datas
    data_abertura: Optional[str] = None
    data_fechamento: Optional[str] = None
    data_previsao_entrega: Optional[str] = None
    data_conclusao: Optional[str] = None

    # Atendimento
    tipo: Optional[str] = "servico"
    tipo_atendimento: Optional[str] = None
    atendente: Optional[str] = None
    atendente_nome: Optional[str] = None
    empresa: Optional[str] = None
    empresa_emissora: Optional[str] = None  # "locadora" | "construtora"
    responsavel_id: Optional[str] = None
    responsavel_nome: Optional[str] = None
    maquina_id: Optional[str] = None
    maquina_nome: Optional[str] = None

    # Itens da OS (opcional — pode ser populado direto na criação)
    itens: Optional[List[dict]] = None

    # Vínculos opcionais (multi-seleção)
    frotas_ids: Optional[List[str]] = None
    maquinas_ids: Optional[List[str]] = None
    fornecedores_ids: Optional[List[str]] = None

    # Valores
    valor_total: Optional[float] = 0
    valor_principal: Optional[float] = 0  # Valor base antes dos extras
    valores_extras: Optional[List[dict]] = None  # [{descricao, valor}]
    valor_desconto: Optional[float] = 0
    valor_subtotal: Optional[float] = 0
    valor_antecipado: Optional[float] = 0

    # Pagamento
    forma_pagamento: Optional[str] = None
    condicao_pagamento: Optional[str] = None

    # Status
    status: str = "em_aberto"  # em_aberto, em_andamento, concluida, cancelada
    confirmada: bool = False
    prioridade: Optional[str] = "media"

    # Tipo Financeiro (refletir no dashboard)
    tipo_financeiro: Optional[str] = None  # a_pagar, a_receber, nenhum

    # Descrição / observações / notas
    descricao: Optional[str] = None
    observacao_servicos: Optional[str] = None
    observacoes: Optional[str] = None
    notas_gerais: Optional[str] = None

class OrdemServicoItemCreate(BaseModel):
    produto_id: Optional[str] = None
    codigo_interno: Optional[str] = None
    descricao: str
    fabricante: Optional[str] = None
    unidade: str = "UN"
    quantidade: float = 1
    valor_unitario: float = 0
    desconto_percent: Optional[float] = 0
    valor_total: Optional[float] = 0

class OrdemServicoResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    numero: int
    numero_contrato: Optional[str] = None
    cliente_id: Optional[str] = None
    cliente_nome: Optional[str] = None
    cliente_fantasia: Optional[str] = None
    obra: Optional[str] = None
    prisma: Optional[str] = None
    data_abertura: str
    data_previsao_entrega: Optional[str] = None
    data_conclusao: Optional[str] = None
    tipo_atendimento: Optional[str] = None
    atendente: Optional[str] = None
    empresa: Optional[str] = None
    valor_total: Optional[float] = 0
    valor_antecipado: Optional[float] = 0
    valor_restante: Optional[float] = 0
    status: str
    confirmada: bool = False
    tipo_financeiro: Optional[str] = None  # a_pagar, a_receber, nenhum
    descricao: Optional[str] = None
    observacoes: Optional[str] = None
    itens: Optional[List[dict]] = []
    created_at: str

# --- Plano de Contas (2 níveis) ---
class PlanoContaCreate(BaseModel):
    codigo: Optional[str] = None
    nome: str
    tipo: str  # receita ou despesa
    nivel: int = 1  # 1 = categoria pai, 2 = subcategoria
    pai_id: Optional[str] = None  # ID da categoria pai se for nível 2
    descricao: Optional[str] = None

class PlanoContaResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    codigo: Optional[str] = None
    nome: str
    tipo: str
    nivel: int
    pai_id: Optional[str] = None
    pai_nome: Optional[str] = None
    descricao: Optional[str] = None
    created_at: str

# --- Centro de Custo ---
class CentroCustoCreate(BaseModel):
    model_config = ConfigDict(extra="allow")
    codigo: Optional[str] = None
    nome: str
    descricao: Optional[str] = None
    status: str = "ativo"
    # Dados de empresa (opcionais — usados quando o centro for emissor de OS/notas)
    eh_empresa_emissora: bool = False
    cnpj: Optional[str] = None
    razao_social: Optional[str] = None
    fantasia: Optional[str] = None
    inscricao_estadual: Optional[str] = None
    inscricao_municipal: Optional[str] = None
    telefone: Optional[str] = None
    celular: Optional[str] = None
    email: Optional[str] = None
    endereco: Optional[str] = None
    bairro: Optional[str] = None
    cidade: Optional[str] = None
    uf: Optional[str] = None
    cep: Optional[str] = None
    logo_base64: Optional[str] = None  # Logo em base64 para uso em PDFs


class CentroCustoResponse(BaseModel):
    model_config = ConfigDict(extra="allow")
    id: str
    codigo: Optional[str] = None
    nome: str
    descricao: Optional[str] = None
    status: str
    created_at: str

# --- Formas de Pagamento ---
class FormaPagamentoCreate(BaseModel):
    codigo: Optional[str] = None
    nome: str
    tipo: str = "outros"  # dinheiro, pix, cartao_debito, cartao_credito, boleto, cheque, transferencia, outros
    taxa: Optional[float] = 0
    prazo_recebimento: Optional[int] = 0  # dias
    conta_bancaria: Optional[str] = None
    ativo: bool = True
    observacoes: Optional[str] = None

class FormaPagamentoResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    codigo: Optional[str] = None
    nome: str
    tipo: str
    taxa: Optional[float] = 0
    prazo_recebimento: Optional[int] = 0
    conta_bancaria: Optional[str] = None
    ativo: bool
    observacoes: Optional[str] = None
    created_at: str

# --- Movimentação de Contas ---
class MovimentacaoContaCreate(BaseModel):
    """Modelo para movimentação entre contas/centros de custo"""
    tipo: str  # entrada, saida, transferencia
    descricao: str
    valor: float
    data_movimentacao: str
    
    # Origem (de onde sai o dinheiro)
    conta_bancaria_origem_id: Optional[str] = None
    conta_bancaria_origem_nome: Optional[str] = None
    centro_custo_origem_id: Optional[str] = None
    centro_custo_origem_nome: Optional[str] = None
    
    # Destino (para onde vai o dinheiro)
    conta_bancaria_destino_id: Optional[str] = None
    conta_bancaria_destino_nome: Optional[str] = None
    centro_custo_destino_id: Optional[str] = None
    centro_custo_destino_nome: Optional[str] = None
    
    # Classificação
    categoria: str = "outros"  # cancelamento_nf, estorno, devolucao, transferencia_interna, ajuste, outros
    documento_referencia: Optional[str] = None  # NF, recibo, etc
    
    observacoes: Optional[str] = None

class MovimentacaoContaResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    numero: int
    tipo: str
    descricao: str
    valor: float
    data_movimentacao: str
    conta_bancaria_origem_id: Optional[str] = None
    conta_bancaria_origem_nome: Optional[str] = None
    centro_custo_origem_id: Optional[str] = None
    centro_custo_origem_nome: Optional[str] = None
    conta_bancaria_destino_id: Optional[str] = None
    conta_bancaria_destino_nome: Optional[str] = None
    centro_custo_destino_id: Optional[str] = None
    centro_custo_destino_nome: Optional[str] = None
    categoria: str
    documento_referencia: Optional[str] = None
    observacoes: Optional[str] = None
    created_by: Optional[str] = None
    created_at: str

# --- Importação Manual de NF ---
class ImportacaoManualNFCreate(BaseModel):
    """Modelo para importação manual de NF quando SEFAZ falha"""
    tipo_nota: str  # nfe ou nfse
    
    # Dados da nota
    numero_nota: str
    serie: Optional[str] = "1"
    chave_acesso: Optional[str] = None
    data_emissao: str
    
    # Emitente
    cnpj_emitente: str
    razao_social_emitente: str
    uf_emitente: Optional[str] = None
    
    # Destinatário
    cnpj_destinatario: Optional[str] = None
    razao_social_destinatario: Optional[str] = None
    
    # Valores
    valor_total: float
    valor_produtos: Optional[float] = None
    valor_servicos: Optional[float] = None
    valor_frete: Optional[float] = 0
    valor_desconto: Optional[float] = 0
    
    # Classificação
    centro_custo_id: Optional[str] = None
    centro_custo_nome: Optional[str] = None
    plano_conta_id: Optional[str] = None
    plano_conta_nome: Optional[str] = None
    
    # Arquivos (em base64)
    xml_base64: Optional[str] = None
    pdf_base64: Optional[str] = None
    
    observacoes: Optional[str] = None

# --- Contas Bancárias ---
class ContaBancariaCreate(BaseModel):
    nome: str  # Nome identificador da conta (ex: "Conta Principal", "Conta Poupança")
    banco: str  # Nome do banco
    codigo_banco: Optional[str] = None  # Código do banco (ex: 001, 341)
    agencia: str
    agencia_digito: Optional[str] = None
    conta: str
    conta_digito: Optional[str] = None
    tipo_conta: str = "corrente"  # corrente, poupanca, investimento, caixa
    titular: Optional[str] = None
    cpf_cnpj_titular: Optional[str] = None
    chave_pix: Optional[str] = None
    tipo_chave_pix: Optional[str] = None  # cpf, cnpj, email, telefone, aleatoria
    saldo_inicial: Optional[float] = 0
    saldo_atual: Optional[float] = 0
    ativo: bool = True
    cor: Optional[str] = "#3B82F6"  # Cor para identificação visual
    observacoes: Optional[str] = None

class ContaBancariaResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    nome: str
    banco: str
    codigo_banco: Optional[str] = None
    agencia: str
    agencia_digito: Optional[str] = None
    conta: str
    conta_digito: Optional[str] = None
    tipo_conta: str
    titular: Optional[str] = None
    cpf_cnpj_titular: Optional[str] = None
    chave_pix: Optional[str] = None
    tipo_chave_pix: Optional[str] = None
    saldo_inicial: Optional[float] = 0
    saldo_atual: Optional[float] = 0
    ativo: bool
    cor: Optional[str] = "#3B82F6"
    observacoes: Optional[str] = None
    created_at: str

# --- Aluguéis de Máquinas ---
class AluguelCreate(BaseModel):
    # Máquina
    maquina_id: str
    maquina_nome: Optional[str] = None
    maquina_placa: Optional[str] = None
    
    # Cliente/Locatário
    cliente_nome: str
    cliente_telefone: Optional[str] = None
    cliente_documento: Optional[str] = None  # CPF/CNPJ
    
    # Contrato
    numero_contrato: Optional[str] = None
    
    # Período
    tipo_periodo: str  # diaria, semanal, quinzenal, mensal, semestral, anual, hora, outro
    periodo_especificado: Optional[str] = None  # quando tipo_periodo = outro
    
    # Datas
    data_entrega: str  # data que a máquina foi entregue
    data_vencimento: str  # data que deve ser devolvida/paga
    data_devolucao: Optional[str] = None  # data real de devolução
    
    # Valores
    valor: float
    valor_caucao: Optional[float] = 0  # valor de caução/garantia
    
    # Local
    local_entrega: Optional[str] = None
    
    # Status e observações
    status: str = "ativo"  # ativo, finalizado, cancelado
    observacoes: Optional[str] = None
    
    # Gerar conta a receber
    gerar_conta_receber: bool = True

class AluguelResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    numero: int
    maquina_id: str
    maquina_nome: Optional[str] = None
    maquina_placa: Optional[str] = None
    cliente_nome: str
    cliente_telefone: Optional[str] = None
    cliente_documento: Optional[str] = None
    tipo_periodo: str
    periodo_especificado: Optional[str] = None
    data_entrega: str
    data_vencimento: str
    data_devolucao: Optional[str] = None
    valor: float
    valor_caucao: Optional[float] = 0
    local_entrega: Optional[str] = None
    status: str
    observacoes: Optional[str] = None
    conta_receber_id: Optional[str] = None
    created_at: str

# --- Configurações do Sistema ---
class ConfiguracaoNotificacao(BaseModel):
    prazo_dias: int = 7  # dias antes do vencimento para notificar

# ============ ADMIN ENDPOINTS ============

# --- Funções auxiliares para auto-incremento ---
async def get_next_sequence(collection_name: str) -> int:
    result = await db.counters.find_one_and_update(
        {"_id": collection_name},
        {"$inc": {"seq": 1}},
        upsert=True,
        return_document=True
    )
    return result["seq"]

# --- Dashboard Admin ---
@api_router.get("/admin/dashboard")
async def get_admin_dashboard(current_user: dict = Depends(get_current_user)):
    from datetime import timedelta
    hoje = datetime.now(timezone.utc)
    hoje_str = hoje.strftime("%Y-%m-%d")
    inicio_mes = hoje.replace(day=1).strftime("%Y-%m-%d")
    inicio_ano = hoje.replace(month=1, day=1).strftime("%Y-%m-%d")
    proxima_semana = (hoje + timedelta(days=7)).strftime("%Y-%m-%d")
    
    # ===== CONTAS A PAGAR =====
    # Em aberto
    contas_pagar_abertas = await db.contas_pagar.find({"status": "em_aberto"}, {"_id": 0}).to_list(5000)
    total_pagar_aberto = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_pagar_abertas)
    total_pagar_mes = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_pagar_abertas if (c.get("data_vencimento") or "") >= inicio_mes)
    total_pagar_ano = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_pagar_abertas if (c.get("data_vencimento") or "") >= inicio_ano)
    lista_pagar_vencidas = [c for c in contas_pagar_abertas if (c.get("data_vencimento") or "") < hoje_str and c.get("data_vencimento")]
    contas_pagar_vencidas = len(lista_pagar_vencidas)
    total_pagar_vencidas_valor = sum(c.get("valor_final") or c.get("valor", 0) for c in lista_pagar_vencidas)
    
    # Quitadas
    contas_pagar_quitadas = await db.contas_pagar.find({"status": "quitada"}, {"_id": 0}).to_list(5000)
    total_pagar_quitado = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_pagar_quitadas)
    total_pagar_quitado_mes = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_pagar_quitadas if (c.get("data_pagamento") or c.get("data_vencimento") or "") >= inicio_mes)
    total_pagar_quitado_ano = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_pagar_quitadas if (c.get("data_pagamento") or c.get("data_vencimento") or "") >= inicio_ano)
    
    # ===== CONTAS A RECEBER =====
    # Em aberto
    contas_receber_abertas = await db.contas_receber.find({"status": "em_aberto"}, {"_id": 0}).to_list(5000)
    total_receber_aberto = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_receber_abertas)
    total_receber_mes = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_receber_abertas if (c.get("data_vencimento") or "") >= inicio_mes)
    total_receber_ano = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_receber_abertas if (c.get("data_vencimento") or "") >= inicio_ano)
    lista_receber_vencidas = [c for c in contas_receber_abertas if (c.get("data_vencimento") or "") < hoje_str and c.get("data_vencimento")]
    contas_receber_vencidas = len(lista_receber_vencidas)
    total_receber_vencidas_valor = sum(c.get("valor_final") or c.get("valor", 0) for c in lista_receber_vencidas)
    
    # Quitadas/Recebidas
    contas_receber_quitadas = await db.contas_receber.find({"status": "quitada"}, {"_id": 0}).to_list(5000)
    total_receber_quitado = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_receber_quitadas)
    total_receber_quitado_mes = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_receber_quitadas if (c.get("data_recebimento") or c.get("data_vencimento") or "") >= inicio_mes)
    total_receber_quitado_ano = sum(c.get("valor_final") or c.get("valor", 0) for c in contas_receber_quitadas if (c.get("data_recebimento") or c.get("data_vencimento") or "") >= inicio_ano)
    
    # ===== ORDENS DE SERVIÇO COM TIPO FINANCEIRO =====
    os_a_pagar = await db.ordens_servico.find({"tipo_financeiro": "a_pagar", "status": {"$ne": "cancelada"}}, {"_id": 0}).to_list(1000)
    os_a_receber = await db.ordens_servico.find({"tipo_financeiro": "a_receber", "status": {"$ne": "cancelada"}}, {"_id": 0}).to_list(1000)
    total_os_pagar = sum(o.get("valor_total", 0) for o in os_a_pagar)
    total_os_receber = sum(o.get("valor_total", 0) for o in os_a_receber)
    
    # ===== COUNTS =====
    cadastros = await db.cadastros.count_documents({})
    fornecedores = await db.cadastros.count_documents({"tipo_cadastro": {"$in": ["fornecedor", "ambos"]}})
    produtos = await db.produtos_admin.count_documents({})
    ordens_abertas = await db.ordens_servico.count_documents({"status": {"$in": ["em_aberto", "em_andamento"]}})
    notas_emitidas = await db.notas_fiscais.count_documents({}) if await db.list_collection_names() else 0
    
    # ===== PRÓXIMOS VENCIMENTOS =====
    proximas_pagar = await db.contas_pagar.find({
        "status": "em_aberto",
        "data_vencimento": {"$lte": proxima_semana}
    }, {"_id": 0}).sort("data_vencimento", 1).limit(5).to_list(5)
    
    proximas_receber = await db.contas_receber.find({
        "status": "em_aberto",
        "data_vencimento": {"$lte": proxima_semana}
    }, {"_id": 0}).sort("data_vencimento", 1).limit(5).to_list(5)
    
    contas_proximas = []
    for c in proximas_pagar:
        contas_proximas.append({
            "tipo": "pagar",
            "descricao": c.get("descricao"),
            "valor": c.get("valor_final") or c.get("valor"),
            "vencimento": c.get("data_vencimento")
        })
    for c in proximas_receber:
        contas_proximas.append({
            "tipo": "receber",
            "descricao": c.get("descricao"),
            "valor": c.get("valor_final") or c.get("valor"),
            "vencimento": c.get("data_vencimento")
        })
    
    contas_proximas.sort(key=lambda x: x.get("vencimento", ""))
    
    # Saldo previsto = (receber aberto + OS receber) - (pagar aberto + OS pagar)
    saldo_previsto = (total_receber_aberto + total_os_receber) - (total_pagar_aberto + total_os_pagar)
    
    return {
        "stats": {
            "totalPagar": total_pagar_aberto,
            "totalReceber": total_receber_aberto,
            "saldoPrevisto": saldo_previsto,
            "contasVencidas": contas_pagar_vencidas + contas_receber_vencidas,
            "contasPagarVencidas": contas_pagar_vencidas,
            "contasReceberVencidas": contas_receber_vencidas,
            "notasEmitidas": 0,
            "fornecedores": fornecedores,
            "cadastros": cadastros,
            "produtos": produtos,
            "ordensAbertas": ordens_abertas
        },
        "aPagar": {
            "total": total_pagar_aberto,
            "mes": total_pagar_mes,
            "ano": total_pagar_ano,
            "vencidas": contas_pagar_vencidas,
            "osValor": total_os_pagar
        },
        "aReceber": {
            "total": total_receber_aberto,
            "mes": total_receber_mes,
            "ano": total_receber_ano,
            "vencidas": contas_receber_vencidas,
            "osValor": total_os_receber
        },
        "quitados": {
            "pagar": {
                "total": total_pagar_quitado,
                "mes": total_pagar_quitado_mes,
                "ano": total_pagar_quitado_ano
            },
            "receber": {
                "total": total_receber_quitado,
                "mes": total_receber_quitado_mes,
                "ano": total_receber_quitado_ano
            }
        },
        "contasProximas": contas_proximas[:10],
        "vencidas": {
            "pagar": {
                "quantidade": contas_pagar_vencidas,
                "valor": total_pagar_vencidas_valor,
                "lista": sorted(lista_pagar_vencidas, key=lambda x: x.get("data_vencimento", ""))[:50]
            },
            "receber": {
                "quantidade": contas_receber_vencidas,
                "valor": total_receber_vencidas_valor,
                "lista": sorted(lista_receber_vencidas, key=lambda x: x.get("data_vencimento", ""))[:50]
            },
            "totalQuantidade": contas_pagar_vencidas + contas_receber_vencidas,
            "totalValor": total_pagar_vencidas_valor + total_receber_vencidas_valor
        }
    }

# --- Consulta de CNPJ via BrasilAPI (Receita Federal) ---
import httpx

@api_router.get("/consulta/cnpj/{cnpj}")
async def consulta_cnpj(
    cnpj: str,
    current_user: dict = Depends(get_current_user)
):
    """Consulta dados de empresa via CNPJ usando BrasilAPI (dados da Receita Federal)"""
    
    # Remove caracteres não numéricos
    cnpj_limpo = ''.join(filter(str.isdigit, cnpj))
    
    if len(cnpj_limpo) != 14:
        raise HTTPException(status_code=400, detail="CNPJ inválido. Deve conter 14 dígitos.")
    
    try:
        async with httpx.AsyncClient() as client:
            # Tentar BrasilAPI primeiro (mais confiável)
            response = await client.get(
                f"https://brasilapi.com.br/api/cnpj/v1/{cnpj_limpo}",
                timeout=15
            )
            
            if response.status_code == 200:
                data = response.json()
                
                # Mapear dados da BrasilAPI
                resultado = {
                    "razao_social": data.get("razao_social", ""),
                    "nome_fantasia": data.get("nome_fantasia", ""),
                    "cnpj": cnpj_limpo,
                    "inscricao_estadual": "",  # BrasilAPI não retorna IE
                    "telefone": data.get("ddd_telefone_1", ""),
                    "email": data.get("email", ""),
                    "cep": data.get("cep", ""),
                    "endereco": data.get("logradouro", ""),
                    "numero": data.get("numero", ""),
                    "complemento": data.get("complemento", ""),
                    "bairro": data.get("bairro", ""),
                    "cidade": data.get("municipio", ""),
                    "uf": data.get("uf", ""),
                    "situacao": data.get("descricao_situacao_cadastral", ""),
                    "atividade_principal": data.get("cnae_fiscal_descricao", ""),
                    "capital_social": str(data.get("capital_social", "")) if data.get("capital_social") else "",
                    "data_abertura": data.get("data_inicio_atividade", ""),
                    "natureza_juridica": data.get("natureza_juridica", ""),
                    "porte": data.get("porte", ""),
                }
                
                await create_audit_log(
                    user=current_user,
                    action="consultar cnpj",
                    entity_type="cnpj",
                    entity_id=cnpj_limpo,
                    entity_name=resultado.get("razao_social", cnpj_limpo),
                    module="Administrativo"
                )
                
                return {"success": True, "data": resultado}
            
            elif response.status_code == 404:
                raise HTTPException(status_code=404, detail="CNPJ não encontrado na base da Receita Federal")
            else:
                raise HTTPException(status_code=500, detail="Erro ao consultar CNPJ")
                
    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="Tempo de consulta excedido. Tente novamente.")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erro ao consultar CNPJ: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao consultar CNPJ: {str(e)}")

@api_router.get("/consulta/cep/{cep}")
async def consulta_cep(
    cep: str,
    current_user: dict = Depends(get_current_user)
):
    """Consulta endereço via CEP usando ViaCEP"""
    
    # Remove caracteres não numéricos
    cep_limpo = ''.join(filter(str.isdigit, cep))
    
    if len(cep_limpo) != 8:
        raise HTTPException(status_code=400, detail="CEP inválido. Deve conter 8 dígitos.")
    
    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(f"https://viacep.com.br/ws/{cep_limpo}/json/", timeout=10)
            
            if response.status_code != 200:
                raise HTTPException(status_code=404, detail="CEP não encontrado")
            
            data = response.json()
            
            if data.get("erro"):
                raise HTTPException(status_code=404, detail="CEP não encontrado")
            
            resultado = {
                "cep": data.get("cep", "").replace("-", ""),
                "endereco": data.get("logradouro", ""),
                "complemento": data.get("complemento", ""),
                "bairro": data.get("bairro", ""),
                "cidade": data.get("localidade", ""),
                "uf": data.get("uf", ""),
                "ibge": data.get("ibge", ""),
                "ddd": data.get("ddd", "")
            }
            
            return {"success": True, "data": resultado}
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erro ao consultar CEP: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao consultar CEP: {str(e)}")

# --- Cadastro de Clientes/Fornecedores ---
@api_router.get("/admin/cadastros")
async def get_cadastros(
    tipo: Optional[str] = None,
    status: Optional[str] = None,
    search: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if tipo:
        query["tipo_cadastro"] = tipo
    if status:
        query["status"] = status
    if search:
        query["$or"] = [
            {"nome_razao": {"$regex": search, "$options": "i"}},
            {"apelido_fantasia": {"$regex": search, "$options": "i"}},
            {"cpf_cnpj": {"$regex": search, "$options": "i"}},
            {"cidade": {"$regex": search, "$options": "i"}}
        ]
    
    cadastros = await db.cadastros.find(query, {"_id": 0}).sort("nome_razao", 1).to_list(1000)
    return cadastros

@api_router.post("/admin/cadastros")
async def create_cadastro(data: CadastroCreate, current_user: dict = Depends(get_current_user)):
    codigo = await get_next_sequence("cadastros")
    cadastro = {
        "id": str(uuid.uuid4()),
        "codigo": codigo,
        **data.model_dump(),
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.cadastros.insert_one(cadastro)
    await create_audit_log(current_user, "create", "cadastro", cadastro["id"], data.nome_razao)
    del cadastro["_id"]
    return cadastro

@api_router.get("/admin/cadastros/{id}")
async def get_cadastro(id: str, current_user: dict = Depends(get_current_user)):
    cadastro = await db.cadastros.find_one({"id": id}, {"_id": 0})
    if not cadastro:
        raise HTTPException(status_code=404, detail="Cadastro não encontrado")
    return cadastro

@api_router.put("/admin/cadastros/{id}")
async def update_cadastro(id: str, data: CadastroCreate, current_user: dict = Depends(get_current_user)):
    cadastro = await db.cadastros.find_one({"id": id}, {"_id": 0})
    if not cadastro:
        raise HTTPException(status_code=404, detail="Cadastro não encontrado")
    
    update_data = data.model_dump()
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.cadastros.update_one({"id": id}, {"$set": update_data})
    await create_audit_log(current_user, "update", "cadastro", id, data.nome_razao)
    
    updated = await db.cadastros.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.delete("/admin/cadastros/{id}")
async def delete_cadastro(id: str, current_user: dict = Depends(get_current_user)):
    cadastro = await db.cadastros.find_one({"id": id}, {"_id": 0})
    if not cadastro:
        raise HTTPException(status_code=404, detail="Cadastro não encontrado")
    
    await db.cadastros.delete_one({"id": id})
    await create_audit_log(current_user, "delete", "cadastro", id, cadastro["nome_razao"])
    return {"message": "Cadastro excluído"}

# --- Anexos de Cadastros ---
CADASTROS_ANEXOS_DIR = ROOT_DIR / "uploads" / "cadastros"
CADASTROS_ANEXOS_DIR.mkdir(parents=True, exist_ok=True)

@api_router.post("/admin/cadastros/{id}/anexos")
async def upload_cadastro_anexo(
    id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload anexo para cadastro"""
    cadastro = await db.cadastros.find_one({"id": id}, {"_id": 0})
    if not cadastro:
        raise HTTPException(status_code=404, detail="Cadastro não encontrado")
    
    content = await file.read()
    max_size = 50 * 1024 * 1024
    if len(content) > max_size:
        raise HTTPException(status_code=400, detail="Arquivo muito grande. Máximo: 50MB")
    
    ext = Path(file.filename).suffix.lower() if file.filename else ''
    unique_filename = f"{id}_{uuid.uuid4()}{ext}"
    file_path = CADASTROS_ANEXOS_DIR / unique_filename
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    anexo = {
        "id": str(uuid.uuid4()),
        "filename": unique_filename,
        "original_name": file.filename,
        "size": len(content),
        "uploaded_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.cadastros.update_one({"id": id}, {"$push": {"anexos": anexo}})
    await create_audit_log(current_user, "upload anexo", "cadastro", id, f"{cadastro['nome_razao']} - {file.filename}")
    
    return {"message": "Anexo adicionado", "anexo": anexo}

@api_router.delete("/admin/cadastros/{id}/anexos/{anexo_id}")
async def delete_cadastro_anexo(
    id: str,
    anexo_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Excluir anexo de cadastro"""
    cadastro = await db.cadastros.find_one({"id": id}, {"_id": 0})
    if not cadastro:
        raise HTTPException(status_code=404, detail="Cadastro não encontrado")
    
    anexo = next((a for a in cadastro.get("anexos", []) if a["id"] == anexo_id), None)
    if not anexo:
        raise HTTPException(status_code=404, detail="Anexo não encontrado")
    
    file_path = CADASTROS_ANEXOS_DIR / anexo["filename"]
    if file_path.exists():
        file_path.unlink()
    
    await db.cadastros.update_one({"id": id}, {"$pull": {"anexos": {"id": anexo_id}}})
    await create_audit_log(current_user, "excluir anexo", "cadastro", id, cadastro["nome_razao"])
    
    return {"message": "Anexo excluído"}

@api_router.get("/admin/cadastros/{id}/anexos/{anexo_id}/download")
async def download_cadastro_anexo(
    id: str,
    anexo_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Download de anexo"""
    cadastro = await db.cadastros.find_one({"id": id}, {"_id": 0})
    if not cadastro:
        raise HTTPException(status_code=404, detail="Cadastro não encontrado")
    
    anexo = next((a for a in cadastro.get("anexos", []) if a["id"] == anexo_id), None)
    if not anexo:
        raise HTTPException(status_code=404, detail="Anexo não encontrado")
    
    file_path = CADASTROS_ANEXOS_DIR / anexo["filename"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
    
    return FileResponse(
        path=str(file_path),
        filename=anexo.get("original_name", anexo["filename"]),
        media_type="application/octet-stream"
    )

# --- Contas a Pagar (Completo) ---
# ============ CONTAS A PAGAR / RECEBER ============
# (Endpoints /admin/contas-pagar/* e /admin/contas-receber/* extraídos para
#  /app/backend/routes/financeiro.py — Refactor Sessão 32 Parte 2)

# --- Produtos (Completo com dados fiscais) ---
@api_router.get("/admin/produtos")
async def get_produtos(
    grupo: Optional[str] = None,
    subgrupo: Optional[str] = None,
    fabricante: Optional[str] = None,
    status: Optional[str] = None,
    estoque_baixo: Optional[bool] = None,
    search: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if grupo:
        query["grupo"] = grupo
    if subgrupo:
        query["subgrupo"] = subgrupo
    if fabricante:
        query["fabricante"] = fabricante
    if status:
        query["status"] = status
    if estoque_baixo:
        query["$expr"] = {"$lte": ["$estoque_atual", "$estoque_minimo"]}
    if search:
        query["$or"] = [
            {"descricao": {"$regex": search, "$options": "i"}},
            {"codigo_interno": {"$regex": search, "$options": "i"}},
            {"codigo_fabricante": {"$regex": search, "$options": "i"}},
            {"fabricante": {"$regex": search, "$options": "i"}},
            {"aplicacao": {"$regex": search, "$options": "i"}}
        ]
    
    produtos = await db.produtos_admin.find(query, {"_id": 0}).sort("descricao", 1).to_list(1000)
    return produtos

@api_router.post("/admin/produtos")
async def create_produto(data: ProdutoCreate, current_user: dict = Depends(get_current_user)):
    # Gerar código interno se não fornecido
    codigo_interno = data.codigo_interno
    if not codigo_interno:
        seq = await get_next_sequence("produtos")
        codigo_interno = f"P{seq:06d}"
    
    # Calcular margem de lucro
    margem = 0
    if data.preco_custo and data.preco_custo > 0 and data.preco_venda:
        margem = ((data.preco_venda - data.preco_custo) / data.preco_custo) * 100
    
    produto = {
        "id": str(uuid.uuid4()),
        "codigo_interno": codigo_interno,
        **data.model_dump(exclude={"codigo_interno", "margem_lucro"}),
        "margem_lucro": round(margem, 2),
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.produtos_admin.insert_one(produto)
    await create_audit_log(current_user, "create", "produto", produto["id"], data.descricao)
    del produto["_id"]
    return produto

@api_router.get("/admin/produtos/{id}")
async def get_produto(id: str, current_user: dict = Depends(get_current_user)):
    produto = await db.produtos_admin.find_one({"id": id}, {"_id": 0})
    if not produto:
        raise HTTPException(status_code=404, detail="Produto não encontrado")
    return produto

@api_router.put("/admin/produtos/{id}")
async def update_produto(id: str, data: ProdutoCreate, current_user: dict = Depends(get_current_user)):
    produto = await db.produtos_admin.find_one({"id": id}, {"_id": 0})
    if not produto:
        raise HTTPException(status_code=404, detail="Produto não encontrado")
    
    # Calcular margem de lucro
    margem = 0
    if data.preco_custo and data.preco_custo > 0 and data.preco_venda:
        margem = ((data.preco_venda - data.preco_custo) / data.preco_custo) * 100
    
    update_data = data.model_dump()
    update_data["margem_lucro"] = round(margem, 2)
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.produtos_admin.update_one({"id": id}, {"$set": update_data})
    await create_audit_log(current_user, "update", "produto", id, data.descricao)
    
    updated = await db.produtos_admin.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.delete("/admin/produtos/{id}")
async def delete_produto(id: str, current_user: dict = Depends(get_current_user)):
    produto = await db.produtos_admin.find_one({"id": id}, {"_id": 0})
    if not produto:
        raise HTTPException(status_code=404, detail="Produto não encontrado")
    
    await db.produtos_admin.delete_one({"id": id})
    await create_audit_log(current_user, "delete", "produto", id, produto["descricao"])
    return {"message": "Produto excluído"}

# --- Anexos de Produtos ---
PRODUTOS_ANEXOS_DIR = ROOT_DIR / "uploads" / "produtos"
PRODUTOS_ANEXOS_DIR.mkdir(parents=True, exist_ok=True)

@api_router.post("/admin/produtos/{id}/anexos")
async def upload_produto_anexo(
    id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload anexo para produto"""
    produto = await db.produtos_admin.find_one({"id": id}, {"_id": 0})
    if not produto:
        raise HTTPException(status_code=404, detail="Produto não encontrado")
    
    content = await file.read()
    max_size = 50 * 1024 * 1024
    if len(content) > max_size:
        raise HTTPException(status_code=400, detail="Arquivo muito grande. Máximo: 50MB")
    
    ext = Path(file.filename).suffix.lower() if file.filename else ''
    unique_filename = f"{id}_{uuid.uuid4()}{ext}"
    file_path = PRODUTOS_ANEXOS_DIR / unique_filename
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    anexo = {
        "id": str(uuid.uuid4()),
        "filename": unique_filename,
        "original_name": file.filename,
        "size": len(content),
        "uploaded_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.produtos_admin.update_one({"id": id}, {"$push": {"anexos": anexo}})
    await create_audit_log(current_user, "upload anexo", "produto", id, f"{produto['descricao']} - {file.filename}")
    
    return {"message": "Anexo adicionado", "anexo": anexo}

@api_router.delete("/admin/produtos/{id}/anexos/{anexo_id}")
async def delete_produto_anexo(
    id: str,
    anexo_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Excluir anexo de produto"""
    produto = await db.produtos_admin.find_one({"id": id}, {"_id": 0})
    if not produto:
        raise HTTPException(status_code=404, detail="Produto não encontrado")
    
    anexo = next((a for a in produto.get("anexos", []) if a["id"] == anexo_id), None)
    if not anexo:
        raise HTTPException(status_code=404, detail="Anexo não encontrado")
    
    file_path = PRODUTOS_ANEXOS_DIR / anexo["filename"]
    if file_path.exists():
        file_path.unlink()
    
    await db.produtos_admin.update_one({"id": id}, {"$pull": {"anexos": {"id": anexo_id}}})
    await create_audit_log(current_user, "excluir anexo", "produto", id, produto["descricao"])
    
    return {"message": "Anexo excluído"}

@api_router.get("/admin/produtos/{id}/anexos/{anexo_id}/download")
async def download_produto_anexo(
    id: str,
    anexo_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Download de anexo de produto"""
    produto = await db.produtos_admin.find_one({"id": id}, {"_id": 0})
    if not produto:
        raise HTTPException(status_code=404, detail="Produto não encontrado")
    
    anexo = next((a for a in produto.get("anexos", []) if a["id"] == anexo_id), None)
    if not anexo:
        raise HTTPException(status_code=404, detail="Anexo não encontrado")
    
    file_path = PRODUTOS_ANEXOS_DIR / anexo["filename"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
    
    return FileResponse(
        path=str(file_path),
        filename=anexo.get("original_name", anexo["filename"]),
        media_type="application/octet-stream"
    )

# --- Ordens de Serviço (Completo) ---
@api_router.get("/admin/ordens-servico")
async def get_ordens_servico(
    status: Optional[str] = None,
    cliente_id: Optional[str] = None,
    obra: Optional[str] = None,
    data_inicio: Optional[str] = None,
    data_fim: Optional[str] = None,
    confirmada: Optional[bool] = None,
    search: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if status:
        query["status"] = status
    if cliente_id:
        query["cliente_id"] = cliente_id
    if obra:
        query["obra"] = {"$regex": obra, "$options": "i"}
    if confirmada is not None:
        query["confirmada"] = confirmada
    if data_inicio and data_fim:
        query["data_abertura"] = {"$gte": data_inicio, "$lte": data_fim}
    if search:
        query["$or"] = [
            {"cliente_nome": {"$regex": search, "$options": "i"}},
            {"cliente_fantasia": {"$regex": search, "$options": "i"}},
            {"obra": {"$regex": search, "$options": "i"}},
            {"descricao": {"$regex": search, "$options": "i"}},
            {"numero_contrato": {"$regex": search, "$options": "i"}}
        ]
    
    ordens = await db.ordens_servico.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    
    # Calcular valor restante
    for ordem in ordens:
        valor_total = ordem.get("valor_total", 0) or 0
        valor_antecipado = ordem.get("valor_antecipado", 0) or 0
        ordem["valor_restante"] = valor_total - valor_antecipado
    
    return ordens

@api_router.post("/admin/ordens-servico")
async def create_ordem_servico(data: OrdemServicoCreate, current_user: dict = Depends(get_current_user)):
    numero = await get_next_sequence("ordens_servico")

    payload = data.model_dump()
    # Garante itens como lista (preserva caso já tenha vindo no payload)
    if not isinstance(payload.get("itens"), list):
        payload["itens"] = []

    ordem = {
        "id": str(uuid.uuid4()),
        "numero": numero,
        **payload,
        "data_abertura": data.data_abertura or datetime.now(timezone.utc).strftime("%Y-%m-%d"),
        "valor_restante": (data.valor_total or 0) - (data.valor_antecipado or 0),
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.ordens_servico.insert_one(ordem)
    await create_audit_log(current_user, "create", "ordem_servico", ordem["id"], f"OS-{numero}")
    del ordem["_id"]
    return ordem

@api_router.get("/admin/ordens-servico/{id}")
async def get_ordem_servico(id: str, current_user: dict = Depends(get_current_user)):
    ordem = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    if not ordem:
        raise HTTPException(status_code=404, detail="Ordem de serviço não encontrada")
    
    valor_total = ordem.get("valor_total", 0) or 0
    valor_antecipado = ordem.get("valor_antecipado", 0) or 0
    ordem["valor_restante"] = valor_total - valor_antecipado
    
    return ordem

@api_router.put("/admin/ordens-servico/{id}")
async def update_ordem_servico(id: str, data: OrdemServicoCreate, current_user: dict = Depends(get_current_user)):
    ordem = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    if not ordem:
        raise HTTPException(status_code=404, detail="Ordem de serviço não encontrada")
    
    update_data = data.model_dump()
    update_data["valor_restante"] = (data.valor_total or 0) - (data.valor_antecipado or 0)
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.ordens_servico.update_one({"id": id}, {"$set": update_data})
    await create_audit_log(current_user, "update", "ordem_servico", id, f"OS-{ordem['numero']}")
    
    updated = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.post("/admin/ordens-servico/{id}/itens")
async def add_item_ordem_servico(id: str, data: OrdemServicoItemCreate, current_user: dict = Depends(get_current_user)):
    ordem = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    if not ordem:
        raise HTTPException(status_code=404, detail="Ordem de serviço não encontrada")
    
    valor_total = data.quantidade * data.valor_unitario
    if data.desconto_percent:
        valor_total = valor_total * (1 - data.desconto_percent / 100)
    
    item = {
        "id": str(uuid.uuid4()),
        **data.model_dump(),
        "valor_total": round(valor_total, 2)
    }
    
    await db.ordens_servico.update_one({"id": id}, {"$push": {"itens": item}})
    
    # Recalcular valor total da OS
    ordem_atualizada = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    novo_total = sum(i.get("valor_total", 0) for i in ordem_atualizada.get("itens", []))
    await db.ordens_servico.update_one({"id": id}, {"$set": {"valor_total": novo_total}})
    
    return {"message": "Item adicionado", "item": item}

@api_router.delete("/admin/ordens-servico/{id}/itens/{item_id}")
async def remove_item_ordem_servico(id: str, item_id: str, current_user: dict = Depends(get_current_user)):
    ordem = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    if not ordem:
        raise HTTPException(status_code=404, detail="Ordem de serviço não encontrada")
    
    await db.ordens_servico.update_one({"id": id}, {"$pull": {"itens": {"id": item_id}}})
    
    # Recalcular valor total da OS
    ordem_atualizada = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    novo_total = sum(i.get("valor_total", 0) for i in ordem_atualizada.get("itens", []))
    await db.ordens_servico.update_one({"id": id}, {"$set": {"valor_total": novo_total}})
    
    return {"message": "Item removido"}

class StatusOSUpdate(BaseModel):
    status: str

@api_router.patch("/admin/ordens-servico/{id}/status")
async def update_ordem_status(id: str, data: StatusOSUpdate, current_user: dict = Depends(get_current_user)):
    ordem = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    if not ordem:
        raise HTTPException(status_code=404, detail="Ordem de serviço não encontrada")
    
    update_fields = {"status": data.status}
    if data.status == "concluida":
        update_fields["data_conclusao"] = datetime.now(timezone.utc).strftime("%Y-%m-%d")
    
    await db.ordens_servico.update_one({"id": id}, {"$set": update_fields})
    await create_audit_log(current_user, "update", "ordem_servico", id, f"OS-{ordem['numero']} - Status: {data.status}")
    return {"message": "Status atualizado"}

@api_router.patch("/admin/ordens-servico/{id}/confirmar")
async def confirmar_ordem_servico(id: str, current_user: dict = Depends(get_current_user)):
    ordem = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    if not ordem:
        raise HTTPException(status_code=404, detail="Ordem de serviço não encontrada")
    
    await db.ordens_servico.update_one({"id": id}, {"$set": {"confirmada": True}})
    await create_audit_log(current_user, "update", "ordem_servico", id, f"OS-{ordem['numero']} - CONFIRMADA")
    return {"message": "Ordem confirmada"}

@api_router.delete("/admin/ordens-servico/{id}")
async def delete_ordem_servico(id: str, current_user: dict = Depends(get_current_user)):
    ordem = await db.ordens_servico.find_one({"id": id}, {"_id": 0})
    if not ordem:
        raise HTTPException(status_code=404, detail="Ordem de serviço não encontrada")
    
    await db.ordens_servico.delete_one({"id": id})
    await create_audit_log(current_user, "delete", "ordem_servico", id, f"OS-{ordem['numero']}")
    return {"message": "Ordem de serviço excluída"}

# --- Plano de Contas (2 níveis) ---
@api_router.get("/admin/plano-contas")
async def get_plano_contas(
    tipo: Optional[str] = None,
    nivel: Optional[int] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if tipo:
        query["tipo"] = tipo
    if nivel:
        query["nivel"] = nivel
    
    contas = await db.plano_contas.find(query, {"_id": 0}).sort([("tipo", 1), ("codigo", 1)]).to_list(1000)
    
    # Adicionar nome do pai se for nível 2
    for conta in contas:
        if conta.get("pai_id"):
            pai = await db.plano_contas.find_one({"id": conta["pai_id"]}, {"_id": 0})
            if pai:
                conta["pai_nome"] = pai.get("nome")
    
    return contas

@api_router.post("/admin/plano-contas")
async def create_plano_conta(data: PlanoContaCreate, current_user: dict = Depends(get_current_user)):
    conta = {
        "id": str(uuid.uuid4()),
        **data.model_dump(),
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.plano_contas.insert_one(conta)
    await create_audit_log(current_user, "create", "plano_conta", conta["id"], data.nome)
    del conta["_id"]
    return conta

@api_router.put("/admin/plano-contas/{id}")
async def update_plano_conta(id: str, data: PlanoContaCreate, current_user: dict = Depends(get_current_user)):
    conta = await db.plano_contas.find_one({"id": id}, {"_id": 0})
    if not conta:
        raise HTTPException(status_code=404, detail="Conta não encontrada")
    
    update_data = data.model_dump()
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.plano_contas.update_one({"id": id}, {"$set": update_data})
    await create_audit_log(current_user, "update", "plano_conta", id, data.nome)
    
    updated = await db.plano_contas.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.delete("/admin/plano-contas/{id}")
async def delete_plano_conta(id: str, current_user: dict = Depends(get_current_user)):
    conta = await db.plano_contas.find_one({"id": id}, {"_id": 0})
    if not conta:
        raise HTTPException(status_code=404, detail="Conta não encontrada")
    
    # Excluir subcontas primeiro (cascata)
    subcontas = await db.plano_contas.find({"pai_id": id}, {"_id": 0}).to_list(100)
    for sub in subcontas:
        await db.plano_contas.delete_one({"id": sub["id"]})
        await create_audit_log(current_user, "delete", "plano_conta", sub["id"], f"Subconta: {sub['nome']}")
    
    # Excluir a conta principal
    await db.plano_contas.delete_one({"id": id})
    await create_audit_log(current_user, "delete", "plano_conta", id, conta["nome"])
    
    msg = f"Conta excluída" + (f" junto com {len(subcontas)} subconta(s)" if subcontas else "")
    return {"message": msg}

# --- Centro de Custo ---
@api_router.get("/admin/centros-custo")
async def get_centros_custo(current_user: dict = Depends(get_current_user)):
    centros = await db.centros_custo.find({}, {"_id": 0}).sort("nome", 1).to_list(1000)
    return centros

@api_router.post("/admin/centros-custo")
async def create_centro_custo(data: CentroCustoCreate, current_user: dict = Depends(get_current_user)):
    centro = {
        "id": str(uuid.uuid4()),
        **data.model_dump(),
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.centros_custo.insert_one(centro)
    await create_audit_log(current_user, "create", "centro_custo", centro["id"], data.nome)
    del centro["_id"]
    return centro

@api_router.put("/admin/centros-custo/{id}")
async def update_centro_custo(id: str, data: CentroCustoCreate, current_user: dict = Depends(get_current_user)):
    centro = await db.centros_custo.find_one({"id": id}, {"_id": 0})
    if not centro:
        raise HTTPException(status_code=404, detail="Centro de custo não encontrado")
    
    update_data = data.model_dump()
    await db.centros_custo.update_one({"id": id}, {"$set": update_data})
    await create_audit_log(current_user, "update", "centro_custo", id, data.nome)
    
    updated = await db.centros_custo.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.delete("/admin/centros-custo/{id}")
async def delete_centro_custo(id: str, current_user: dict = Depends(get_current_user)):
    centro = await db.centros_custo.find_one({"id": id}, {"_id": 0})
    if not centro:
        raise HTTPException(status_code=404, detail="Centro de custo não encontrado")
    
    await db.centros_custo.delete_one({"id": id})
    await create_audit_log(current_user, "delete", "centro_custo", id, centro["nome"])
    return {"message": "Centro de custo excluído"}

# --- Formas de Pagamento ---
@api_router.get("/admin/formas-pagamento")
async def get_formas_pagamento(
    ativo: Optional[bool] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if ativo is not None:
        query["ativo"] = ativo
    
    formas = await db.formas_pagamento.find(query, {"_id": 0}).sort("nome", 1).to_list(1000)
    return formas

@api_router.post("/admin/formas-pagamento")
async def create_forma_pagamento(data: FormaPagamentoCreate, current_user: dict = Depends(get_current_user)):
    forma = {
        "id": str(uuid.uuid4()),
        **data.model_dump(),
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.formas_pagamento.insert_one(forma)
    await create_audit_log(current_user, "create", "forma_pagamento", forma["id"], data.nome)
    del forma["_id"]
    return forma

@api_router.put("/admin/formas-pagamento/{id}")
async def update_forma_pagamento(id: str, data: FormaPagamentoCreate, current_user: dict = Depends(get_current_user)):
    forma = await db.formas_pagamento.find_one({"id": id}, {"_id": 0})
    if not forma:
        raise HTTPException(status_code=404, detail="Forma de pagamento não encontrada")
    
    update_data = data.model_dump()
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.formas_pagamento.update_one({"id": id}, {"$set": update_data})
    await create_audit_log(current_user, "update", "forma_pagamento", id, data.nome)
    
    updated = await db.formas_pagamento.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.delete("/admin/formas-pagamento/{id}")
async def delete_forma_pagamento(id: str, current_user: dict = Depends(get_current_user)):
    forma = await db.formas_pagamento.find_one({"id": id}, {"_id": 0})
    if not forma:
        raise HTTPException(status_code=404, detail="Forma de pagamento não encontrada")
    
    await db.formas_pagamento.delete_one({"id": id})
    await create_audit_log(current_user, "delete", "forma_pagamento", id, forma["nome"])
    return {"message": "Forma de pagamento excluída"}

# --- Contas Bancárias ---
@api_router.get("/admin/contas-bancarias")
async def get_contas_bancarias(
    ativo: Optional[bool] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if ativo is not None:
        query["ativo"] = ativo
    
    contas = await db.contas_bancarias.find(query, {"_id": 0}).sort("nome", 1).to_list(1000)
    return contas

@api_router.post("/admin/contas-bancarias")
async def create_conta_bancaria(data: ContaBancariaCreate, current_user: dict = Depends(get_current_user)):
    conta = {
        "id": str(uuid.uuid4()),
        **data.model_dump(),
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.contas_bancarias.insert_one(conta)
    await create_audit_log(current_user, "create", "conta_bancaria", conta["id"], data.nome)
    del conta["_id"]
    return conta

@api_router.get("/admin/contas-bancarias/{id}")
async def get_conta_bancaria(id: str, current_user: dict = Depends(get_current_user)):
    conta = await db.contas_bancarias.find_one({"id": id}, {"_id": 0})
    if not conta:
        raise HTTPException(status_code=404, detail="Conta bancária não encontrada")
    return conta

@api_router.put("/admin/contas-bancarias/{id}")
async def update_conta_bancaria(id: str, data: ContaBancariaCreate, current_user: dict = Depends(get_current_user)):
    conta = await db.contas_bancarias.find_one({"id": id}, {"_id": 0})
    if not conta:
        raise HTTPException(status_code=404, detail="Conta bancária não encontrada")
    
    update_data = data.model_dump()
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.contas_bancarias.update_one({"id": id}, {"$set": update_data})
    await create_audit_log(current_user, "update", "conta_bancaria", id, data.nome)
    
    updated = await db.contas_bancarias.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.delete("/admin/contas-bancarias/{id}")
async def delete_conta_bancaria(id: str, current_user: dict = Depends(get_current_user)):
    conta = await db.contas_bancarias.find_one({"id": id}, {"_id": 0})
    if not conta:
        raise HTTPException(status_code=404, detail="Conta bancária não encontrada")
    
    await db.contas_bancarias.delete_one({"id": id})
    await create_audit_log(current_user, "delete", "conta_bancaria", id, conta["nome"])
    return {"message": "Conta bancária excluída"}

@api_router.patch("/admin/contas-bancarias/{id}/saldo")
async def update_saldo_conta_bancaria(id: str, saldo: float, current_user: dict = Depends(get_current_user)):
    conta = await db.contas_bancarias.find_one({"id": id}, {"_id": 0})
    if not conta:
        raise HTTPException(status_code=404, detail="Conta bancária não encontrada")
    
    await db.contas_bancarias.update_one({"id": id}, {"$set": {"saldo_atual": saldo, "updated_at": datetime.now(timezone.utc).isoformat()}})
    await create_audit_log(current_user, "update", "conta_bancaria_saldo", id, f"Saldo: {saldo}")
    
    updated = await db.contas_bancarias.find_one({"id": id}, {"_id": 0})
    return updated


# ==================== MOVIMENTAÇÃO DE CONTAS ====================

@api_router.get("/admin/movimentacoes")
async def get_movimentacoes(
    tipo: Optional[str] = None,
    categoria: Optional[str] = None,
    conta_bancaria_id: Optional[str] = None,
    centro_custo_id: Optional[str] = None,
    data_inicio: Optional[str] = None,
    data_fim: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Lista todas as movimentações de contas"""
    query = {}
    if tipo:
        query["tipo"] = tipo
    if categoria:
        query["categoria"] = categoria
    if conta_bancaria_id:
        query["$or"] = [
            {"conta_bancaria_origem_id": conta_bancaria_id},
            {"conta_bancaria_destino_id": conta_bancaria_id}
        ]
    if centro_custo_id:
        if "$or" in query:
            query["$and"] = [
                {"$or": query["$or"]},
                {"$or": [
                    {"centro_custo_origem_id": centro_custo_id},
                    {"centro_custo_destino_id": centro_custo_id}
                ]}
            ]
            del query["$or"]
        else:
            query["$or"] = [
                {"centro_custo_origem_id": centro_custo_id},
                {"centro_custo_destino_id": centro_custo_id}
            ]
    if data_inicio:
        query["data_movimentacao"] = {"$gte": data_inicio}
    if data_fim:
        if "data_movimentacao" in query:
            query["data_movimentacao"]["$lte"] = data_fim
        else:
            query["data_movimentacao"] = {"$lte": data_fim}
    
    movimentacoes = await db.movimentacoes_contas.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return movimentacoes


@api_router.get("/admin/movimentacoes/{id}")
async def get_movimentacao(id: str, current_user: dict = Depends(get_current_user)):
    """Retorna uma movimentação específica"""
    mov = await db.movimentacoes_contas.find_one({"id": id}, {"_id": 0})
    if not mov:
        raise HTTPException(status_code=404, detail="Movimentação não encontrada")
    return mov


@api_router.post("/admin/movimentacoes")
async def create_movimentacao(data: MovimentacaoContaCreate, current_user: dict = Depends(get_current_user)):
    """Cria uma nova movimentação de conta"""
    
    if data.valor <= 0:
        raise HTTPException(status_code=400, detail="Valor deve ser maior que zero")
    
    numero = await get_next_sequence("movimentacoes_contas")
    
    mov = {
        "id": str(uuid.uuid4()),
        "numero": numero,
        **data.model_dump(),
        "created_by": current_user.get("name", current_user.get("email", "")),
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    # Atualizar saldos das contas bancárias
    if data.tipo == "saida" or data.tipo == "transferencia":
        if data.conta_bancaria_origem_id:
            conta_origem = await db.contas_bancarias.find_one({"id": data.conta_bancaria_origem_id})
            if conta_origem:
                novo_saldo = (conta_origem.get("saldo_atual", 0) or 0) - data.valor
                await db.contas_bancarias.update_one(
                    {"id": data.conta_bancaria_origem_id},
                    {"$set": {"saldo_atual": novo_saldo, "updated_at": datetime.now(timezone.utc).isoformat()}}
                )
    
    if data.tipo == "entrada" or data.tipo == "transferencia":
        if data.conta_bancaria_destino_id:
            conta_destino = await db.contas_bancarias.find_one({"id": data.conta_bancaria_destino_id})
            if conta_destino:
                novo_saldo = (conta_destino.get("saldo_atual", 0) or 0) + data.valor
                await db.contas_bancarias.update_one(
                    {"id": data.conta_bancaria_destino_id},
                    {"$set": {"saldo_atual": novo_saldo, "updated_at": datetime.now(timezone.utc).isoformat()}}
                )
    
    await db.movimentacoes_contas.insert_one(mov)
    
    await create_audit_log(
        current_user, "create", "movimentacao_conta", mov["id"],
        f"{data.tipo.upper()}: {data.descricao} - R$ {data.valor:.2f}"
    )
    
    del mov["_id"]
    return mov


@api_router.delete("/admin/movimentacoes/{id}")
async def delete_movimentacao(id: str, current_user: dict = Depends(get_current_user)):
    """Exclui uma movimentação e reverte os saldos"""
    mov = await db.movimentacoes_contas.find_one({"id": id}, {"_id": 0})
    if not mov:
        raise HTTPException(status_code=404, detail="Movimentação não encontrada")
    
    # Reverter saldos
    if mov.get("tipo") == "saida" or mov.get("tipo") == "transferencia":
        if mov.get("conta_bancaria_origem_id"):
            conta_origem = await db.contas_bancarias.find_one({"id": mov["conta_bancaria_origem_id"]})
            if conta_origem:
                novo_saldo = (conta_origem.get("saldo_atual", 0) or 0) + mov["valor"]
                await db.contas_bancarias.update_one(
                    {"id": mov["conta_bancaria_origem_id"]},
                    {"$set": {"saldo_atual": novo_saldo}}
                )
    
    if mov.get("tipo") == "entrada" or mov.get("tipo") == "transferencia":
        if mov.get("conta_bancaria_destino_id"):
            conta_destino = await db.contas_bancarias.find_one({"id": mov["conta_bancaria_destino_id"]})
            if conta_destino:
                novo_saldo = (conta_destino.get("saldo_atual", 0) or 0) - mov["valor"]
                await db.contas_bancarias.update_one(
                    {"id": mov["conta_bancaria_destino_id"]},
                    {"$set": {"saldo_atual": novo_saldo}}
                )
    
    await db.movimentacoes_contas.delete_one({"id": id})
    
    await create_audit_log(
        current_user, "delete", "movimentacao_conta", id,
        f"Movimentação excluída: {mov['descricao']}"
    )
    
    return {"message": "Movimentação excluída e saldos revertidos"}


# ==================== IMPORTAÇÃO MANUAL DE NF ====================

@api_router.post("/nf/importar-manual")
async def importar_nf_manual(data: ImportacaoManualNFCreate, current_user: dict = Depends(get_current_user)):
    """Importa uma NF manualmente quando a SEFAZ falha"""
    
    # Verificar se a nota já existe (pela chave de acesso ou número)
    filtro_existente = {}
    if data.chave_acesso:
        filtro_existente["chave_acesso"] = data.chave_acesso
    else:
        filtro_existente["$and"] = [
            {"numero_nota": data.numero_nota},
            {"cnpj_emitente": data.cnpj_emitente.replace(".", "").replace("/", "").replace("-", "")}
        ]
    
    if data.tipo_nota == "nfe":
        existente = await db.nfe_importadas.find_one(filtro_existente)
    else:
        existente = await db.nfse_importadas.find_one(filtro_existente)
    
    if existente:
        raise HTTPException(status_code=400, detail="Esta nota fiscal já foi importada anteriormente")
    
    nf_id = str(uuid.uuid4())
    cnpj_limpo = data.cnpj_emitente.replace(".", "").replace("/", "").replace("-", "")
    
    nf_doc = {
        "id": nf_id,
        "numero_nota": data.numero_nota,
        "serie": data.serie or "1",
        "chave_acesso": data.chave_acesso,
        "data_emissao": data.data_emissao,
        "cnpj_emitente": cnpj_limpo,
        "razao_social_emitente": data.razao_social_emitente,
        "uf_emitente": data.uf_emitente,
        "cnpj_destinatario": data.cnpj_destinatario,
        "razao_social_destinatario": data.razao_social_destinatario,
        "valor_total": data.valor_total,
        "valor_produtos": data.valor_produtos or data.valor_total,
        "valor_servicos": data.valor_servicos,
        "valor_frete": data.valor_frete or 0,
        "valor_desconto": data.valor_desconto or 0,
        "centro_custo_id": data.centro_custo_id,
        "centro_custo_nome": data.centro_custo_nome,
        "plano_conta_id": data.plano_conta_id,
        "plano_conta_nome": data.plano_conta_nome,
        "xml_base64": data.xml_base64,
        "pdf_base64": data.pdf_base64,
        "observacoes": data.observacoes,
        "importacao_manual": True,
        "status": "nova",
        "itens": [],
        "created_by": current_user.get("name", current_user.get("email", "")),
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    if data.tipo_nota == "nfe":
        await db.nfe_importadas.insert_one(nf_doc)
        entity_type = "nfe_manual"
    else:
        await db.nfse_importadas.insert_one(nf_doc)
        entity_type = "nfse_manual"
    
    await create_audit_log(
        current_user, "create", entity_type, nf_id,
        f"NF {data.numero_nota} - {data.razao_social_emitente} - R$ {data.valor_total:.2f}"
    )
    
    return {
        "message": f"{'NF-e' if data.tipo_nota == 'nfe' else 'NFS-e'} importada manualmente com sucesso",
        "id": nf_id,
        "numero_nota": data.numero_nota
    }


class XMLExtractRequest(BaseModel):
    xml_base64: str


@api_router.post("/nf/extrair-xml")
async def extrair_dados_xml(data: XMLExtractRequest, current_user: dict = Depends(get_current_user)):
    """Extrai dados de um arquivo XML de NF-e (portalfiscal) ou NFS-e (ABRASF)
    para preenchimento automático no upload manual."""
    from xml.etree import ElementTree as ET

    try:
        xml_content = base64.b64decode(data.xml_base64).decode('utf-8', errors='replace')
        root = ET.fromstring(xml_content)

        # ============ Detecta NFS-e (ABRASF) e parseia ============
        # Heurística: presença de <CompNfse>/<InfNfse>/<Nfse> ou namespace ABRASF.
        is_nfse = False
        for el in root.iter():
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag in ("InfNfse", "CompNfse", "Nfse", "ConsultarNfseServicoTomadoResposta"):
                is_nfse = True
                break

        def first_text(elem, tag_name):
            """Busca o primeiro descendente com tag local==tag_name e retorna texto."""
            for child in elem.iter():
                ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if ctag == tag_name and child.text:
                    return child.text.strip()
            return ""

        if is_nfse:
            # Encontra o nó da nota
            inf_nfse = None
            for el in root.iter():
                tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
                if tag == "InfNfse":
                    inf_nfse = el
                    break
            base = inf_nfse if inf_nfse is not None else root

            # PRESTADOR (emitente)
            prestador = None
            tomador = None
            for el in base.iter():
                tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
                if tag == "PrestadorServico":
                    prestador = el
                elif tag == "TomadorServico":
                    tomador = el

            cnpj_prest = first_text(prestador, "Cnpj") if prestador is not None else first_text(base, "Cnpj")
            razao_prest = first_text(prestador, "RazaoSocial") if prestador is not None else ""
            cnpj_tom = first_text(tomador, "Cnpj") if tomador is not None else ""
            razao_tom = first_text(tomador, "RazaoSocial") if tomador is not None else ""

            valor_servicos_str = first_text(base, "ValorServicos") or first_text(base, "ValorLiquidoNfse") or "0"
            valor_total_str = valor_servicos_str
            try:
                valor_servicos = float(valor_servicos_str.replace(",", "."))
                valor_total = float(valor_total_str.replace(",", "."))
            except (ValueError, TypeError):
                valor_servicos = valor_total = 0.0

            data_emissao_raw = first_text(base, "DataEmissao") or ""
            data_emissao = data_emissao_raw[:10] if data_emissao_raw else ""

            discriminacao = first_text(base, "Discriminacao") or ""
            codigo_verif = first_text(base, "CodigoVerificacao") or ""
            numero = first_text(base, "Numero") or ""

            return {
                "sucesso": True,
                "tipo_nota": "nfse",
                "numero_nota": numero,
                "serie": first_text(base, "Serie") or "1",
                "chave_acesso": codigo_verif,
                "data_emissao": data_emissao,
                "cnpj_emitente": cnpj_prest,
                "razao_social_emitente": razao_prest,
                "uf_emitente": "",
                "ie_emitente": first_text(prestador, "InscricaoMunicipal") if prestador is not None else "",
                "cnpj_destinatario": cnpj_tom,
                "razao_social_destinatario": razao_tom,
                "valor_total": valor_total,
                "valor_produtos": 0.0,
                "valor_servicos": valor_servicos,
                "valor_frete": 0.0,
                "valor_desconto": 0.0,
                "observacoes": discriminacao[:500] if discriminacao else "",
                "itens": [],
            }

        # ============ NF-e (portalfiscal — mantido) ============
        ns = {
            'nfe': 'http://www.portalfiscal.inf.br/nfe',
            'ns': 'http://www.portalfiscal.inf.br/nfe'
        }
        nfe = root.find('.//nfe:NFe', ns) or root.find('.//NFe') or root
        inf_nfe = nfe.find('.//nfe:infNFe', ns) or nfe.find('.//infNFe') or root.find('.//infNFe')

        resultado = {
            "sucesso": True,
            "tipo_nota": "nfe",
            "numero_nota": "",
            "serie": "1",
            "chave_acesso": "",
            "data_emissao": "",
            "cnpj_emitente": "",
            "razao_social_emitente": "",
            "uf_emitente": "",
            "ie_emitente": "",
            "cnpj_destinatario": "",
            "razao_social_destinatario": "",
            "valor_total": 0,
            "valor_produtos": 0,
            "valor_frete": 0,
            "valor_desconto": 0,
            "itens": []
        }
        
        if inf_nfe is None:
            # Pode ser um resumo de NF-e (resNFe)
            res_nfe = root.find('.//nfe:resNFe', ns) or root.find('.//resNFe')
            if res_nfe is not None:
                chave = res_nfe.findtext('.//nfe:chNFe', '', ns) or res_nfe.findtext('.//chNFe', '')
                resultado["chave_acesso"] = chave
                resultado["cnpj_emitente"] = res_nfe.findtext('.//nfe:CNPJ', '', ns) or res_nfe.findtext('.//CNPJ', '')
                resultado["razao_social_emitente"] = res_nfe.findtext('.//nfe:xNome', '', ns) or res_nfe.findtext('.//xNome', '')
                valor = res_nfe.findtext('.//nfe:vNF', '0', ns) or res_nfe.findtext('.//vNF', '0')
                resultado["valor_total"] = float(valor) if valor else 0
                resultado["valor_produtos"] = resultado["valor_total"]
                data_emissao = res_nfe.findtext('.//nfe:dhEmi', '', ns) or res_nfe.findtext('.//dhEmi', '')
                resultado["data_emissao"] = data_emissao[:10] if data_emissao else ""
                
                # Extrair número e série da chave de acesso
                if len(chave) >= 34:
                    resultado["numero_nota"] = chave[25:34].lstrip('0')
                    resultado["serie"] = chave[22:25].lstrip('0') or "1"
                    resultado["uf_emitente"] = chave[0:2]
                
                return resultado
            else:
                return {"sucesso": False, "erro": "Estrutura do XML não reconhecida. Não foi possível encontrar os dados da NF-e."}
        
        # Extrair chave de acesso
        chave = inf_nfe.get('Id', '').replace('NFe', '') if inf_nfe.get('Id') else ''
        resultado["chave_acesso"] = chave
        
        # Extrair UF da chave
        if len(chave) >= 2:
            resultado["uf_emitente"] = chave[0:2]
        
        # Dados de identificação (ide)
        ide = inf_nfe.find('.//nfe:ide', ns) or inf_nfe.find('.//ide')
        if ide is not None:
            resultado["numero_nota"] = ide.findtext('.//nfe:nNF', '', ns) or ide.findtext('.//nNF', '')
            resultado["serie"] = ide.findtext('.//nfe:serie', '1', ns) or ide.findtext('.//serie', '1')
            data_emissao = ide.findtext('.//nfe:dhEmi', '', ns) or ide.findtext('.//dhEmi', '')
            resultado["data_emissao"] = data_emissao[:10] if data_emissao else ""
        
        # Dados do emitente (emit)
        emit = inf_nfe.find('.//nfe:emit', ns) or inf_nfe.find('.//emit')
        if emit is not None:
            resultado["cnpj_emitente"] = emit.findtext('.//nfe:CNPJ', '', ns) or emit.findtext('.//CNPJ', '')
            resultado["razao_social_emitente"] = emit.findtext('.//nfe:xNome', '', ns) or emit.findtext('.//xNome', '')
            resultado["ie_emitente"] = emit.findtext('.//nfe:IE', '', ns) or emit.findtext('.//IE', '')
            
            # Endereço do emitente para UF
            ender_emit = emit.find('.//nfe:enderEmit', ns) or emit.find('.//enderEmit')
            if ender_emit is not None:
                uf = ender_emit.findtext('.//nfe:UF', '', ns) or ender_emit.findtext('.//UF', '')
                if uf:
                    resultado["uf_emitente"] = uf
        
        # Dados do destinatário (dest)
        dest = inf_nfe.find('.//nfe:dest', ns) or inf_nfe.find('.//dest')
        if dest is not None:
            resultado["cnpj_destinatario"] = dest.findtext('.//nfe:CNPJ', '', ns) or dest.findtext('.//CNPJ', '') or dest.findtext('.//nfe:CPF', '', ns) or dest.findtext('.//CPF', '')
            resultado["razao_social_destinatario"] = dest.findtext('.//nfe:xNome', '', ns) or dest.findtext('.//xNome', '')
        
        # Totais
        total = inf_nfe.find('.//nfe:total', ns) or inf_nfe.find('.//total')
        if total is not None:
            icms_tot = total.find('.//nfe:ICMSTot', ns) or total.find('.//ICMSTot')
            if icms_tot is not None:
                valor_nf = icms_tot.findtext('.//nfe:vNF', '0', ns) or icms_tot.findtext('.//vNF', '0')
                valor_prod = icms_tot.findtext('.//nfe:vProd', '0', ns) or icms_tot.findtext('.//vProd', '0')
                valor_frete = icms_tot.findtext('.//nfe:vFrete', '0', ns) or icms_tot.findtext('.//vFrete', '0')
                valor_desc = icms_tot.findtext('.//nfe:vDesc', '0', ns) or icms_tot.findtext('.//vDesc', '0')
                
                resultado["valor_total"] = float(valor_nf) if valor_nf else 0
                resultado["valor_produtos"] = float(valor_prod) if valor_prod else 0
                resultado["valor_frete"] = float(valor_frete) if valor_frete else 0
                resultado["valor_desconto"] = float(valor_desc) if valor_desc else 0
        
        # Itens da nota
        det_list = inf_nfe.findall('.//nfe:det', ns) or inf_nfe.findall('.//det')
        for det in det_list:
            prod = det.find('.//nfe:prod', ns) or det.find('.//prod')
            if prod is not None:
                item = {
                    "codigo": prod.findtext('.//nfe:cProd', '', ns) or prod.findtext('.//cProd', ''),
                    "descricao": prod.findtext('.//nfe:xProd', '', ns) or prod.findtext('.//xProd', ''),
                    "ncm": prod.findtext('.//nfe:NCM', '', ns) or prod.findtext('.//NCM', ''),
                    "cfop": prod.findtext('.//nfe:CFOP', '', ns) or prod.findtext('.//CFOP', ''),
                    "unidade": prod.findtext('.//nfe:uCom', '', ns) or prod.findtext('.//uCom', ''),
                    "quantidade": float(prod.findtext('.//nfe:qCom', '0', ns) or prod.findtext('.//qCom', '0')),
                    "valor_unitario": float(prod.findtext('.//nfe:vUnCom', '0', ns) or prod.findtext('.//vUnCom', '0')),
                    "valor_total": float(prod.findtext('.//nfe:vProd', '0', ns) or prod.findtext('.//vProd', '0'))
                }
                resultado["itens"].append(item)
        
        return resultado
        
    except ET.ParseError as e:
        return {"sucesso": False, "erro": f"Erro ao processar XML: Formato inválido - {str(e)}"}
    except Exception as e:
        logging.error(f"Erro ao extrair dados do XML: {str(e)}")
        return {"sucesso": False, "erro": f"Erro ao processar XML: {str(e)}"}


# --- Aluguéis de Máquinas ---
@api_router.get("/admin/alugueis")
async def get_alugueis(
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    query = {}
    if status:
        query["status"] = status
    
    alugueis = await db.alugueis.find(query, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return alugueis

@api_router.get("/admin/alugueis/{id}")
async def get_aluguel(id: str, current_user: dict = Depends(get_current_user)):
    aluguel = await db.alugueis.find_one({"id": id}, {"_id": 0})
    if not aluguel:
        raise HTTPException(status_code=404, detail="Aluguel não encontrado")
    return aluguel

@api_router.post("/admin/alugueis")
async def create_aluguel(data: AluguelCreate, current_user: dict = Depends(get_current_user)):
    numero = await get_next_sequence("alugueis")
    
    aluguel = {
        "id": str(uuid.uuid4()),
        "numero": numero,
        **data.model_dump(exclude={"gerar_conta_receber"}),
        "conta_receber_id": None,
        "status": "ativo",  # ativo, finalizado, cancelado
        "created_by": current_user["id"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    # Alterar status da máquina para "operacional"
    if data.maquina_id:
        await db.machines.update_one(
            {"id": data.maquina_id},
            {"$set": {"status": "operacional"}}
        )
    
    # Gerar conta a receber automaticamente se solicitado
    if data.gerar_conta_receber:
        periodo_label = {
            "diaria": "Diária", "semanal": "Semanal", "quinzenal": "Quinzenal",
            "mensal": "Mensal", "semestral": "Semestral", "anual": "Anual",
            "hora": "Por Hora", "outro": data.periodo_especificado or "Outro"
        }.get(data.tipo_periodo, data.tipo_periodo)
        
        conta_receber = {
            "id": str(uuid.uuid4()),
            "numero": await get_next_sequence("contas_receber"),
            "cliente_nome": data.cliente_nome,
            "documento": data.cliente_documento,
            "descricao": f"Aluguel #{numero} - {data.maquina_nome or 'Máquina'} ({periodo_label})",
            "valor": data.valor,
            "valor_final": data.valor,
            "data_emissao": data.data_entrega,
            "data_vencimento": data.data_vencimento,
            "status": "em_aberto",
            "forma_pagamento": "boleto",
            "observacoes": f"Gerado automaticamente do aluguel #{numero}",
            "aluguel_id": aluguel["id"],
            "created_by": current_user["id"],
            "created_at": datetime.now(timezone.utc).isoformat()
        }
        await db.contas_receber.insert_one(conta_receber)
        aluguel["conta_receber_id"] = conta_receber["id"]
        del conta_receber["_id"]
    
    await db.alugueis.insert_one(aluguel)
    await create_audit_log(current_user, "create", "aluguel", aluguel["id"], f"Aluguel #{numero}")
    del aluguel["_id"]
    return aluguel

@api_router.put("/admin/alugueis/{id}")
async def update_aluguel(id: str, data: AluguelCreate, current_user: dict = Depends(get_current_user)):
    aluguel = await db.alugueis.find_one({"id": id}, {"_id": 0})
    if not aluguel:
        raise HTTPException(status_code=404, detail="Aluguel não encontrado")
    
    update_data = data.model_dump(exclude={"gerar_conta_receber"})
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    await db.alugueis.update_one({"id": id}, {"$set": update_data})
    
    # Atualizar conta a receber se existir
    if aluguel.get("conta_receber_id"):
        await db.contas_receber.update_one(
            {"id": aluguel["conta_receber_id"]},
            {"$set": {
                "cliente_nome": data.cliente_nome,
                "valor": data.valor,
                "valor_final": data.valor,
                "data_vencimento": data.data_vencimento
            }}
        )
    
    await create_audit_log(current_user, "update", "aluguel", id, f"Aluguel #{aluguel['numero']}")
    
    updated = await db.alugueis.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.patch("/admin/alugueis/{id}/status")
async def update_aluguel_status(id: str, status_data: dict, current_user: dict = Depends(get_current_user)):
    aluguel = await db.alugueis.find_one({"id": id}, {"_id": 0})
    if not aluguel:
        raise HTTPException(status_code=404, detail="Aluguel não encontrado")
    
    new_status = status_data.get("status")
    update_data = {"status": new_status, "updated_at": datetime.now(timezone.utc).isoformat()}
    
    if new_status == "finalizado" and status_data.get("data_devolucao"):
        update_data["data_devolucao"] = status_data.get("data_devolucao")
    
    await db.alugueis.update_one({"id": id}, {"$set": update_data})
    
    # Se finalizado ou cancelado, voltar status da máquina para "patio"
    if new_status in ["finalizado", "cancelado"] and aluguel.get("maquina_id"):
        await db.machines.update_one(
            {"id": aluguel["maquina_id"]},
            {"$set": {"status": "patio"}}
        )
    
    # Se finalizado, marcar conta a receber como quitada
    if new_status == "finalizado" and aluguel.get("conta_receber_id"):
        await db.contas_receber.update_one(
            {"id": aluguel["conta_receber_id"]},
            {"$set": {"status": "quitada", "data_recebimento": datetime.now(timezone.utc).strftime("%Y-%m-%d")}}
        )
    
    await create_audit_log(current_user, "update", "aluguel", id, f"Status: {new_status}")
    
    updated = await db.alugueis.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.delete("/admin/alugueis/{id}")
async def delete_aluguel(id: str, current_user: dict = Depends(get_current_user)):
    aluguel = await db.alugueis.find_one({"id": id}, {"_id": 0})
    if not aluguel:
        raise HTTPException(status_code=404, detail="Aluguel não encontrado")
    
    # Excluir conta a receber associada
    if aluguel.get("conta_receber_id"):
        await db.contas_receber.delete_one({"id": aluguel["conta_receber_id"]})
    
    await db.alugueis.delete_one({"id": id})
    await create_audit_log(current_user, "delete", "aluguel", id, f"Aluguel #{aluguel['numero']}")
    return {"message": "Aluguel excluído"}

# --- Upload contrato de aluguel ---
CONTRATOS_DIR = ROOT_DIR / "uploads" / "contratos"
CONTRATOS_DIR.mkdir(parents=True, exist_ok=True)

@api_router.post("/admin/alugueis/{id}/contrato")
async def upload_contrato(
    id: str, 
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload de arquivo de contrato para um aluguel"""
    aluguel = await db.alugueis.find_one({"id": id}, {"_id": 0})
    if not aluguel:
        raise HTTPException(status_code=404, detail="Aluguel não encontrado")
    
    # Validate file type
    allowed_extensions = ['.pdf', '.doc', '.docx', '.jpg', '.jpeg', '.png']
    ext = Path(file.filename).suffix.lower() if file.filename else ''
    if ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Tipo de arquivo não permitido. Use: {', '.join(allowed_extensions)}")
    
    # Read file content
    content = await file.read()
    max_size = 50 * 1024 * 1024  # 50MB
    if len(content) > max_size:
        raise HTTPException(status_code=400, detail="Arquivo muito grande. Máximo: 50MB")
    
    # Save file
    unique_filename = f"{id}_{uuid.uuid4()}{ext}"
    file_path = CONTRATOS_DIR / unique_filename
    
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Update aluguel
    await db.alugueis.update_one(
        {"id": id},
        {"$set": {
            "contrato_arquivo": unique_filename,
            "contrato_nome": file.filename,
            "contrato_uploaded_at": datetime.now(timezone.utc).isoformat()
        }}
    )
    
    await create_audit_log(current_user, "upload", "contrato", id, f"Contrato do aluguel #{aluguel.get('numero', id)}")
    return {"message": "Contrato anexado com sucesso", "filename": unique_filename}

@api_router.get("/admin/alugueis/{id}/contrato/download")
async def download_contrato(
    id: str,
    current_user: dict = Depends(get_current_user)
):
    """Download de arquivo de contrato"""
    aluguel = await db.alugueis.find_one({"id": id}, {"_id": 0})
    if not aluguel:
        raise HTTPException(status_code=404, detail="Aluguel não encontrado")
    
    if not aluguel.get("contrato_arquivo"):
        raise HTTPException(status_code=404, detail="Contrato não encontrado")
    
    file_path = CONTRATOS_DIR / aluguel["contrato_arquivo"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
    
    return FileResponse(
        path=str(file_path),
        filename=aluguel.get("contrato_nome", aluguel["contrato_arquivo"]),
        media_type="application/octet-stream"
    )

# --- IMÓVEIS PARA LOCAÇÃO ---
class ImovelCreate(BaseModel):
    tipo_imovel: str = "apartamento"
    descricao: str
    endereco: str
    numero: Optional[str] = ""
    complemento: Optional[str] = ""
    bairro: Optional[str] = ""
    cidade: Optional[str] = ""
    estado: Optional[str] = "TO"
    cep: Optional[str] = ""
    area_m2: Optional[float] = 0
    quartos: Optional[int] = 0
    banheiros: Optional[int] = 0
    vagas_garagem: Optional[int] = 0
    cliente_nome: str
    cliente_telefone: Optional[str] = ""
    cliente_documento: Optional[str] = ""
    numero_contrato: Optional[str] = ""
    tipo_periodo: str = "mensal"
    periodo_especificado: Optional[str] = ""
    data_inicio: str
    data_vencimento: Optional[str] = ""
    valor_aluguel: float
    valor_condominio: Optional[float] = 0
    valor_iptu: Optional[float] = 0
    valor_caucao: Optional[float] = 0
    dia_vencimento: Optional[int] = 10
    observacoes: Optional[str] = ""
    gerar_conta_receber: bool = True

@api_router.get("/admin/imoveis")
async def get_imoveis(current_user: dict = Depends(get_current_user)):
    imoveis = await db.imoveis.find({}, {"_id": 0}).sort("created_at", -1).to_list(1000)
    return imoveis

@api_router.get("/admin/imoveis/{id}")
async def get_imovel(id: str, current_user: dict = Depends(get_current_user)):
    imovel = await db.imoveis.find_one({"id": id}, {"_id": 0})
    if not imovel:
        raise HTTPException(status_code=404, detail="Imóvel não encontrado")
    return imovel

@api_router.post("/admin/imoveis")
async def create_imovel(imovel: ImovelCreate, current_user: dict = Depends(get_current_user)):
    imovel_dict = imovel.model_dump()
    imovel_dict["id"] = str(uuid.uuid4())
    imovel_dict["status"] = "ativo" if imovel_dict["cliente_nome"] else "pendente"
    imovel_dict["created_at"] = datetime.now(timezone.utc).isoformat()
    imovel_dict["created_by"] = current_user.get("user_id")
    
    # Gerar conta a receber
    if imovel_dict["gerar_conta_receber"] and imovel_dict["valor_aluguel"] > 0:
        valor_total = imovel_dict["valor_aluguel"] + imovel_dict.get("valor_condominio", 0) + imovel_dict.get("valor_iptu", 0)
        
        # Calcular próximo vencimento
        hoje = datetime.now()
        dia_venc = imovel_dict.get("dia_vencimento", 10)
        if hoje.day > dia_venc:
            prox_mes = hoje.month + 1 if hoje.month < 12 else 1
            prox_ano = hoje.year if hoje.month < 12 else hoje.year + 1
        else:
            prox_mes = hoje.month
            prox_ano = hoje.year
        
        data_vencimento = f"{prox_ano}-{prox_mes:02d}-{dia_venc:02d}"
        
        conta_receber = {
            "id": str(uuid.uuid4()),
            "descricao": f"Aluguel - {imovel_dict['descricao']}",
            "cliente_nome": imovel_dict["cliente_nome"],
            "cliente_documento": imovel_dict.get("cliente_documento", ""),
            "valor": valor_total,
            "valor_final": valor_total,
            "data_vencimento": data_vencimento,
            "status": "em_aberto",
            "origem": "imovel",
            "imovel_id": imovel_dict["id"],
            "created_at": datetime.now(timezone.utc).isoformat(),
            "created_by": current_user.get("user_id")
        }
        await db.contas_receber.insert_one(conta_receber)
        imovel_dict["conta_receber_id"] = conta_receber["id"]
    
    await db.imoveis.insert_one(imovel_dict)
    await create_audit_log(current_user, "create", "imovel", imovel_dict["id"], f"Imóvel: {imovel_dict['descricao']}")
    
    created = await db.imoveis.find_one({"id": imovel_dict["id"]}, {"_id": 0})
    return created

@api_router.put("/admin/imoveis/{id}")
async def update_imovel(id: str, imovel: ImovelCreate, current_user: dict = Depends(get_current_user)):
    existing = await db.imoveis.find_one({"id": id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Imóvel não encontrado")
    
    update_data = imovel.model_dump()
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.imoveis.update_one({"id": id}, {"$set": update_data})
    await create_audit_log(current_user, "update", "imovel", id, f"Imóvel: {update_data['descricao']}")
    
    updated = await db.imoveis.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.patch("/admin/imoveis/{id}/status")
async def update_imovel_status(id: str, status_data: dict, current_user: dict = Depends(get_current_user)):
    imovel = await db.imoveis.find_one({"id": id}, {"_id": 0})
    if not imovel:
        raise HTTPException(status_code=404, detail="Imóvel não encontrado")
    
    new_status = status_data.get("status")
    await db.imoveis.update_one({"id": id}, {"$set": {"status": new_status, "updated_at": datetime.now(timezone.utc).isoformat()}})
    await create_audit_log(current_user, "update", "imovel", id, f"Status: {new_status}")
    
    updated = await db.imoveis.find_one({"id": id}, {"_id": 0})
    return updated

@api_router.delete("/admin/imoveis/{id}")
async def delete_imovel(id: str, current_user: dict = Depends(get_current_user)):
    imovel = await db.imoveis.find_one({"id": id}, {"_id": 0})
    if not imovel:
        raise HTTPException(status_code=404, detail="Imóvel não encontrado")
    
    if imovel.get("conta_receber_id"):
        await db.contas_receber.delete_one({"id": imovel["conta_receber_id"]})
    
    await db.imoveis.delete_one({"id": id})
    await create_audit_log(current_user, "delete", "imovel", id, f"Imóvel: {imovel['descricao']}")
    return {"message": "Imóvel excluído"}

CONTRATOS_IMOVEIS_DIR = ROOT_DIR / "uploads" / "contratos_imoveis"
CONTRATOS_IMOVEIS_DIR.mkdir(parents=True, exist_ok=True)

@api_router.post("/admin/imoveis/{id}/contrato")
async def upload_contrato_imovel(id: str, file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    imovel = await db.imoveis.find_one({"id": id}, {"_id": 0})
    if not imovel:
        raise HTTPException(status_code=404, detail="Imóvel não encontrado")
    
    ext = Path(file.filename).suffix.lower()
    allowed_extensions = [".pdf", ".jpg", ".jpeg", ".png", ".doc", ".docx"]
    if ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Tipo não permitido. Use: {', '.join(allowed_extensions)}")
    
    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Arquivo muito grande. Máximo: 50MB")
    
    unique_filename = f"{id}_{uuid.uuid4()}{ext}"
    file_path = CONTRATOS_IMOVEIS_DIR / unique_filename
    with open(file_path, "wb") as f:
        f.write(content)
    
    await db.imoveis.update_one({"id": id}, {"$set": {
        "contrato_arquivo": unique_filename,
        "contrato_nome": file.filename,
        "contrato_uploaded_at": datetime.now(timezone.utc).isoformat()
    }})
    
    return {"message": "Contrato anexado", "filename": unique_filename}

# --- Buscar máquinas do sistema de gerenciamento ---
@api_router.get("/admin/maquinas-disponiveis")
async def get_maquinas_disponiveis(current_user: dict = Depends(get_current_user)):
    """Retorna lista de máquinas cadastradas no sistema de gerenciamento"""
    maquinas = await db.machines.find({}, {"_id": 0}).sort("name", 1).to_list(1000)
    return maquinas

# --- Notificações ---
@api_router.get("/admin/notificacoes")
async def get_notificacoes(
    prazo_dias: int = 7,
    current_user: dict = Depends(get_current_user)
):
    """Retorna itens com vencimento próximo"""
    from datetime import timedelta
    hoje = datetime.now(timezone.utc)
    hoje_str = hoje.strftime("%Y-%m-%d")
    limite = (hoje + timedelta(days=prazo_dias)).strftime("%Y-%m-%d")
    
    notificacoes = []
    
    # Contas a Pagar próximas do vencimento
    contas_pagar = await db.contas_pagar.find({
        "status": "em_aberto",
        "data_vencimento": {"$lte": limite}
    }, {"_id": 0}).to_list(100)
    
    for c in contas_pagar:
        vencida = c.get("data_vencimento", "") < hoje_str
        notificacoes.append({
            "tipo": "conta_pagar",
            "id": c.get("id"),
            "titulo": c.get("descricao") or f"Conta #{c.get('numero')}",
            "subtitulo": c.get("fornecedor_nome"),
            "valor": c.get("valor_final") or c.get("valor"),
            "data": c.get("data_vencimento"),
            "vencida": vencida,
            "urgencia": "alta" if vencida else ("media" if c.get("data_vencimento", "") <= (hoje + timedelta(days=3)).strftime("%Y-%m-%d") else "baixa"),
            "link": "/administrativo/a-pagar"
        })
    
    # Contas a Receber próximas do vencimento
    contas_receber = await db.contas_receber.find({
        "status": "em_aberto",
        "data_vencimento": {"$lte": limite}
    }, {"_id": 0}).to_list(100)
    
    for c in contas_receber:
        vencida = c.get("data_vencimento", "") < hoje_str
        notificacoes.append({
            "tipo": "conta_receber",
            "id": c.get("id"),
            "titulo": c.get("descricao") or f"Conta #{c.get('numero')}",
            "subtitulo": c.get("cliente_nome"),
            "valor": c.get("valor_final") or c.get("valor"),
            "data": c.get("data_vencimento"),
            "vencida": vencida,
            "urgencia": "alta" if vencida else ("media" if c.get("data_vencimento", "") <= (hoje + timedelta(days=3)).strftime("%Y-%m-%d") else "baixa"),
            "link": "/administrativo/a-receber"
        })
    
    # Ordens de Serviço com previsão de entrega próxima
    ordens = await db.ordens_servico.find({
        "status": {"$in": ["em_aberto", "em_andamento"]},
        "data_previsao_entrega": {"$lte": limite, "$ne": None, "$ne": ""}
    }, {"_id": 0}).to_list(100)
    
    for o in ordens:
        vencida = o.get("data_previsao_entrega", "") < hoje_str
        notificacoes.append({
            "tipo": "ordem_servico",
            "id": o.get("id"),
            "titulo": f"OS #{o.get('numero')} - {o.get('descricao', 'Sem descrição')[:50]}",
            "subtitulo": o.get("cliente_nome") or o.get("cliente_fantasia"),
            "valor": o.get("valor_total"),
            "data": o.get("data_previsao_entrega"),
            "vencida": vencida,
            "urgencia": "alta" if vencida else ("media" if o.get("data_previsao_entrega", "") <= (hoje + timedelta(days=3)).strftime("%Y-%m-%d") else "baixa"),
            "link": "/administrativo/ordens-servico"
        })
    
    # Aluguéis com vencimento próximo
    alugueis = await db.alugueis.find({
        "status": "ativo",
        "data_vencimento": {"$lte": limite}
    }, {"_id": 0}).to_list(100)
    
    for a in alugueis:
        vencida = a.get("data_vencimento", "") < hoje_str
        notificacoes.append({
            "tipo": "aluguel",
            "id": a.get("id"),
            "titulo": f"Aluguel #{a.get('numero')} - {a.get('maquina_nome', 'Máquina')}",
            "subtitulo": a.get("cliente_nome"),
            "valor": a.get("valor"),
            "data": a.get("data_vencimento"),
            "vencida": vencida,
            "urgencia": "alta" if vencida else ("media" if a.get("data_vencimento", "") <= (hoje + timedelta(days=3)).strftime("%Y-%m-%d") else "baixa"),
            "link": "/administrativo/alugueis"
        })
    
    # Ordenar por data e urgência
    notificacoes.sort(key=lambda x: (
        0 if x["urgencia"] == "alta" else (1 if x["urgencia"] == "media" else 2),
        x.get("data", "9999-99-99")
    ))
    
    # Resumo
    resumo = {
        "total": len(notificacoes),
        "vencidas": len([n for n in notificacoes if n["vencida"]]),
        "alta": len([n for n in notificacoes if n["urgencia"] == "alta"]),
        "media": len([n for n in notificacoes if n["urgencia"] == "media"]),
        "baixa": len([n for n in notificacoes if n["urgencia"] == "baixa"]),
        "por_tipo": {
            "conta_pagar": len([n for n in notificacoes if n["tipo"] == "conta_pagar"]),
            "conta_receber": len([n for n in notificacoes if n["tipo"] == "conta_receber"]),
            "ordem_servico": len([n for n in notificacoes if n["tipo"] == "ordem_servico"]),
            "aluguel": len([n for n in notificacoes if n["tipo"] == "aluguel"])
        }
    }
    
    return {
        "resumo": resumo,
        "notificacoes": notificacoes,
        "prazo_dias": prazo_dias
    }

@api_router.get("/admin/notificacoes/contagem")
async def get_notificacoes_contagem(
    prazo_dias: int = 7,
    current_user: dict = Depends(get_current_user)
):
    """Retorna apenas a contagem de notificações (para o badge)"""
    from datetime import timedelta
    hoje = datetime.now(timezone.utc)
    hoje_str = hoje.strftime("%Y-%m-%d")
    limite = (hoje + timedelta(days=prazo_dias)).strftime("%Y-%m-%d")
    
    count_pagar = await db.contas_pagar.count_documents({
        "status": "em_aberto",
        "data_vencimento": {"$lte": limite}
    })
    
    count_receber = await db.contas_receber.count_documents({
        "status": "em_aberto",
        "data_vencimento": {"$lte": limite}
    })
    
    count_os = await db.ordens_servico.count_documents({
        "status": {"$in": ["em_aberto", "em_andamento"]},
        "data_previsao_entrega": {"$lte": limite, "$ne": None, "$ne": ""}
    })
    
    count_alugueis = await db.alugueis.count_documents({
        "status": "ativo",
        "data_vencimento": {"$lte": limite}
    })
    
    # Contar vencidas
    count_vencidas = await db.contas_pagar.count_documents({
        "status": "em_aberto",
        "data_vencimento": {"$lt": hoje_str}
    }) + await db.contas_receber.count_documents({
        "status": "em_aberto",
        "data_vencimento": {"$lt": hoje_str}
    }) + await db.alugueis.count_documents({
        "status": "ativo",
        "data_vencimento": {"$lt": hoje_str}
    })
    
    total = count_pagar + count_receber + count_os + count_alugueis
    
    return {
        "total": total,
        "vencidas": count_vencidas,
        "prazo_dias": prazo_dias
    }


# ============ PAINEL ADMINISTRATIVO (GESTÃO DE USUÁRIOS) ============

@api_router.get("/admin-panel/users")
async def get_all_users(current_user: dict = Depends(get_current_user)):
    """Lista todos os usuários da plataforma"""
    users = []
    cursor = db.users.find({}, {"password": 0})
    async for user in cursor:
        users.append({
            "id": user.get("id"),
            "name": user.get("name"),
            "email": user.get("email"),
            "role": user.get("role", "gerenciamento"),
            "created_at": user.get("created_at"),
            "last_login": user.get("last_login")
        })
    return users

class UserCreateAdmin(BaseModel):
    name: str
    email: EmailStr
    password: str
    role: str = "gerenciamento"  # gerenciamento, administrativo, ambos, admin

@api_router.post("/admin-panel/users")
async def create_user_admin(data: UserCreateAdmin, current_user: dict = Depends(get_current_user)):
    """Cria um novo usuário (apenas via painel admin)"""
    # Check if email already exists
    existing = await db.users.find_one({"email": data.email})
    if existing:
        raise HTTPException(status_code=400, detail="Email já cadastrado")
    
    # Validate role - includes all combination roles and custom roles like 'programador'
    valid_roles = [
        "gerenciamento", "administrativo", "rh", "ambos", 
        "ambos_rh", "gerenciamento_rh", "administrativo_rh", 
        "admin", "programador"
    ]
    if data.role not in valid_roles:
        raise HTTPException(status_code=400, detail=f"Tipo de acesso inválido. Opções: {', '.join(valid_roles)}")
    
    # Hash password
    hashed_password = bcrypt.hashpw(data.password.encode('utf-8'), bcrypt.gensalt())
    
    user_id = str(uuid.uuid4())
    user_doc = {
        "id": user_id,
        "name": data.name,
        "email": data.email,
        "password": hashed_password.decode('utf-8'),
        "role": data.role,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "last_login": None
    }
    
    await db.users.insert_one(user_doc)
    
    # Traduzir role para português
    role_labels = {
        "gerenciamento": "Gerenciamento",
        "administrativo": "Administrativo",
        "rh": "RH",
        "ambos": "Gerenciamento + Administrativo",
        "ambos_rh": "Ger + Admin + RH",
        "gerenciamento_rh": "Gerenciamento + RH",
        "administrativo_rh": "Administrativo + RH",
        "admin": "Administrador",
        "programador": "Programador"
    }
    
    # Registrar na auditoria
    await db.audit_logs.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": current_user["id"],
        "user_name": current_user["name"],
        "action": f"Criou usuário: {data.name}",
        "details": f"Email: {data.email}\nTipo de Acesso: {role_labels.get(data.role, data.role)}",
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    return {"message": "Usuário criado com sucesso", "id": user_id}

class UserRoleUpdate(BaseModel):
    role: str

@api_router.patch("/admin-panel/users/{user_id}/role")
async def update_user_role(user_id: str, data: UserRoleUpdate, current_user: dict = Depends(get_current_user)):
    """Atualiza o role/permissões de um usuário"""
    # Roles com acesso total (admin e programador)
    admin_roles = ["admin", "programador"]
    if current_user.get("role") not in admin_roles:
        raise HTTPException(status_code=403, detail="Apenas administradores podem alterar permissões")
    
    # All valid roles including combination roles and custom roles
    valid_roles = [
        "gerenciamento", "administrativo", "rh", "ambos", 
        "ambos_rh", "gerenciamento_rh", "administrativo_rh", 
        "admin", "programador"
    ]
    if data.role not in valid_roles:
        raise HTTPException(status_code=400, detail=f"Role inválido. Opções: {', '.join(valid_roles)}")
    
    user = await db.users.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")
    
    old_role = user.get("role", "gerenciamento")
    
    await db.users.update_one(
        {"id": user_id},
        {"$set": {"role": data.role}}
    )
    
    # Registrar na auditoria
    await create_audit_log(
        user=current_user,
        action="alterar permissão",
        entity_type="usuário",
        entity_id=user_id,
        entity_name=user.get("name"),
        details=f"Permissão alterada de '{old_role}' para '{data.role}'",
        module="Painel Admin"
    )
    
    return {"message": "Permissão atualizada com sucesso", "new_role": data.role}

@api_router.delete("/admin-panel/users/{user_id}")
async def delete_user_admin(user_id: str, current_user: dict = Depends(get_current_user)):
    """Exclui um usuário"""
    if user_id == current_user["id"]:
        raise HTTPException(status_code=400, detail="Você não pode excluir sua própria conta")
    
    user = await db.users.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")
    
    await db.users.delete_one({"id": user_id})
    
    # Registrar na auditoria
    await db.audit_logs.insert_one({
        "id": str(uuid.uuid4()),
        "user_id": current_user["id"],
        "user_name": current_user["name"],
        "action": f"Excluiu usuário: {user.get('name')}",
        "details": f"Email: {user.get('email')}",
        "created_at": datetime.now(timezone.utc).isoformat()
    })
    
    return {"message": "Usuário excluído com sucesso"}


class UserPasswordReset(BaseModel):
    new_password: Optional[str] = None  # Se não vier, gera uma senha aleatória


@api_router.post("/admin-panel/users/{user_id}/reset-password")
async def reset_user_password(
    user_id: str,
    data: UserPasswordReset,
    current_user: dict = Depends(get_current_user),
):
    """Admin reseta a senha de um usuário. Se `new_password` não vier no body,
    é gerada uma senha aleatória de 12 caracteres e devolvida na resposta para que
    o admin compartilhe com o usuário. A nova senha é armazenada com bcrypt."""
    admin_roles = ["admin", "programador"]
    if current_user.get("role") not in admin_roles:
        raise HTTPException(status_code=403, detail="Apenas administradores podem resetar senhas")

    user = await db.users.find_one({"id": user_id})
    if not user:
        raise HTTPException(status_code=404, detail="Usuário não encontrado")

    # Validar nova senha (mínimo 6 caracteres) ou gerar uma
    if data.new_password:
        if len(data.new_password) < 6:
            raise HTTPException(status_code=400, detail="A senha deve ter pelo menos 6 caracteres")
        nova_senha = data.new_password
        gerada = False
    else:
        import secrets
        import string
        alfabeto = string.ascii_letters + string.digits
        nova_senha = "".join(secrets.choice(alfabeto) for _ in range(12))
        gerada = True

    await db.users.update_one(
        {"id": user_id},
        {"$set": {
            "password": hash_password(nova_senha),
            "password_reset_at": datetime.now(timezone.utc).isoformat(),
            "password_reset_by": current_user.get("id"),
        }},
    )

    await create_audit_log(
        user=current_user,
        action="resetar senha",
        entity_type="usuário",
        entity_id=user_id,
        entity_name=user.get("name"),
        details=f"Senha do usuário '{user.get('email')}' foi resetada pelo administrador"
                + (" (senha gerada automaticamente)" if gerada else " (senha definida manualmente)"),
        module="Painel Admin",
    )

    return {
        "message": "Senha resetada com sucesso. Compartilhe a nova senha com o usuário com segurança.",
        "new_password": nova_senha,
        "gerada_automaticamente": gerada,
    }


@api_router.get("/admin-panel/audit-logs")
async def get_audit_logs(current_user: dict = Depends(get_current_user)):
    """Retorna todos os logs de auditoria"""
    logs = []
    cursor = db.audit_logs.find({}).sort("created_at", -1).limit(500)
    async for log in cursor:
        logs.append({
            "id": log.get("id"),
            "user_id": log.get("user_id"),
            "user_name": log.get("user_name"),
            "user_email": log.get("user_email", ""),
            "action": log.get("action"),
            "entity_type": log.get("entity_type", ""),
            "entity_name": log.get("entity_name", ""),
            "module": log.get("module", "Sistema"),
            "details": log.get("details"),
            "created_at": log.get("created_at")
        })
    return logs

@api_router.get("/admin-panel/users/{user_id}/activities")
async def get_user_activities(user_id: str, current_user: dict = Depends(get_current_user)):
    """Retorna as atividades de um usuário específico"""
    activities = []
    cursor = db.audit_logs.find({"user_id": user_id}).sort("created_at", -1).limit(100)
    async for log in cursor:
        activities.append({
            "id": log.get("id"),
            "user_id": log.get("user_id"),
            "user_name": log.get("user_name"),
            "action": log.get("action"),
            "entity_type": log.get("entity_type", ""),
            "entity_name": log.get("entity_name", ""),
            "module": log.get("module", "Sistema"),
            "details": log.get("details"),
            "created_at": log.get("created_at")
        })
    return activities


# ============ DATABASE MANAGER ROUTES (ADMIN ONLY) ============

# Roles com acesso administrativo total
ADMIN_ROLES = ["admin", "programador"]

def require_admin(user: dict):
    """Verifica se o usuário é admin ou programador"""
    if user.get("role") not in ADMIN_ROLES:
        raise HTTPException(status_code=403, detail="Acesso negado. Somente administradores podem acessar este recurso.")

# Lista de coleções disponíveis para gerenciamento
ALLOWED_COLLECTIONS = [
    "users", "machines", "maintenances", "categories", "stock_items", 
    "stock_categories", "stock_movements", "obras", "contas_pagar", 
    "contas_receber", "cadastros", "produtos_admin", "ordens_servico", 
    "alugueis", "audit_logs", "usage_logs"
]

@api_router.get("/admin-panel/database/{collection_name}")
async def get_collection_documents(
    collection_name: str, 
    page: int = 1, 
    limit: int = 20, 
    search: str = "",
    current_user: dict = Depends(get_current_user)
):
    """Retorna documentos de uma coleção específica com paginação e busca"""
    require_admin(current_user)
    
    if collection_name not in ALLOWED_COLLECTIONS:
        raise HTTPException(status_code=400, detail=f"Coleção '{collection_name}' não permitida")
    
    collection = db[collection_name]
    skip = (page - 1) * limit
    
    # Build search query
    query = {}
    if search:
        # Busca em campos comuns - usar $exists para evitar erros
        query = {
            "$or": [
                {"name": {"$regex": search, "$options": "i"}},
                {"email": {"$regex": search, "$options": "i"}},
                {"titulo": {"$regex": search, "$options": "i"}},
                {"descricao": {"$regex": search, "$options": "i"}},
                {"action": {"$regex": search, "$options": "i"}},
                {"id": {"$regex": search, "$options": "i"}}
            ]
        }
    
    try:
        # Get total count
        total = await collection.count_documents(query if search else {})
        
        # Get documents - tentar ordenar por created_at, se falhar usar _id
        documents = []
        try:
            cursor = collection.find(query if search else {}, {"_id": 0}).sort("created_at", -1).skip(skip).limit(limit)
            async for doc in cursor:
                documents.append(doc)
        except Exception:
            # Fallback sem ordenação
            cursor = collection.find(query if search else {}, {"_id": 0}).skip(skip).limit(limit)
            async for doc in cursor:
                documents.append(doc)
        
        return {
            "documents": documents,
            "total": total,
            "page": page,
            "limit": limit
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao buscar documentos: {str(e)}")

@api_router.post("/admin-panel/database/{collection_name}")
async def create_document(
    collection_name: str,
    document: dict,
    current_user: dict = Depends(get_current_user)
):
    """Cria um novo documento em uma coleção"""
    require_admin(current_user)
    
    if collection_name not in ALLOWED_COLLECTIONS:
        raise HTTPException(status_code=400, detail=f"Coleção '{collection_name}' não permitida")
    
    collection = db[collection_name]
    
    # Adiciona ID e timestamp se não existirem
    if "id" not in document:
        document["id"] = str(uuid.uuid4())
    if "created_at" not in document:
        document["created_at"] = datetime.now(timezone.utc).isoformat()
    
    await collection.insert_one(document)
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="criar",
        entity_type=f"documento ({collection_name})",
        entity_id=document["id"],
        entity_name=document.get("name") or document.get("email") or document.get("titulo") or document["id"][:8],
        details=f"Documento criado via painel admin"
    )
    
    # Remove _id antes de retornar
    document.pop("_id", None)
    return {"message": "Documento criado com sucesso", "document": document}

@api_router.put("/admin-panel/database/{collection_name}/{doc_id}")
async def update_document(
    collection_name: str,
    doc_id: str,
    document: dict,
    current_user: dict = Depends(get_current_user)
):
    """Atualiza um documento existente"""
    require_admin(current_user)
    
    if collection_name not in ALLOWED_COLLECTIONS:
        raise HTTPException(status_code=400, detail=f"Coleção '{collection_name}' não permitida")
    
    collection = db[collection_name]
    
    # Encontra o documento
    existing = await collection.find_one({"id": doc_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    
    # Remove campos que não devem ser alterados diretamente
    document.pop("_id", None)
    
    # Atualiza o documento
    await collection.update_one({"id": doc_id}, {"$set": document})
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="editar",
        entity_type=f"documento ({collection_name})",
        entity_id=doc_id,
        entity_name=document.get("name") or document.get("email") or document.get("titulo") or doc_id[:8],
        details=f"Documento editado via painel admin"
    )
    
    return {"message": "Documento atualizado com sucesso"}

@api_router.delete("/admin-panel/database/{collection_name}/{doc_id}")
async def delete_document(
    collection_name: str,
    doc_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Exclui um documento de uma coleção"""
    require_admin(current_user)
    
    if collection_name not in ALLOWED_COLLECTIONS:
        raise HTTPException(status_code=400, detail=f"Coleção '{collection_name}' não permitida")
    
    # Proteção especial para usuários admin
    if collection_name == "users":
        user_doc = await db.users.find_one({"id": doc_id}, {"_id": 0})
        if user_doc and user_doc.get("id") == current_user["id"]:
            raise HTTPException(status_code=400, detail="Você não pode excluir sua própria conta")
    
    collection = db[collection_name]
    
    # Encontra o documento para log
    existing = await collection.find_one({"id": doc_id}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    
    # Exclui o documento
    await collection.delete_one({"id": doc_id})
    
    # Audit log
    await create_audit_log(
        user=current_user,
        action="excluir",
        entity_type=f"documento ({collection_name})",
        entity_id=doc_id,
        entity_name=existing.get("name") or existing.get("email") or existing.get("titulo") or doc_id[:8],
        details=f"Documento excluído via painel admin"
    )
    
    return {"message": "Documento excluído com sucesso"}


# ============ CHATBOT ROUTES (AI ASSISTANT) ============

from pydantic import BaseModel as PydanticBaseModel

class ChatMessage(PydanticBaseModel):
    message: str
    module: str = "gerenciamento"  # gerenciamento ou administrativo

class ChatResponse(PydanticBaseModel):
    response: str
    context_used: List[str] = []

async def get_full_platform_context() -> str:
    """Coleta TODAS as informações de TODAS as coleções do banco de dados"""
    context_parts = []
    
    context_parts.append("=" * 60)
    context_parts.append("BANCO DE DADOS COMPLETO - CRA CONSTRUTORA")
    context_parts.append("=" * 60)
    
    # ========== USUÁRIOS ==========
    users = await db.users.find({}, {"_id": 0, "password": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nUSUÁRIOS DO SISTEMA ({len(users)} registros)\n{'='*40}")
    for u in users:
        context_parts.append(f"- Nome: {u.get('name')} | Email: {u.get('email')} | Tipo: {u.get('role', 'gerenciamento')} | Criado: {u.get('created_at', '-')[:10] if u.get('created_at') else '-'}")
    
    # ========== CATEGORIAS DE MÁQUINAS ==========
    categories = await db.categories.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nCATEGORIAS DE MÁQUINAS ({len(categories)} registros)\n{'='*40}")
    for c in categories:
        context_parts.append(f"- ID: {c.get('id', '-')[:8]} | Nome: {c.get('name')} | Descrição: {c.get('description', '-')}")
    
    # ========== MÁQUINAS ==========
    machines = await db.machines.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nMÁQUINAS CADASTRADAS ({len(machines)} registros)\n{'='*40}")
    for m in machines:
        status = "Operacional" if m.get("status") == "operational" else "Em manutenção"
        context_parts.append(f"- ID: {m.get('id', '-')[:8]} | Nome: {m.get('name')} | Placa: {m.get('plate')} | Marca: {m.get('brand', '-')} | Modelo: {m.get('model', '-')} | Ano: {m.get('year', '-')} | Status: {status} | Horas desde troca óleo: {m.get('hours_since_oil_change', 0)}")
    
    # ========== MANUTENÇÕES ==========
    maintenances = await db.maintenances.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nMANUTENÇÕES ({len(maintenances)} registros)\n{'='*40}")
    # Calcular totais
    total_valor = sum(m.get('part_value', 0) for m in maintenances)
    preventivas = [m for m in maintenances if m.get('maintenance_type') == 'preventiva']
    corretivas = [m for m in maintenances if m.get('maintenance_type') == 'corretiva']
    context_parts.append(f"RESUMO: Total gasto: R$ {total_valor:.2f} | Preventivas: {len(preventivas)} | Corretivas: {len(corretivas)}")
    for m in maintenances:
        tipo = "Preventiva" if m.get("maintenance_type") == "preventiva" else "Corretiva"
        context_parts.append(f"- Peça: {m.get('part_name')} | Tipo: {tipo} | Valor: R$ {m.get('part_value', 0):.2f} | Data: {m.get('replacement_date', '-')} | Máquina ID: {m.get('machine_id', '-')[:8] if m.get('machine_id') else '-'} | Troca óleo: {'Sim' if m.get('is_oil_change') else 'Não'}")
    
    # ========== ESTOQUE - CATEGORIAS ==========
    stock_categories = await db.stock_categories.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nCATEGORIAS DE ESTOQUE ({len(stock_categories)} registros)\n{'='*40}")
    for c in stock_categories:
        context_parts.append(f"- Nome: {c.get('name')}")
    
    # ========== ESTOQUE - ITENS ==========
    stock_items = await db.stock_items.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nITENS DE ESTOQUE ({len(stock_items)} registros)\n{'='*40}")
    low_stock = [i for i in stock_items if i.get("quantity", 0) <= i.get("min_quantity", 0)]
    context_parts.append(f"ALERTA: {len(low_stock)} itens com estoque baixo!")
    for i in stock_items:
        status = "⚠️ BAIXO" if i.get("quantity", 0) <= i.get("min_quantity", 0) else "OK"
        context_parts.append(f"- Nome: {i.get('name')} | Código: {i.get('code', '-')} | Categoria: {i.get('category', '-')} | Qtd: {i.get('quantity', 0)} {i.get('unit', 'un')} | Mínimo: {i.get('min_quantity', 0)} | Preço unit: R$ {i.get('unit_price', 0):.2f} | Local: {i.get('location', '-')} | Status: {status}")
    
    # ========== ESTOQUE - MOVIMENTAÇÕES ==========
    stock_movements = await db.stock_movements.find({}, {"_id": 0}).sort("created_at", -1).to_list(100)
    context_parts.append(f"\n\n{'='*40}\nMOVIMENTAÇÕES DE ESTOQUE (últimas {len(stock_movements)})\n{'='*40}")
    for m in stock_movements:
        tipo = "ENTRADA" if m.get("movement_type") == "entrada" else "SAÍDA"
        context_parts.append(f"- {tipo}: {m.get('quantity')} unidades | Item ID: {m.get('item_id', '-')[:8] if m.get('item_id') else '-'} | Motivo: {m.get('reason', '-')} | Data: {m.get('created_at', '-')[:10] if m.get('created_at') else '-'}")
    
    # ========== OBRAS ==========
    obras = await db.obras.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nOBRAS/PROJETOS ({len(obras)} registros)\n{'='*40}")
    for o in obras:
        status_map = {"em_andamento": "Em andamento", "concluida": "Concluída", "pausada": "Pausada"}
        status = status_map.get(o.get("status", ""), o.get("status", ""))
        context_parts.append(f"- Nome: {o.get('name')} | Local: {o.get('location', '-')} | Status: {status} | Início: {o.get('start_date', '-')} | Fim: {o.get('end_date', '-')} | Descrição: {o.get('description', '-')}")
    
    # ========== REGISTROS DE USO (HORÍMETRO) ==========
    usage_logs = await db.usage_logs.find({}, {"_id": 0}).sort("created_at", -1).to_list(100)
    context_parts.append(f"\n\n{'='*40}\nREGISTROS DE USO/HORÍMETRO (últimos {len(usage_logs)})\n{'='*40}")
    for u in usage_logs:
        context_parts.append(f"- Máquina ID: {u.get('machine_id', '-')[:8] if u.get('machine_id') else '-'} | Horas: {u.get('hours', 0)} | Data: {u.get('created_at', '-')[:10] if u.get('created_at') else '-'} | Obs: {u.get('notes', '-')}")
    
    # ========== CADASTROS (CLIENTES/FORNECEDORES) ==========
    cadastros = await db.cadastros.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nCADASTROS - CLIENTES/FORNECEDORES ({len(cadastros)} registros)\n{'='*40}")
    clientes = [c for c in cadastros if c.get('tipo') == 'cliente']
    fornecedores = [c for c in cadastros if c.get('tipo') == 'fornecedor']
    context_parts.append(f"RESUMO: {len(clientes)} clientes | {len(fornecedores)} fornecedores")
    for c in cadastros:
        context_parts.append(f"- Tipo: {c.get('tipo', 'cliente').upper()} | Nome/Razão: {c.get('nome_razao')} | CPF/CNPJ: {c.get('cpf_cnpj', '-')} | Tel: {c.get('telefone', '-')} | Email: {c.get('email', '-')} | Cidade: {c.get('cidade', '-')}/{c.get('estado', '-')}")
    
    # ========== CONTAS A PAGAR ==========
    contas_pagar = await db.contas_pagar.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nCONTAS A PAGAR ({len(contas_pagar)} registros)\n{'='*40}")
    pendentes_pagar = [c for c in contas_pagar if c.get("status") == "pendente"]
    quitadas_pagar = [c for c in contas_pagar if c.get("status") == "quitada"]
    total_pendente_pagar = sum(c.get("valor", 0) for c in pendentes_pagar)
    total_quitado_pagar = sum(c.get("valor", 0) for c in quitadas_pagar)
    context_parts.append(f"RESUMO: Pendentes: {len(pendentes_pagar)} (R$ {total_pendente_pagar:.2f}) | Quitadas: {len(quitadas_pagar)} (R$ {total_quitado_pagar:.2f})")
    for c in contas_pagar:
        context_parts.append(f"- Descrição: {c.get('descricao')} | Valor: R$ {c.get('valor', 0):.2f} | Vencimento: {c.get('data_vencimento', '-')} | Status: {c.get('status', 'pendente').upper()} | Fornecedor: {c.get('fornecedor_nome', '-')} | Categoria: {c.get('plano_conta_nome', '-')}")
    
    # ========== CONTAS A RECEBER ==========
    contas_receber = await db.contas_receber.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nCONTAS A RECEBER ({len(contas_receber)} registros)\n{'='*40}")
    pendentes_receber = [c for c in contas_receber if c.get("status") == "pendente"]
    quitadas_receber = [c for c in contas_receber if c.get("status") == "quitada"]
    total_pendente_receber = sum(c.get("valor", 0) for c in pendentes_receber)
    total_quitado_receber = sum(c.get("valor", 0) for c in quitadas_receber)
    context_parts.append(f"RESUMO: Pendentes: {len(pendentes_receber)} (R$ {total_pendente_receber:.2f}) | Recebidas: {len(quitadas_receber)} (R$ {total_quitado_receber:.2f})")
    for c in contas_receber:
        context_parts.append(f"- Descrição: {c.get('descricao')} | Valor: R$ {c.get('valor', 0):.2f} | Vencimento: {c.get('data_vencimento', '-')} | Status: {c.get('status', 'pendente').upper()} | Cliente: {c.get('cliente_nome', '-')} | Categoria: {c.get('plano_conta_nome', '-')}")
    
    # ========== PRODUTOS ==========
    produtos = await db.produtos_admin.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nPRODUTOS ({len(produtos)} registros)\n{'='*40}")
    for p in produtos:
        context_parts.append(f"- Código: {p.get('codigo', '-')} | Descrição: {p.get('descricao')} | Unidade: {p.get('unidade', '-')} | Preço: R$ {p.get('preco', 0):.2f} | Estoque: {p.get('estoque', 0)}")
    
    # ========== ORDENS DE SERVIÇO ==========
    ordens = await db.ordens_servico.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nORDENS DE SERVIÇO ({len(ordens)} registros)\n{'='*40}")
    abertas = [o for o in ordens if o.get("status") in ["aberta", "em_andamento"]]
    concluidas = [o for o in ordens if o.get("status") == "concluida"]
    context_parts.append(f"RESUMO: Abertas/Em andamento: {len(abertas)} | Concluídas: {len(concluidas)}")
    for o in ordens:
        context_parts.append(f"- OS Nº {o.get('numero')} | Descrição: {o.get('descricao', '-')[:60]} | Cliente: {o.get('cliente_nome', '-')} | Valor: R$ {o.get('valor_total', 0):.2f} | Status: {o.get('status', '-').upper()} | Tipo Financeiro: {o.get('tipo_financeiro', '-')}")
    
    # ========== PLANO DE CONTAS ==========
    plano_contas = await db.plano_contas.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nPLANO DE CONTAS ({len(plano_contas)} registros)\n{'='*40}")
    for p in plano_contas:
        tipo = "📥 RECEITA" if p.get("tipo") == "receita" else "📤 DESPESA"
        context_parts.append(f"- {tipo} | Código: {p.get('codigo', '-')} | Nome: {p.get('nome')} | Pai: {p.get('pai_nome', 'Raiz')}")
    
    # ========== CENTROS DE CUSTO ==========
    centros_custo = await db.centros_custo.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nCENTROS DE CUSTO ({len(centros_custo)} registros)\n{'='*40}")
    for c in centros_custo:
        context_parts.append(f"- Código: {c.get('codigo', '-')} | Nome: {c.get('nome')} | Descrição: {c.get('descricao', '-')}")
    
    # ========== FORMAS DE PAGAMENTO ==========
    formas_pagamento = await db.formas_pagamento.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nFORMAS DE PAGAMENTO ({len(formas_pagamento)} registros)\n{'='*40}")
    for f in formas_pagamento:
        context_parts.append(f"- Nome: {f.get('nome')} | Descrição: {f.get('descricao', '-')}")
    
    # ========== ALUGUÉIS DE MÁQUINAS ==========
    alugueis = await db.alugueis.find({}, {"_id": 0}).to_list(500)
    context_parts.append(f"\n\n{'='*40}\nALUGUÉIS DE MÁQUINAS ({len(alugueis)} registros)\n{'='*40}")
    ativos = [a for a in alugueis if a.get("status") == "ativo"]
    finalizados = [a for a in alugueis if a.get("status") == "finalizado"]
    total_alugueis = sum(a.get("valor_total", 0) for a in alugueis)
    context_parts.append(f"RESUMO: Ativos: {len(ativos)} | Finalizados: {len(finalizados)} | Valor total: R$ {total_alugueis:.2f}")
    for a in alugueis:
        context_parts.append(f"- Máquina: {a.get('maquina_nome', '-')} | Cliente: {a.get('cliente_nome', '-')} | Tel: {a.get('cliente_telefone', '-')} | Período: {a.get('tipo_periodo', '-')} | Valor: R$ {a.get('valor_total', 0):.2f} | Status: {a.get('status', '-').upper()} | Entrega: {a.get('data_entrega', '-')} | Vencimento: {a.get('data_vencimento', '-')}")
    
    # ========== LOGS DE AUDITORIA (ÚLTIMOS) ==========
    audit_logs = await db.audit_logs.find({}, {"_id": 0}).sort("created_at", -1).to_list(50)
    context_parts.append(f"\n\n{'='*40}\nÚLTIMAS ATIVIDADES/AUDITORIA ({len(audit_logs)} registros)\n{'='*40}")
    for a in audit_logs:
        context_parts.append(f"- {a.get('created_at', '-')[:16] if a.get('created_at') else '-'} | Usuário: {a.get('user_name', '-')} | Ação: {a.get('action', '-')} | Módulo: {a.get('module', '-')}")
    
    # ========== RESUMO FINANCEIRO GERAL ==========
    context_parts.append(f"\n\n{'='*60}")
    context_parts.append("RESUMO FINANCEIRO GERAL")
    context_parts.append(f"{'='*60}")
    context_parts.append(f"Total gasto em manutenções: R$ {total_valor:.2f}")
    context_parts.append(f"Contas a pagar pendentes: R$ {total_pendente_pagar:.2f}")
    context_parts.append(f"Contas a receber pendentes: R$ {total_pendente_receber:.2f}")
    context_parts.append(f"Total em aluguéis: R$ {total_alugueis:.2f}")
    saldo_projetado = total_pendente_receber - total_pendente_pagar
    context_parts.append(f"Saldo projetado (a receber - a pagar): R$ {saldo_projetado:.2f}")
    
    return "\n".join(context_parts)

@api_router.post("/chatbot/ask", response_model=ChatResponse)
async def chatbot_ask(chat: ChatMessage, current_user: dict = Depends(get_current_user)):
    """Endpoint do chatbot que responde perguntas sobre a plataforma"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    
    try:
        # Coletar contexto COMPLETO da plataforma
        platform_context = await get_full_platform_context()
        
        # Sistema message com contexto completo
        system_message = f"""Você é o assistente virtual inteligente da CRA Construtora.
Você tem ACESSO COMPLETO E TOTAL a TODAS as informações do banco de dados da plataforma.

DADOS COMPLETOS DO SISTEMA (ATUALIZADOS EM TEMPO REAL):
{platform_context}

SUAS CAPACIDADES:
- Você conhece TODOS os usuários, máquinas, manutenções, estoque, obras, contas, ordens de serviço, aluguéis, etc.
- Você pode calcular totais, médias, fazer comparações e análises
- Você sabe quais itens estão com estoque baixo
- Você conhece o histórico financeiro completo
- Você pode identificar padrões e fazer recomendações

INSTRUÇÕES DE FORMATAÇÃO (MUITO IMPORTANTE):
1. SEMPRE responda em português brasileiro
2. Use QUEBRAS DE LINHA para separar parágrafos e seções
3. Use listas com "•" para enumerar itens (não use asteriscos)
4. Formate valores monetários como R$ 1.234,56
5. Organize informações em seções claras quando houver muitos dados
6. NÃO use markdown com asteriscos (**, *)
7. Seja claro, direto e bem organizado
8. Separe diferentes tipos de informação com linhas em branco

EXEMPLO DE FORMATAÇÃO CORRETA:
Resumo do Estoque:

• Item 1: 50 unidades - R$ 25,90 cada
• Item 2: 30 unidades - R$ 15,00 cada

Total de itens: 2
Valor total em estoque: R$ 1.745,00

INSTRUÇÕES DE CONTEÚDO:
1. SEMPRE use os dados REAIS fornecidos acima
2. Seja ESPECÍFICO - cite nomes, valores, datas exatos
3. Se não encontrar a informação, diga claramente
4. Faça cálculos quando necessário
5. Seja proativo em fornecer informações úteis relacionadas
"""
        
        # Inicializar chat com Gemini
        llm_key = os.environ.get("EMERGENT_LLM_KEY")
        
        llm_chat = LlmChat(
            api_key=llm_key,
            session_id=f"chatbot-{current_user['id']}-{datetime.now().strftime('%Y%m%d%H%M')}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.0-flash")
        
        # Enviar mensagem
        user_message = UserMessage(text=chat.message)
        response = await llm_chat.send_message(user_message)
        
        return ChatResponse(
            response=response,
            context_used=["todos_os_modulos"]
        )
        
    except Exception as e:
        logging.error(f"Erro no chatbot: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Erro ao processar pergunta: {str(e)}")

@api_router.post("/chatbot/ask-with-files", response_model=ChatResponse)
async def chatbot_ask_with_files(
    message: str = Form(...),
    module: str = Form("gerenciamento"),
    files: List[UploadFile] = File(None),
    current_user: dict = Depends(get_current_user)
):
    """Endpoint do chatbot que processa arquivos anexados com extração de conteúdo"""
    from emergentintegrations.llm.chat import LlmChat, UserMessage
    import io
    
    try:
        # Processar arquivos anexados
        files_info = []
        file_contents = []
        
        if files:
            for file in files:
                content = await file.read()
                file_size = len(content)
                filename = file.filename or "arquivo_sem_nome"
                content_type = file.content_type or ""
                
                # Informações básicas do arquivo
                files_info.append({
                    "nome": filename,
                    "tipo": content_type,
                    "tamanho": f"{file_size / 1024:.1f} KB"
                })
                
                # Extrair conteúdo baseado no tipo de arquivo
                extracted_content = None
                
                # Arquivos de texto
                if 'text' in content_type or filename.endswith(('.txt', '.csv', '.json', '.xml', '.md', '.log')):
                    try:
                        text_content = content.decode('utf-8')
                        extracted_content = f"📄 CONTEÚDO DE {filename}:\n{text_content[:5000]}"
                        if len(text_content) > 5000:
                            extracted_content += "\n[...conteúdo truncado...]"
                    except:
                        extracted_content = f"⚠️ Arquivo {filename}: não foi possível decodificar como texto"
                
                # PDFs - extrair texto
                elif filename.lower().endswith('.pdf') or 'pdf' in content_type:
                    try:
                        from PyPDF2 import PdfReader
                        pdf_reader = PdfReader(io.BytesIO(content))
                        pdf_text = ""
                        for i, page in enumerate(pdf_reader.pages[:10]):  # Máximo 10 páginas
                            page_text = page.extract_text()
                            if page_text:
                                pdf_text += f"\n--- Página {i+1} ---\n{page_text}"
                        if pdf_text.strip():
                            extracted_content = f"📑 CONTEÚDO DO PDF {filename}:{pdf_text[:8000]}"
                            if len(pdf_text) > 8000:
                                extracted_content += "\n[...conteúdo truncado...]"
                        else:
                            extracted_content = f"📑 PDF {filename}: {len(pdf_reader.pages)} páginas (texto não extraível - pode ser imagem/escaneado)"
                    except Exception as pdf_err:
                        extracted_content = f"⚠️ PDF {filename}: erro ao extrair texto ({str(pdf_err)[:100]})"
                
                # Documentos Word
                elif filename.lower().endswith(('.docx', '.doc')):
                    try:
                        from docx import Document
                        doc = Document(io.BytesIO(content))
                        doc_text = "\n".join([para.text for para in doc.paragraphs])
                        if doc_text.strip():
                            extracted_content = f"📝 CONTEÚDO DO DOCUMENTO {filename}:\n{doc_text[:5000]}"
                            if len(doc_text) > 5000:
                                extracted_content += "\n[...conteúdo truncado...]"
                        else:
                            extracted_content = f"📝 Documento {filename}: sem texto extraível"
                    except Exception as doc_err:
                        extracted_content = f"⚠️ Documento {filename}: erro ao extrair ({str(doc_err)[:100]})"
                
                # Planilhas Excel
                elif filename.lower().endswith(('.xlsx', '.xls')):
                    try:
                        import openpyxl
                        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
                        excel_content = []
                        for sheet_name in wb.sheetnames[:3]:  # Máximo 3 abas
                            sheet = wb[sheet_name]
                            sheet_data = f"\n--- Aba: {sheet_name} ---\n"
                            row_count = 0
                            for row in sheet.iter_rows(max_row=50, values_only=True):  # Máximo 50 linhas
                                row_values = [str(cell) if cell is not None else "" for cell in row]
                                if any(row_values):
                                    sheet_data += " | ".join(row_values) + "\n"
                                    row_count += 1
                            excel_content.append(sheet_data)
                        extracted_content = f"📊 CONTEÚDO DA PLANILHA {filename}:{''.join(excel_content)}"
                    except Exception as xl_err:
                        extracted_content = f"⚠️ Planilha {filename}: erro ao extrair ({str(xl_err)[:100]})"
                
                # Imagens - descrever que é uma imagem
                elif content_type.startswith('image/') or filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp')):
                    try:
                        from PIL import Image
                        img = Image.open(io.BytesIO(content))
                        extracted_content = f"🖼️ IMAGEM {filename}: formato {img.format}, dimensões {img.width}x{img.height} pixels, modo {img.mode}"
                    except:
                        extracted_content = f"🖼️ Imagem {filename}: arquivo de imagem ({file_size / 1024:.1f} KB)"
                
                # Outros arquivos
                else:
                    extracted_content = f"📁 Arquivo {filename}: tipo {content_type or 'desconhecido'} ({file_size / 1024:.1f} KB)"
                
                if extracted_content:
                    file_contents.append(extracted_content)
                
                # Reset file pointer
                await file.seek(0)
        
        # Construir contexto dos arquivos
        files_context = ""
        if files_info:
            files_context = "\n\n═══════════════════════════════════════════════════════════════\n"
            files_context += "ARQUIVOS ANEXADOS PELO USUÁRIO:\n"
            files_context += "═══════════════════════════════════════════════════════════════\n"
            for info in files_info:
                files_context += f"• {info['nome']} ({info['tipo'] or 'tipo desconhecido'}, {info['tamanho']})\n"
            
            if file_contents:
                files_context += "\n" + "\n\n".join(file_contents)
            files_context += "\n═══════════════════════════════════════════════════════════════\n"
        
        # Coletar contexto da plataforma
        platform_context = await get_full_platform_context()
        
        # Sistema message com contexto
        system_message = f"""Você é o assistente virtual inteligente da CRA Construtora.
Você tem ACESSO COMPLETO E TOTAL a TODAS as informações do banco de dados da plataforma.

DADOS COMPLETOS DO SISTEMA:
{platform_context}

{files_context}

SUAS CAPACIDADES:
- Você pode analisar arquivos anexados (PDFs, documentos Word, planilhas Excel, textos)
- Você pode extrair e analisar informações de documentos
- Você pode comparar dados dos arquivos com dados da plataforma
- Você conhece TODOS os usuários, máquinas, manutenções, estoque, obras, contas, etc.

INSTRUÇÕES IMPORTANTES:
1. SEMPRE responda em português brasileiro
2. Use QUEBRAS DE LINHA para separar parágrafos
3. Use listas com "•" para enumerar itens
4. Formate valores monetários como R$ 1.234,56
5. Se o usuário anexou arquivos, ANALISE O CONTEÚDO EXTRAÍDO e faça comentários úteis
6. Se for uma planilha ou documento, resuma as informações principais
7. Se for uma imagem, descreva o que sabe sobre ela
8. NÃO use markdown com asteriscos (**, *)
9. Relacione os dados dos arquivos com os dados da plataforma quando relevante
"""
        
        # Inicializar chat com Gemini
        llm_key = os.environ.get("EMERGENT_LLM_KEY")
        
        llm_chat = LlmChat(
            api_key=llm_key,
            session_id=f"chatbot-files-{current_user['id']}-{datetime.now().strftime('%Y%m%d%H%M%S')}",
            system_message=system_message
        ).with_model("gemini", "gemini-2.0-flash")
        
        # Mensagem do usuário
        user_text = message if message else "Analise os arquivos que anexei e me diga o que encontrou"
        if files_info:
            user_text += f"\n\n[Arquivos anexados: {', '.join([f['nome'] for f in files_info])}]"
        
        user_message = UserMessage(text=user_text)
        response = await llm_chat.send_message(user_message)
        
        return ChatResponse(
            response=response,
            context_used=["arquivos", "todos_os_modulos"]
        )
        
    except Exception as e:
        logging.error(f"Erro no chatbot com arquivos: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Erro ao processar arquivos: {str(e)}")


# ============ EXPORTAÇÃO (PDF/Excel/OFX/Recibo/Duplicata) ============
# (Endpoints /export/* extraídos para /app/backend/routes/exports_all.py
#  — Refactor Sessão 32 Fase 2 Parte 1)


ALLOWED_EXTENSIONS = {'.pdf', '.png', '.jpg', '.jpeg', '.gif', '.webp'}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

@api_router.post("/attachments/upload")
async def upload_attachment(
    file: UploadFile = File(...),
    entity_type: str = Form(...),  # contas_pagar, contas_receber, ordens_servico, etc
    entity_id: str = Form(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload de anexo para uma entidade (conta, OS, etc)"""
    
    # Validar extensão
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Tipo de arquivo não permitido. Permitidos: {', '.join(ALLOWED_EXTENSIONS)}")
    
    # Validar tamanho
    contents = await file.read()
    if len(contents) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="Arquivo muito grande. Máximo: 10MB")
    
    # Gerar nome único
    file_id = str(uuid.uuid4())
    filename = f"{file_id}{file_ext}"
    file_path = UPLOAD_DIR / filename
    
    # Salvar arquivo
    with open(file_path, "wb") as f:
        f.write(contents)
    
    # Inferir content-type se não veio no upload (garante canPreview no frontend)
    ext_mime = {
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".gif": "image/gif", ".webp": "image/webp",
    }
    file_type = file.content_type or ext_mime.get(file_ext) or "application/octet-stream"

    # Criar registro do anexo
    attachment = {
        "id": file_id,
        "filename": file.filename,
        "stored_filename": filename,
        "file_type": file_type,
        "file_size": len(contents),
        "entity_type": entity_type,
        "entity_id": entity_id,
        "uploaded_by": current_user["id"],
        "uploaded_by_name": current_user["name"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.attachments.insert_one(attachment)
    
    # Atualizar a entidade com referência ao anexo
    collection = db[entity_type]
    await collection.update_one(
        {"id": entity_id},
        {"$push": {"anexos": file_id}}
    )
    
    # Auditoria
    await create_audit_log(
        user=current_user,
        action="anexar",
        entity_type=f"arquivo em {entity_type}",
        entity_id=entity_id,
        entity_name=file.filename,
        details=f"Arquivo anexado: {file.filename}",
        module="Administrativo"
    )
    
    attachment.pop("_id", None)
    return {"message": "Arquivo enviado com sucesso", "attachment": attachment}

@api_router.get("/attachments/download/{file_id}")
async def download_attachment(file_id: str, current_user: dict = Depends(get_current_user)):
    """Download de um anexo"""
    attachment = await db.attachments.find_one({"id": file_id}, {"_id": 0})
    if not attachment:
        raise HTTPException(status_code=404, detail="Anexo não encontrado")
    
    file_path = UPLOAD_DIR / attachment["stored_filename"]
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
    
    # Infere content-type para anexos legados (file_type vazio)
    ext_mime = {
        ".pdf": "application/pdf",
        ".png": "image/png",
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".gif": "image/gif", ".webp": "image/webp",
    }
    stored = attachment.get("stored_filename", "")
    ext = Path(stored).suffix.lower()
    media_type = attachment.get("file_type") or ext_mime.get(ext) or "application/octet-stream"

    from fastapi.responses import FileResponse
    return FileResponse(
        path=str(file_path),
        filename=attachment["filename"],
        media_type=media_type,
    )

@api_router.get("/attachments/{entity_type}/{entity_id}")
async def get_attachments(entity_type: str, entity_id: str, current_user: dict = Depends(get_current_user)):
    """Lista anexos de uma entidade"""
    attachments = await db.attachments.find(
        {"entity_type": entity_type, "entity_id": entity_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    return attachments

@api_router.delete("/attachments/{file_id}")
async def delete_attachment(file_id: str, current_user: dict = Depends(get_current_user)):
    """Remove um anexo"""
    attachment = await db.attachments.find_one({"id": file_id}, {"_id": 0})
    if not attachment:
        raise HTTPException(status_code=404, detail="Anexo não encontrado")
    
    # Remover arquivo físico
    file_path = UPLOAD_DIR / attachment["stored_filename"]
    if file_path.exists():
        file_path.unlink()
    
    # Remover registro do banco
    await db.attachments.delete_one({"id": file_id})
    
    # Remover referência da entidade
    collection = db[attachment["entity_type"]]
    await collection.update_one(
        {"id": attachment["entity_id"]},
        {"$pull": {"anexos": file_id}}
    )
    
    # Auditoria
    await create_audit_log(
        user=current_user,
        action="remover anexo",
        entity_type=f"arquivo de {attachment['entity_type']}",
        entity_id=attachment["entity_id"],
        entity_name=attachment["filename"],
        details=f"Anexo removido: {attachment['filename']}",
        module="Administrativo"
    )
    
    return {"message": "Anexo removido com sucesso"}


# ============ TASK/MESSAGE SYSTEM ============

# Directory for task attachments (100MB max)
TASK_UPLOAD_DIR = ROOT_DIR / "task_uploads"
TASK_UPLOAD_DIR.mkdir(exist_ok=True)
UPLOAD_DIR = ROOT_DIR / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

class TaskCreate(BaseModel):
    target_system: str  # "gerenciamento" or "administrativo"
    priority: str  # "baixa", "media", "alta"
    title: str
    message: str

class TaskResponse(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str
    target_system: str
    priority: str
    title: str
    message: str
    attachments: List[dict] = []
    created_by_id: str
    created_by_name: str
    created_at: str
    read: bool = False
    read_at: Optional[str] = None
    read_by: Optional[str] = None

@api_router.post("/admin-panel/tasks")
async def create_task(
    task: TaskCreate,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Create a new task/message for a specific system"""
    current_user = await get_current_user(credentials)
    
    if current_user.get("role") not in ADMIN_ROLES:
        raise HTTPException(status_code=403, detail="Apenas administradores podem criar tarefas")
    
    if task.target_system not in ["gerenciamento", "administrativo", "rh"]:
        raise HTTPException(status_code=400, detail="Sistema alvo inválido")
    
    if task.priority not in ["baixa", "media", "alta"]:
        raise HTTPException(status_code=400, detail="Prioridade inválida")
    
    task_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    
    task_doc = {
        "id": task_id,
        "target_system": task.target_system,
        "priority": task.priority,
        "title": task.title,
        "message": task.message,
        "attachments": [],
        "created_by_id": current_user["id"],
        "created_by_name": current_user["name"],
        "created_at": now,
        "read": False,
        "read_at": None,
        "read_by": None
    }
    
    await db.tasks.insert_one(task_doc)
    
    # Auditoria
    await create_audit_log(
        user=current_user,
        action="criar tarefa",
        entity_type="tarefa",
        entity_id=task_id,
        entity_name=task.title,
        details=f"Tarefa criada para {task.target_system} - Prioridade: {task.priority}",
        module="Painel Admin"
    )
    
    task_doc.pop("_id", None)
    return task_doc

@api_router.post("/admin-panel/tasks/{task_id}/attachments")
async def upload_task_attachment(
    task_id: str,
    file: UploadFile = File(...),
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Upload attachment to a task (max 100MB)"""
    current_user = await get_current_user(credentials)
    
    if current_user.get("role") not in ADMIN_ROLES:
        raise HTTPException(status_code=403, detail="Apenas administradores podem adicionar anexos")
    
    task = await db.tasks.find_one({"id": task_id})
    if not task:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada")
    
    # Read file content
    content = await file.read()
    
    # Check file size (100MB max)
    max_size = 100 * 1024 * 1024  # 100MB
    if len(content) > max_size:
        raise HTTPException(status_code=400, detail="Arquivo muito grande. Máximo: 100MB")
    
    # Create task directory
    task_dir = TASK_UPLOAD_DIR / task_id
    task_dir.mkdir(exist_ok=True)
    
    # Generate unique filename
    ext = Path(file.filename).suffix if file.filename else ""
    unique_filename = f"{uuid.uuid4()}{ext}"
    file_path = task_dir / unique_filename
    
    # Save file
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Add attachment info to task
    attachment = {
        "id": str(uuid.uuid4()),
        "original_name": file.filename,
        "filename": unique_filename,
        "content_type": file.content_type,
        "size": len(content),
        "uploaded_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.tasks.update_one(
        {"id": task_id},
        {"$push": {"attachments": attachment}}
    )
    
    return {"message": "Anexo adicionado com sucesso", "attachment": attachment}

@api_router.get("/admin-panel/tasks")
async def list_all_tasks(
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """List all tasks (admin only)"""
    current_user = await get_current_user(credentials)
    
    if current_user.get("role") not in ADMIN_ROLES:
        raise HTTPException(status_code=403, detail="Apenas administradores podem ver todas as tarefas")
    
    tasks = await db.tasks.find().sort("created_at", -1).to_list(1000)
    
    for task in tasks:
        task.pop("_id", None)
    
    return tasks

@api_router.get("/tasks")
async def list_tasks_for_system(
    system: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """List tasks for a specific system"""
    current_user = await get_current_user(credentials)
    
    if system not in ["gerenciamento", "administrativo", "rh"]:
        raise HTTPException(status_code=400, detail="Sistema inválido")
    
    tasks = await db.tasks.find({"target_system": system}).sort("created_at", -1).to_list(1000)
    
    for task in tasks:
        task.pop("_id", None)
    
    return tasks

@api_router.get("/tasks/unread-count")
async def get_unread_task_count(
    system: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Get count of unread tasks for a system"""
    await get_current_user(credentials)
    
    if system not in ["gerenciamento", "administrativo", "rh"]:
        raise HTTPException(status_code=400, detail="Sistema inválido")
    
    count = await db.tasks.count_documents({"target_system": system, "read": False})
    
    return {"count": count}

@api_router.patch("/tasks/{task_id}/read")
async def mark_task_as_read(
    task_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Mark a task as read"""
    current_user = await get_current_user(credentials)
    
    task = await db.tasks.find_one({"id": task_id})
    if not task:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada")
    
    now = datetime.now(timezone.utc).isoformat()
    
    await db.tasks.update_one(
        {"id": task_id},
        {"$set": {"read": True, "read_at": now, "read_by": current_user["name"]}}
    )
    
    return {"message": "Tarefa marcada como lida"}

@api_router.get("/tasks/{task_id}/attachments/{filename}")
async def download_task_attachment(
    task_id: str,
    filename: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Download a task attachment"""
    await get_current_user(credentials)
    
    task = await db.tasks.find_one({"id": task_id})
    if not task:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada")
    
    # Find attachment
    attachment = None
    for att in task.get("attachments", []):
        if att["filename"] == filename:
            attachment = att
            break
    
    if not attachment:
        raise HTTPException(status_code=404, detail="Anexo não encontrado")
    
    file_path = TASK_UPLOAD_DIR / task_id / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
    
    from fastapi.responses import FileResponse
    return FileResponse(
        path=str(file_path),
        filename=attachment["original_name"],
        media_type=attachment.get("content_type", "application/octet-stream")
    )

@api_router.delete("/admin-panel/tasks/{task_id}")
async def delete_task(
    task_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Delete a task (admin only)"""
    current_user = await get_current_user(credentials)
    
    if current_user.get("role") not in ADMIN_ROLES:
        raise HTTPException(status_code=403, detail="Apenas administradores podem excluir tarefas")
    
    task = await db.tasks.find_one({"id": task_id})
    if not task:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada")
    
    # Delete attachments directory
    task_dir = TASK_UPLOAD_DIR / task_id
    if task_dir.exists():
        import shutil
        shutil.rmtree(task_dir)
    
    await db.tasks.delete_one({"id": task_id})
    
    # Auditoria
    await create_audit_log(
        user=current_user,
        action="excluir tarefa",
        entity_type="tarefa",
        entity_id=task_id,
        entity_name=task["title"],
        details=f"Tarefa excluída: {task['title']}",
        module="Painel Admin"
    )
    
    return {"message": "Tarefa excluída com sucesso"}


# ============ RH SYSTEM (RECURSOS HUMANOS) ============
# Movido para: /app/backend/routes/rh.py
# Incluído via: api_router.include_router(rh_router)

# ============ STORAGE SYSTEM (FILE MANAGER) ============

# Directory for storage system
STORAGE_DIR = ROOT_DIR / "storage"
STORAGE_DIR.mkdir(exist_ok=True)

class FolderCreate(BaseModel):
    name: str
    parent_path: str = "/"
    password: Optional[str] = None  # Senha opcional para proteger a pasta

class FolderPasswordCheck(BaseModel):
    path: str
    password: str

class FolderPasswordSet(BaseModel):
    path: str
    password: Optional[str] = None  # None para remover senha

class RenameItem(BaseModel):
    path: str
    new_name: str

@api_router.get("/storage/list")
async def list_storage_items(
    path: str = "/",
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """List files and folders in a path"""
    await get_current_user(credentials)
    
    # Normalize path
    if not path.startswith("/"):
        path = "/" + path
    
    # Build absolute path
    abs_path = STORAGE_DIR / path.lstrip("/")
    
    if not abs_path.exists():
        abs_path.mkdir(parents=True, exist_ok=True)
    
    # Buscar todas as pastas protegidas de uma vez
    protected_folders = await db.folder_passwords.find({}, {"_id": 0, "path": 1}).to_list(1000)
    protected_paths = {f["path"] for f in protected_folders}
    
    items = []
    try:
        for entry in abs_path.iterdir():
            rel_path = "/" + str(entry.relative_to(STORAGE_DIR)).replace("\\", "/")
            
            if entry.is_dir():
                # Count items inside folder
                items_count = len(list(entry.iterdir())) if entry.exists() else 0
                items.append({
                    "name": entry.name,
                    "type": "folder",
                    "path": rel_path,
                    "items_count": items_count,
                    "modified_at": datetime.fromtimestamp(entry.stat().st_mtime, tz=timezone.utc).isoformat(),
                    "has_password": rel_path in protected_paths
                })
            else:
                items.append({
                    "name": entry.name,
                    "type": "file",
                    "path": rel_path,
                    "size": entry.stat().st_size,
                    "modified_at": datetime.fromtimestamp(entry.stat().st_mtime, tz=timezone.utc).isoformat()
                })
    except Exception as e:
        logger.error(f"Error listing storage: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao listar arquivos: {str(e)}")
    
    # Sort: folders first, then files, both alphabetically
    items.sort(key=lambda x: (0 if x["type"] == "folder" else 1, x["name"].lower()))
    
    return items

@api_router.post("/storage/folder")
async def create_folder(
    data: FolderCreate,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Create a new folder"""
    current_user = await get_current_user(credentials)
    
    # Validate folder name
    if not data.name or "/" in data.name or "\\" in data.name:
        raise HTTPException(status_code=400, detail="Nome de pasta inválido")
    
    # Normalize parent path
    parent = data.parent_path if data.parent_path.startswith("/") else "/" + data.parent_path
    
    # Build full path
    full_path = STORAGE_DIR / parent.lstrip("/") / data.name
    folder_path = "/" + str(full_path.relative_to(STORAGE_DIR)).replace("\\", "/")
    
    if full_path.exists():
        raise HTTPException(status_code=400, detail="Pasta já existe")
    
    try:
        full_path.mkdir(parents=True, exist_ok=False)
        
        # Se tiver senha, salvar no MongoDB
        if data.password:
            password_hash = hash_password(data.password)
            await db.folder_passwords.update_one(
                {"path": folder_path},
                {"$set": {
                    "path": folder_path,
                    "password_hash": password_hash,
                    "created_by": current_user["id"],
                    "created_at": datetime.now(timezone.utc).isoformat()
                }},
                upsert=True
            )
        
        # Auditoria
        await create_audit_log(
            user=current_user,
            action="criar pasta",
            entity_type="storage",
            entity_id=data.name,
            entity_name=data.name,
            details=f"Pasta criada em {parent}" + (" (protegida com senha)" if data.password else ""),
            module="Armazenamento"
        )
        
        return {"message": "Pasta criada com sucesso", "path": folder_path, "has_password": bool(data.password)}
    except Exception as e:
        logger.error(f"Error creating folder: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao criar pasta: {str(e)}")

@api_router.post("/storage/folder/check-password")
async def check_folder_password(
    data: FolderPasswordCheck,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Verifica se a senha da pasta está correta"""
    await get_current_user(credentials)
    
    # Normalizar path
    path = data.path if data.path.startswith("/") else "/" + data.path
    
    # Buscar senha da pasta
    folder_record = await db.folder_passwords.find_one({"path": path}, {"_id": 0})
    
    if not folder_record:
        return {"valid": True, "message": "Pasta não possui senha"}
    
    # Verificar senha
    if verify_password(data.password, folder_record["password_hash"]):
        return {"valid": True, "message": "Senha correta"}
    else:
        raise HTTPException(status_code=401, detail="Senha incorreta")

@api_router.post("/storage/folder/set-password")
async def set_folder_password(
    data: FolderPasswordSet,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Define ou remove a senha de uma pasta"""
    current_user = await get_current_user(credentials)
    
    # Normalizar path
    path = data.path if data.path.startswith("/") else "/" + data.path
    
    # Verificar se a pasta existe
    abs_path = STORAGE_DIR / path.lstrip("/")
    if not abs_path.exists() or not abs_path.is_dir():
        raise HTTPException(status_code=404, detail="Pasta não encontrada")
    
    if data.password:
        # Definir nova senha
        password_hash = hash_password(data.password)
        await db.folder_passwords.update_one(
            {"path": path},
            {"$set": {
                "path": path,
                "password_hash": password_hash,
                "updated_by": current_user["id"],
                "updated_at": datetime.now(timezone.utc).isoformat()
            }},
            upsert=True
        )
        
        await create_audit_log(
            user=current_user,
            action="definir senha",
            entity_type="storage",
            entity_id=path,
            entity_name=path.split("/")[-1],
            details=f"Senha definida para pasta {path}",
            module="Armazenamento"
        )
        
        return {"message": "Senha definida com sucesso", "has_password": True}
    else:
        # Remover senha
        await db.folder_passwords.delete_one({"path": path})
        
        await create_audit_log(
            user=current_user,
            action="remover senha",
            entity_type="storage",
            entity_id=path,
            entity_name=path.split("/")[-1],
            details=f"Senha removida da pasta {path}",
            module="Armazenamento"
        )
        
        return {"message": "Senha removida com sucesso", "has_password": False}

@api_router.get("/storage/folder/has-password")
async def check_folder_has_password(
    path: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Verifica se uma pasta possui senha"""
    await get_current_user(credentials)
    
    # Normalizar path
    if not path.startswith("/"):
        path = "/" + path
    
    folder_record = await db.folder_passwords.find_one({"path": path}, {"_id": 0})
    return {"has_password": folder_record is not None}

@api_router.post("/storage/upload")
async def upload_storage_file(
    file: UploadFile = File(...),
    path: str = Form("/"),
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Upload a file to storage"""
    current_user = await get_current_user(credentials)
    
    # Normalize path
    if not path.startswith("/"):
        path = "/" + path
    
    # Build full path
    dir_path = STORAGE_DIR / path.lstrip("/")
    dir_path.mkdir(parents=True, exist_ok=True)
    
    # Read file content
    content = await file.read()
    
    # Save file with original name (handle duplicates)
    filename = file.filename or "arquivo"
    file_path = dir_path / filename
    
    # If file exists, add number suffix
    counter = 1
    base_name = Path(filename).stem
    ext = Path(filename).suffix
    while file_path.exists():
        filename = f"{base_name}_{counter}{ext}"
        file_path = dir_path / filename
        counter += 1
    
    try:
        with open(file_path, "wb") as f:
            f.write(content)
        
        # Auditoria
        await create_audit_log(
            user=current_user,
            action="upload arquivo",
            entity_type="storage",
            entity_id=filename,
            entity_name=filename,
            details=f"Arquivo enviado para {path}",
            module="Armazenamento"
        )
        
        rel_path = "/" + str(file_path.relative_to(STORAGE_DIR)).replace("\\", "/")
        return {
            "message": "Arquivo enviado com sucesso",
            "name": filename,
            "path": rel_path,
            "size": len(content)
        }
    except Exception as e:
        logger.error(f"Error uploading file: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao enviar arquivo: {str(e)}")

@api_router.get("/storage/download")
async def download_storage_file(
    path: str,
    token: Optional[str] = None,
    credentials: Optional[HTTPAuthorizationCredentials] = Depends(optional_security)
):
    """Download a file from storage"""
    # Support both Authorization header and query param token (for iframe/img src)
    if credentials and credentials.credentials:
        try:
            payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=["HS256"])
        except (jwt.ExpiredSignatureError, jwt.InvalidTokenError):
            raise HTTPException(status_code=401, detail="Token inválido")
    elif token:
        try:
            payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        except jwt.ExpiredSignatureError:
            raise HTTPException(status_code=401, detail="Token expirado")
        except jwt.InvalidTokenError:
            raise HTTPException(status_code=401, detail="Token inválido")
    else:
        raise HTTPException(status_code=401, detail="Not authenticated")
    
    # Normalize path
    if not path.startswith("/"):
        path = "/" + path
    
    # Build absolute path
    abs_path = STORAGE_DIR / path.lstrip("/")
    
    if not abs_path.exists() or not abs_path.is_file():
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
    
    # Determine content type
    content_type = "application/octet-stream"
    ext = abs_path.suffix.lower()
    content_types = {
        ".pdf": "application/pdf",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".svg": "image/svg+xml",
        ".txt": "text/plain",
        ".csv": "text/csv",
        ".json": "application/json",
        ".xml": "application/xml",
        ".mp4": "video/mp4",
        ".mp3": "audio/mpeg",
        ".wav": "audio/wav",
        ".zip": "application/zip"
    }
    content_type = content_types.get(ext, content_type)
    
    return FileResponse(
        path=str(abs_path),
        filename=abs_path.name,
        media_type=content_type
    )

@api_router.get("/storage/preview-office")
async def preview_office_file(
    path: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Preview Word/Excel files as HTML"""
    current_user = await get_current_user(credentials)
    
    # Normalize path
    if not path.startswith("/"):
        path = "/" + path
    
    abs_path = STORAGE_DIR / path.lstrip("/")
    
    if not abs_path.exists() or not abs_path.is_file():
        raise HTTPException(status_code=404, detail="Arquivo não encontrado")
    
    ext = abs_path.suffix.lower()
    
    try:
        if ext in ['.doc', '.docx']:
            # Preview Word document
            from docx import Document
            doc = Document(str(abs_path))
            
            html_content = '<div style="font-family: Arial, sans-serif; padding: 20px; max-width: 800px; margin: 0 auto;">'
            for para in doc.paragraphs:
                style = para.style.name if para.style else ''
                if 'Heading 1' in style:
                    html_content += f'<h1 style="color: #333; border-bottom: 2px solid #E31A1A; padding-bottom: 10px;">{para.text}</h1>'
                elif 'Heading 2' in style:
                    html_content += f'<h2 style="color: #444;">{para.text}</h2>'
                elif 'Heading 3' in style:
                    html_content += f'<h3 style="color: #555;">{para.text}</h3>'
                else:
                    html_content += f'<p style="line-height: 1.6; margin: 10px 0;">{para.text}</p>'
            
            # Add tables
            for table in doc.tables:
                html_content += '<table style="width: 100%; border-collapse: collapse; margin: 15px 0;">'
                for row in table.rows:
                    html_content += '<tr>'
                    for cell in row.cells:
                        html_content += f'<td style="border: 1px solid #ddd; padding: 8px;">{cell.text}</td>'
                    html_content += '</tr>'
                html_content += '</table>'
            
            html_content += '</div>'
            return {"html": html_content, "type": "word"}
            
        elif ext in ['.xls', '.xlsx']:
            # Preview Excel file
            import openpyxl
            wb = openpyxl.load_workbook(str(abs_path), data_only=True)
            
            html_content = '<div style="font-family: Arial, sans-serif; padding: 20px;">'
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                html_content += f'<h3 style="color: #E31A1A; margin-top: 20px;">Planilha: {sheet_name}</h3>'
                html_content += '<table style="width: 100%; border-collapse: collapse; margin-bottom: 20px;">'
                
                for row in sheet.iter_rows(max_row=100, max_col=20):  # Limit to prevent huge tables
                    html_content += '<tr>'
                    for cell in row:
                        value = cell.value if cell.value is not None else ''
                        style = 'border: 1px solid #ddd; padding: 8px; '
                        if cell.row == 1:
                            style += 'background: #f5f5f5; font-weight: bold;'
                        html_content += f'<td style="{style}">{value}</td>'
                    html_content += '</tr>'
                
                html_content += '</table>'
            
            html_content += '</div>'
            return {"html": html_content, "type": "excel"}
        else:
            raise HTTPException(status_code=400, detail="Tipo de arquivo não suportado para preview")
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao processar arquivo: {str(e)}")

# Trash directory for storage system
STORAGE_TRASH_DIR = ROOT_DIR / "storage_trash"
STORAGE_TRASH_DIR.mkdir(exist_ok=True)

class MoveRequest(BaseModel):
    source_path: str
    destination_path: str

@api_router.post("/storage/move")
async def move_storage_item(
    data: MoveRequest,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Move a file or folder to another location"""
    current_user = await get_current_user(credentials)
    import shutil
    
    # Normalize paths
    source = data.source_path if data.source_path.startswith("/") else "/" + data.source_path
    dest = data.destination_path if data.destination_path.startswith("/") else "/" + data.destination_path
    
    source_abs = STORAGE_DIR / source.lstrip("/")
    dest_dir = STORAGE_DIR / dest.lstrip("/")
    
    if not source_abs.exists():
        raise HTTPException(status_code=404, detail="Arquivo de origem não encontrado")
    
    if not dest_dir.exists() or not dest_dir.is_dir():
        raise HTTPException(status_code=400, detail="Pasta de destino não existe")
    
    # Check if destination is inside source (cannot move folder into itself)
    if source_abs.is_dir():
        try:
            dest_dir.relative_to(source_abs)
            raise HTTPException(status_code=400, detail="Não é possível mover uma pasta para dentro dela mesma")
        except ValueError:
            pass  # dest is not inside source, this is good
    
    dest_path = dest_dir / source_abs.name
    
    # Handle name conflicts
    if dest_path.exists():
        base = dest_path.stem
        ext = dest_path.suffix
        counter = 1
        while dest_path.exists():
            dest_path = dest_dir / f"{base} ({counter}){ext}"
            counter += 1
    
    try:
        shutil.move(str(source_abs), str(dest_path))
        return {"message": "Item movido com sucesso", "new_path": "/" + str(dest_path.relative_to(STORAGE_DIR))}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao mover: {str(e)}")

@api_router.post("/storage/copy")
async def copy_storage_item(
    data: MoveRequest,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Copy a file or folder to another location"""
    current_user = await get_current_user(credentials)
    import shutil
    
    # Normalize paths
    source = data.source_path if data.source_path.startswith("/") else "/" + data.source_path
    dest = data.destination_path if data.destination_path.startswith("/") else "/" + data.destination_path
    
    source_abs = STORAGE_DIR / source.lstrip("/")
    dest_dir = STORAGE_DIR / dest.lstrip("/")
    
    if not source_abs.exists():
        raise HTTPException(status_code=404, detail="Arquivo de origem não encontrado")
    
    if not dest_dir.exists() or not dest_dir.is_dir():
        raise HTTPException(status_code=400, detail="Pasta de destino não existe")
    
    dest_path = dest_dir / source_abs.name
    
    # Handle name conflicts
    if dest_path.exists():
        base = dest_path.stem
        ext = dest_path.suffix
        counter = 1
        while dest_path.exists():
            dest_path = dest_dir / f"{base} - Cópia ({counter}){ext}"
            counter += 1
    
    try:
        if source_abs.is_dir():
            shutil.copytree(str(source_abs), str(dest_path))
        else:
            shutil.copy2(str(source_abs), str(dest_path))
        return {"message": "Item copiado com sucesso", "new_path": "/" + str(dest_path.relative_to(STORAGE_DIR))}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao copiar: {str(e)}")

@api_router.delete("/storage/delete")
async def delete_storage_item(
    path: str,
    permanent: bool = False,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Move a file or folder to trash (or delete permanently if permanent=True)"""
    current_user = await get_current_user(credentials)
    
    # Normalize path
    if not path.startswith("/"):
        path = "/" + path
    
    # Build absolute path
    abs_path = STORAGE_DIR / path.lstrip("/")
    
    if not abs_path.exists():
        raise HTTPException(status_code=404, detail="Item não encontrado")
    
    item_name = abs_path.name
    is_folder = abs_path.is_dir()
    
    try:
        if permanent:
            # Delete permanently
            if is_folder:
                import shutil
                shutil.rmtree(abs_path)
            else:
                abs_path.unlink()
            action = "excluir permanentemente"
        else:
            # Move to trash
            import shutil
            trash_item_id = str(uuid.uuid4())
            trash_path = STORAGE_TRASH_DIR / trash_item_id
            
            # Move the item
            shutil.move(str(abs_path), str(trash_path))
            
            # Save metadata to database
            trash_record = {
                "id": trash_item_id,
                "original_name": item_name,
                "original_path": path,
                "type": "folder" if is_folder else "file",
                "deleted_by": current_user["id"],
                "deleted_by_name": current_user["name"],
                "deleted_at": datetime.now(timezone.utc).isoformat()
            }
            
            # Get file size if it's a file
            if not is_folder and trash_path.exists():
                trash_record["size"] = trash_path.stat().st_size
            
            await db.storage_trash.insert_one(trash_record)
            action = "mover para lixeira"
        
        # Auditoria
        await create_audit_log(
            user=current_user,
            action=action,
            entity_type="storage",
            entity_id=item_name,
            entity_name=item_name,
            details=f"{'Pasta' if is_folder else 'Arquivo'}: {path}",
            module="Armazenamento"
        )
        
        return {"message": "Item movido para lixeira" if not permanent else "Item excluído permanentemente"}
    except Exception as e:
        logger.error(f"Error deleting item: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao excluir: {str(e)}")

@api_router.get("/storage/trash")
async def list_trash_items(
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """List all items in trash"""
    await get_current_user(credentials)
    
    items = await db.storage_trash.find({}, {"_id": 0}).sort("deleted_at", -1).to_list(1000)
    return items

@api_router.post("/storage/trash/{item_id}/restore")
async def restore_trash_item(
    item_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Restore an item from trash"""
    current_user = await get_current_user(credentials)
    
    # Find item in trash
    trash_item = await db.storage_trash.find_one({"id": item_id}, {"_id": 0})
    if not trash_item:
        raise HTTPException(status_code=404, detail="Item não encontrado na lixeira")
    
    trash_path = STORAGE_TRASH_DIR / item_id
    if not trash_path.exists():
        # Clean up database record
        await db.storage_trash.delete_one({"id": item_id})
        raise HTTPException(status_code=404, detail="Arquivo não encontrado no sistema")
    
    # Restore to original location
    original_path = trash_item["original_path"]
    restore_path = STORAGE_DIR / original_path.lstrip("/")
    
    # Create parent directories if needed
    restore_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Handle name conflicts
    if restore_path.exists():
        base_name = Path(trash_item["original_name"]).stem
        ext = Path(trash_item["original_name"]).suffix
        counter = 1
        while restore_path.exists():
            new_name = f"{base_name}_{counter}{ext}"
            restore_path = restore_path.parent / new_name
            counter += 1
    
    try:
        import shutil
        shutil.move(str(trash_path), str(restore_path))
        
        # Remove from trash database
        await db.storage_trash.delete_one({"id": item_id})
        
        # Auditoria
        await create_audit_log(
            user=current_user,
            action="restaurar da lixeira",
            entity_type="storage",
            entity_id=trash_item["original_name"],
            entity_name=trash_item["original_name"],
            details=f"Restaurado para: {str(restore_path.relative_to(STORAGE_DIR))}",
            module="Armazenamento"
        )
        
        return {"message": "Item restaurado com sucesso", "restored_path": "/" + str(restore_path.relative_to(STORAGE_DIR)).replace("\\", "/")}
    except Exception as e:
        logger.error(f"Error restoring item: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao restaurar: {str(e)}")

@api_router.delete("/storage/trash/{item_id}")
async def delete_trash_item_permanently(
    item_id: str,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Permanently delete an item from trash"""
    current_user = await get_current_user(credentials)
    
    # Find item in trash
    trash_item = await db.storage_trash.find_one({"id": item_id}, {"_id": 0})
    if not trash_item:
        raise HTTPException(status_code=404, detail="Item não encontrado na lixeira")
    
    trash_path = STORAGE_TRASH_DIR / item_id
    
    try:
        if trash_path.exists():
            if trash_path.is_dir():
                import shutil
                shutil.rmtree(trash_path)
            else:
                trash_path.unlink()
        
        # Remove from trash database
        await db.storage_trash.delete_one({"id": item_id})
        
        # Auditoria
        await create_audit_log(
            user=current_user,
            action="excluir permanentemente",
            entity_type="storage",
            entity_id=trash_item["original_name"],
            entity_name=trash_item["original_name"],
            details=f"Excluído permanentemente da lixeira",
            module="Armazenamento"
        )
        
        return {"message": "Item excluído permanentemente"}
    except Exception as e:
        logger.error(f"Error deleting trash item: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao excluir: {str(e)}")

@api_router.delete("/storage/trash")
async def empty_trash(
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Empty all items from trash"""
    current_user = await get_current_user(credentials)
    
    try:
        import shutil
        
        # Get all items from trash
        trash_items = await db.storage_trash.find({}, {"_id": 0}).to_list(1000)
        count = len(trash_items)
        
        # Delete all files/folders in trash directory
        for item in trash_items:
            trash_path = STORAGE_TRASH_DIR / item["id"]
            if trash_path.exists():
                if trash_path.is_dir():
                    shutil.rmtree(trash_path)
                else:
                    trash_path.unlink()
        
        # Clear trash database
        await db.storage_trash.delete_many({})
        
        # Auditoria
        await create_audit_log(
            user=current_user,
            action="esvaziar lixeira",
            entity_type="storage",
            entity_id="trash",
            entity_name="Lixeira",
            details=f"{count} itens excluídos permanentemente",
            module="Armazenamento"
        )
        
        return {"message": f"Lixeira esvaziada. {count} itens excluídos."}
    except Exception as e:
        logger.error(f"Error emptying trash: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao esvaziar lixeira: {str(e)}")

@api_router.patch("/storage/rename")
async def rename_storage_item(
    data: RenameItem,
    credentials: HTTPAuthorizationCredentials = Depends(security)
):
    """Rename a file or folder in storage"""
    current_user = await get_current_user(credentials)
    
    # Validate new name
    if not data.new_name or "/" in data.new_name or "\\" in data.new_name:
        raise HTTPException(status_code=400, detail="Nome inválido")
    
    # Normalize path
    path = data.path if data.path.startswith("/") else "/" + data.path
    
    # Build absolute paths
    abs_path = STORAGE_DIR / path.lstrip("/")
    new_path = abs_path.parent / data.new_name
    
    if not abs_path.exists():
        raise HTTPException(status_code=404, detail="Item não encontrado")
    
    if new_path.exists():
        raise HTTPException(status_code=400, detail="Já existe um item com esse nome")
    
    old_name = abs_path.name
    
    try:
        abs_path.rename(new_path)
        
        # Auditoria
        await create_audit_log(
            user=current_user,
            action="renomear",
            entity_type="storage",
            entity_id=data.new_name,
            entity_name=data.new_name,
            details=f"Renomeado de '{old_name}' para '{data.new_name}'",
            module="Armazenamento"
        )
        
        return {
            "message": "Item renomeado com sucesso",
            "new_path": "/" + str(new_path.relative_to(STORAGE_DIR)).replace("\\", "/")
        }
    except Exception as e:
        logger.error(f"Error renaming item: {e}")
        raise HTTPException(status_code=500, detail=f"Erro ao renomear: {str(e)}")


# ============ NFE IMPORT ROUTES ============
# (Endpoints de certificados/importadas/downloads extraídos para
#  /app/backend/routes/nfe.py — Refactor Sessão 32)
#  Ainda presentes abaixo: POST /nfe/importar/{id} e POST /nfse/importar/{id}

# ============ IMPORTAÇÃO NF-e / NFS-e (SEFAZ + ABRASF) ============
# (Endpoints POST /nfe/importar/{id} e POST /nfse/importar/{id} extraídos
#  para /app/backend/routes/importacao_nf.py — Refactor Sessão 32 Parte 2)




# ==================== NFS-e ====================
# (Endpoints extraídos para /app/backend/routes/nfse.py — Refactor Sessão 32)

# ==================== EMISSÃO DE NF-e / NFS-e ====================
# (Endpoints /nfe/emitir, /nfse/emitir, /notas-emitidas/*, /nfe/cfops, /nfse/codigos-servico
#  extraídos para /app/backend/routes/emissao_nf.py — Refactor Sessão 32 Parte 2)

# ==================== CONCILIAÇÃO BANCÁRIA ====================
# (Endpoints extraídos para /app/backend/routes/conciliacao.py — Refactor Sessão 32)

# ==================== NF-e DOWNLOADS ====================
# (Endpoints /nfe/importadas/{id}/download-xml e /download-pdf extraídos para
#  /app/backend/routes/nfe.py — Refactor Sessão 32)

# Include modular routers first
api_router.include_router(rh_router)
api_router.include_router(admin_router)
api_router.include_router(chatbot_router)
api_router.include_router(folha_importacao_router)
api_router.include_router(fin_folha_solicitacoes_router)
api_router.include_router(storage_router)
api_router.include_router(export_router)
api_router.include_router(stock_router)
api_router.include_router(obras_router)
api_router.include_router(conciliacao_router)
api_router.include_router(nfse_router)
api_router.include_router(nfe_router)
api_router.include_router(financeiro_router)
api_router.include_router(emissao_router)
api_router.include_router(importacao_router)
api_router.include_router(exports_all_router)
api_router.include_router(dashboard_router)
api_router.include_router(medicoes_router)


# ===== Endpoints de Importação Automática (devem estar antes do include_router) =====
# Nota: As funções scheduler e importação estão no final do arquivo

# Include the router in the main app
app.include_router(api_router)

# Serve uploaded files
app.mount("/uploads", StaticFiles(directory=str(UPLOAD_DIR)), name="uploads")

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ===== SISTEMA DE IMPORTAÇÃO AUTOMÁTICA DE NOTAS =====
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger

scheduler = AsyncIOScheduler(timezone="America/Sao_Paulo")


async def importacao_automatica_notas():
    """Importa automaticamente NF-e e NFS-e de todos os CNPJs cadastrados"""
    logger.info("=" * 50)
    logger.info("INICIANDO IMPORTAÇÃO AUTOMÁTICA DE NOTAS FISCAIS")
    logger.info("=" * 50)
    
    try:
        # Buscar todos os certificados ativos
        certificados = await db.nfe_certificados.find({"ativo": True}).to_list(100)
        logger.info(f"Encontrados {len(certificados)} certificados ativos para importação")
        
        total_nfe_importadas = 0
        total_nfse_importadas = 0
        erros = []
        
        for cert in certificados:
            cert_id = cert.get("id")
            cnpj = cert.get("cnpj", "")
            razao = cert.get("razao_social", cnpj)
            
            logger.info(f"Processando certificado: {razao} (CNPJ: {cnpj})")
            
            # Importar NF-e
            try:
                resultado_nfe = await importar_nfe_automatico(cert_id)
                total_nfe_importadas += resultado_nfe.get("novas", 0)
                logger.info(f"  NF-e: {resultado_nfe.get('novas', 0)} novas importadas")
            except Exception as e:
                erro_msg = f"Erro NF-e {razao}: {str(e)}"
                erros.append(erro_msg)
                logger.error(erro_msg)
            
            # Importar NFS-e (se configurado para a cidade)
            try:
                resultado_nfse = await importar_nfse_automatico(cert_id)
                total_nfse_importadas += resultado_nfse.get("novas", 0)
                logger.info(f"  NFS-e: {resultado_nfse.get('novas', 0)} novas importadas")
            except Exception as e:
                erro_msg = f"Erro NFS-e {razao}: {str(e)}"
                erros.append(erro_msg)
                logger.warning(erro_msg)  # Warning porque nem todos CNPJs têm NFS-e
        
        # Criar registro de log de importação
        log_importacao = {
            "id": str(uuid.uuid4()),
            "tipo": "importacao_automatica",
            "data_hora": datetime.now(timezone.utc).isoformat(),
            "total_certificados": len(certificados),
            "nfe_importadas": total_nfe_importadas,
            "nfse_importadas": total_nfse_importadas,
            "erros": erros,
            "status": "concluido" if not erros else "concluido_com_erros"
        }
        await db.logs_importacao.insert_one(log_importacao)
        
        logger.info("=" * 50)
        logger.info(f"IMPORTAÇÃO AUTOMÁTICA CONCLUÍDA")
        logger.info(f"NF-e importadas: {total_nfe_importadas}")
        logger.info(f"NFS-e importadas: {total_nfse_importadas}")
        logger.info(f"Erros: {len(erros)}")
        logger.info("=" * 50)
        
    except Exception as e:
        logger.error(f"ERRO CRÍTICO na importação automática: {str(e)}")


async def importar_nfe_automatico(certificado_id: str):
    """Importa NF-e automaticamente (versão interna sem auth)"""
    certificado = await db.nfe_certificados.find_one({"id": certificado_id})
    if not certificado or not certificado.get("ativo"):
        return {"novas": 0, "erro": "Certificado inválido ou inativo"}
    
    # Verificar se está bloqueado
    bloqueado_ate = certificado.get("bloqueado_ate")
    if bloqueado_ate:
        bloqueio_dt = datetime.fromisoformat(bloqueado_ate.replace('Z', '+00:00')) if isinstance(bloqueado_ate, str) else bloqueado_ate
        if datetime.now(timezone.utc) < bloqueio_dt:
            return {"novas": 0, "erro": "Certificado bloqueado"}
    
    novas_importadas = 0
    ultimo_nsu_processado = certificado.get("ultimo_nsu", "000000000000000")
    
    try:
        import tempfile
        import gzip
        from xml.etree import ElementTree as ET
        
        cert_data = base64.b64decode(certificado["certificado_base64"])
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pfx') as cert_file:
            cert_file.write(cert_data)
            cert_path = cert_file.name
        
        ns = {
            'nfe': 'http://www.portalfiscal.inf.br/nfe',
            'res': 'http://www.portalfiscal.inf.br/nfe/wsdl/NFeDistribuicaoDFe'
        }
        
        # Usar PyNFe para consultar (mesma abordagem da importação manual)
        from pynfe.processamento.comunicacao import ComunicacaoSefaz
        
        con = ComunicacaoSefaz(
            uf=certificado.get("uf", "TO"),
            certificado=cert_path,
            certificado_senha=certificado.get("senha_certificado", ""),
            homologacao=(certificado.get("ambiente", "producao") == "homologacao")
        )
        
        # Consultar distribuição DFe
        try:
            resposta = con.consulta_distribuicao(
                cnpj=certificado["cnpj"].replace(".", "").replace("/", "").replace("-", ""),
                nsu=int(ultimo_nsu_processado) if ultimo_nsu_processado.isdigit() else 0
            )
            
            if resposta:
                # Processar resposta (pode ser string XML, bytes, ou objeto)
                if hasattr(resposta, 'text'):
                    xml_resposta = resposta.text
                elif hasattr(resposta, 'content'):
                    xml_resposta = resposta.content.decode('utf-8')
                elif isinstance(resposta, bytes):
                    xml_resposta = resposta.decode('utf-8')
                elif isinstance(resposta, str):
                    xml_resposta = resposta
                else:
                    xml_resposta = str(resposta)
                
                logger.info(f"Resposta SEFAZ recebida, tamanho: {len(xml_resposta)} bytes")
                
                root = ET.fromstring(xml_resposta)
                
                cStat = root.findtext('.//cStat') or root.findtext('.//{http://www.portalfiscal.inf.br/nfe}cStat')
                
                if cStat == '138':  # Documento localizado
                    docs = root.findall('.//docZip') or root.findall('.//{http://www.portalfiscal.inf.br/nfe}docZip')
                    
                    for doc in docs:
                        nsu = doc.get('NSU', '')
                        
                        try:
                            content_b64 = doc.text
                            if content_b64:
                                content_gzip = base64.b64decode(content_b64)
                                xml_content = gzip.decompress(content_gzip).decode('utf-8')
                                
                                # Processar XML da NF-e
                                nfe_root = ET.fromstring(xml_content)
                                res_nfe = nfe_root.find('.//nfe:resNFe', ns) or nfe_root.find('.//resNFe')
                                
                                if res_nfe is not None:
                                    chave = res_nfe.findtext('.//nfe:chNFe', '', ns) or res_nfe.findtext('.//chNFe', '')
                                    
                                    # Verificar se já existe
                                    existing = await db.nfe_importadas.find_one({"chave_acesso": chave})
                                    if existing:
                                        continue
                                    
                                    cnpj_emit = res_nfe.findtext('.//nfe:CNPJ', '', ns) or res_nfe.findtext('.//CNPJ', '')
                                    razao_emit = res_nfe.findtext('.//nfe:xNome', '', ns) or res_nfe.findtext('.//xNome', '')
                                    valor = res_nfe.findtext('.//nfe:vNF', '0', ns) or res_nfe.findtext('.//vNF', '0')
                                    data_emissao = res_nfe.findtext('.//nfe:dhEmi', '', ns) or res_nfe.findtext('.//dhEmi', '')
                                    
                                    numero_nf = chave[25:34].lstrip('0') if len(chave) >= 34 else ""
                                    serie = chave[22:25].lstrip('0') if len(chave) >= 25 else "1"
                                    
                                    doc_nfe = {
                                        "id": str(uuid.uuid4()),
                                        "certificado_id": certificado_id,
                                        "cnpj_destinatario": certificado["cnpj"],
                                        "chave_acesso": chave,
                                        "numero_nf": numero_nf or "N/A",
                                        "serie": serie or "1",
                                        "data_emissao": data_emissao[:10] if data_emissao else datetime.now(timezone.utc).isoformat()[:10],
                                        "cnpj_emitente": cnpj_emit,
                                        "razao_social_emitente": razao_emit or "Emitente não identificado",
                                        "valor_total": float(valor) if valor else 0.0,
                                        "itens": [],
                                        "xml_base64": base64.b64encode(xml_content.encode()).decode(),
                                        "nsu": nsu,
                                        "status": "nova",
                                        "importacao_automatica": True,
                                        "created_at": datetime.now(timezone.utc).isoformat()
                                    }
                                    await db.nfe_importadas.insert_one(doc_nfe)
                                    novas_importadas += 1
                                    
                                    if nsu > ultimo_nsu_processado:
                                        ultimo_nsu_processado = nsu
                        except Exception as doc_error:
                            logger.warning(f"Erro ao processar documento NSU {nsu}: {doc_error}")
                    
                    # Atualizar último NSU
                    ultNSU = root.findtext('.//ultNSU') or root.findtext('.//{http://www.portalfiscal.inf.br/nfe}ultNSU')
                    if ultNSU:
                        ultimo_nsu_processado = ultNSU
                else:
                    logger.info(f"SEFAZ retornou cStat={cStat} - Nenhum documento novo")
            
            # Atualizar último NSU no certificado
            await db.nfe_certificados.update_one(
                {"id": certificado_id},
                {"$set": {"ultimo_nsu": ultimo_nsu_processado, "ultima_consulta_auto": datetime.now(timezone.utc).isoformat()}}
            )
            
        except Exception as e:
            logger.warning(f"Erro ao consultar SEFAZ: {str(e)}")
        
        # Limpar arquivo temporário
        try:
            os.unlink(cert_path)
        except:
            pass
        
    except Exception as e:
        logger.error(f"Erro na importação automática NF-e: {str(e)}")
        return {"novas": 0, "erro": str(e)}
    
    return {"novas": novas_importadas}


async def importar_nfse_automatico(certificado_id: str):
    """Importa NFS-e automaticamente usando o webservice municipal configurado (ABRASF)"""
    certificado = await db.nfe_certificados.find_one({"id": certificado_id})
    if not certificado or not certificado.get("ativo"):
        return {"novas": 0, "erro": "Certificado inválido ou inativo"}

    url_nfse = (certificado.get("url_nfse") or "").strip()
    if not url_nfse:
        logger.info(f"NFS-e ignorada para {certificado.get('cnpj', '?')}: URL do webservice não configurada")
        return {"novas": 0}

    try:
        resultado = await importar_nfse(certificado_id, current_user={"name": "scheduler", "role": "admin"})
        novas = resultado.get("novas_nfses", 0) if isinstance(resultado, dict) else 0
        logger.info(f"NFS-e importadas automaticamente para {certificado.get('cnpj', '?')}: {novas}")
        return {"novas": novas}
    except Exception as e:
        logger.error(f"Erro na importação automática NFS-e: {str(e)}")
        return {"novas": 0, "erro": str(e)}


# Endpoints para gerenciar importação automática (usando app diretamente pois está após include_router)
@app.post("/api/nf/importacao-automatica/executar")
async def executar_importacao_automatica(current_user: dict = Depends(get_current_user)):
    """Executa a importação automática manualmente"""
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Apenas administradores podem executar importação automática")
    
    import asyncio
    asyncio.create_task(importacao_automatica_notas())
    
    return {"message": "Importação automática iniciada em background", "status": "processing"}


@app.get("/api/nf/importacao-automatica/status")
async def status_importacao_automatica(current_user: dict = Depends(get_current_user)):
    """Retorna o status da última importação automática"""
    ultimo_log = await db.logs_importacao.find_one(
        {"tipo": "importacao_automatica"},
        sort=[("data_hora", -1)]
    )
    
    if not ultimo_log:
        return {"ultimo_log": None, "scheduler_ativo": scheduler.running}
    
    del ultimo_log["_id"]
    return {
        "ultimo_log": ultimo_log,
        "scheduler_ativo": scheduler.running,
        "proximo_agendamento": "22:00 (horário de Brasília)"
    }


@app.get("/api/nf/importacao-automatica/logs")
async def logs_importacao_automatica(limit: int = 10, current_user: dict = Depends(get_current_user)):
    """Retorna os últimos logs de importação automática"""
    logs = await db.logs_importacao.find(
        {"tipo": "importacao_automatica"},
        {"_id": 0}
    ).sort("data_hora", -1).limit(limit).to_list(limit)
    
    return logs


@app.on_event("startup")
async def startup_event():
    """Inicializa o scheduler de importação automática"""
    logger.info("Iniciando scheduler de importação automática de notas...")
    
    # Inicializar object storage (Emergent)
    try:
        from utils.storage import init_storage
        init_storage()
        logger.info("Object storage inicializado (Emergent)")
    except Exception as e:
        logger.warning(f"Falha ao inicializar object storage: {e}. Uploads ficarão indisponíveis até reinicializar.")
    
    # Bootstrap da Base de Conhecimento do Chat IA do RH
    # Idempotente: só baixa/insere o que ainda não existe na coleção.
    try:
        import asyncio as _asyncio
        from routes.chatbot import bootstrap_knowledge_base

        async def _run_kb_bootstrap():
            try:
                summary = await bootstrap_knowledge_base()
                logger.info(
                    f"[KB] Bootstrap concluído — adicionados: {summary.get('added', [])}, "
                    f"já existentes: {summary.get('already_present', [])}, "
                    f"erros: {summary.get('errors', [])}"
                )
            except Exception as e:
                logger.error(f"[KB] Falha no bootstrap em background: {e}")

        # Roda em background para não bloquear o startup
        _asyncio.create_task(_run_kb_bootstrap())
        logger.info("Bootstrap da Base de Conhecimento do Chat IA agendado em background")
    except Exception as e:
        logger.warning(f"Não foi possível agendar bootstrap KB: {e}")
    
    # Agendar importação diária às 22:00 (horário de Brasília)
    scheduler.add_job(
        importacao_automatica_notas,
        CronTrigger(hour=22, minute=0, timezone="America/Sao_Paulo"),
        id="importacao_automatica_22h",
        replace_existing=True,
        name="Importação Automática de Notas às 22h"
    )
    
    scheduler.start()
    logger.info("Scheduler iniciado - Importação automática agendada para 22:00 (Brasília)")


@app.on_event("shutdown")
async def shutdown_scheduler():
    """Desliga o scheduler"""
    scheduler.shutdown(wait=False)
    logger.info("Scheduler de importação automática encerrado")
